import os
import sys
import re
import json
import hmac
import base64
import csv
import copy
import shutil
import hashlib
import queue
import threading
import time
try:
    import requests  # type: ignore
except ImportError:
    requests = None

import subprocess
from datetime import datetime, timedelta
from time import perf_counter

# Rutas base disponibles antes de cargar cualquier configuración.
APP_DIR = os.path.dirname(sys.executable) if getattr(sys, "frozen", False) else os.path.dirname(os.path.abspath(__file__))
BUNDLE_DIR = getattr(sys, "_MEIPASS", APP_DIR)

# 1. IMPORTAMOS DOTENV PARA SEGURIDAD (Busca archivo .env oculto)
try:
    from dotenv import load_dotenv  # type: ignore
    load_dotenv(os.path.join(APP_DIR, ".env"))
except ImportError:
    pass

try:
    import psycopg2  # type: ignore
    import psycopg2.extras  # type: ignore
    from psycopg2 import pool  # type: ignore
except ImportError:
    psycopg2 = None
    pool = None

from PySide6.QtCore import Qt, QSize, QDate, QTimer, QObject, QEvent, Signal, Slot, QThread, QRectF, QPointF  # type: ignore
from PySide6.QtGui import QPixmap, QColor, QAction, QCursor, QKeySequence, QShortcut, QPainter, QBrush, QPen, QTextDocument, QTextCursor, QIcon, QPageLayout  # type: ignore
from PySide6.QtPdf import QPdfDocument  # type: ignore
from PySide6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog  # type: ignore
from PySide6.QtWidgets import (  # type: ignore
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QComboBox, QPushButton, QTabWidget, QTabBar, QListWidget,
    QListWidgetItem, QSpinBox, QDoubleSpinBox, QGroupBox, QMessageBox,
    QSplitter, QFormLayout, QDialog, QDialogButtonBox, QToolButton, QStyle,
    QDateEdit, QAbstractSpinBox, QFileDialog, QTableWidget, QTableWidgetItem,
    QAbstractItemView, QCompleter, QMenu, QHeaderView, QSizePolicy, QInputDialog,
    QGridLayout, QCheckBox, QScrollArea, QToolTip, QRadioButton
)

from reportlab.pdfgen import canvas as rl_canvas  # type: ignore
from reportlab.lib.pagesizes import letter  # type: ignore
from reportlab.lib.units import inch  # type: ignore
from reportlab.platypus import Table, TableStyle, Paragraph, Frame, Spacer  # type: ignore
from reportlab.lib.styles import ParagraphStyle  # type: ignore
from reportlab.lib.colors import black, white, gray, blue, green, purple, orange, HexColor, Color  # type: ignore
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT  # type: ignore

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

from report_engine import PanelDataService, ReportHTMLRenderer, export_panel_xlsx


APP_TITLE = "Sistema de Facturación Médica"
VERSION = "2.6.2"

# En modo ejecutable PyInstaller, los recursos viven temporalmente en _MEI,
# mientras que reportes, recibos y logs se guardan junto al ejecutable.
BASE_DIR = APP_DIR
REPORTS_DIR = os.path.join(BASE_DIR, "reportes")
PDFS_DIR = os.path.join(BASE_DIR, "recibos")
SALA_STEP = 100.0
IDLE_TIMEOUT_MINUTES = 30 
WARNING_BEFORE_TIMEOUT_MINUTES = 1

# =========================================================
# CONFIGURACIÓN DE BASE DE DATOS EN LA NUBE (Segura con .env)
# =========================================================
try:
    from config_local import DATABASE_URL as LOCAL_DATABASE_URL  # type: ignore
except ImportError:
    LOCAL_DATABASE_URL = ""

DB_URL = os.environ.get("DATABASE_URL") or LOCAL_DATABASE_URL

LOGO_CANDIDATES = [
    os.path.join(BASE_DIR, "logo.jpg"),
    os.path.join(BASE_DIR, "logo.png"),
    os.path.join(BASE_DIR, "assets", "logo.jpg"),
    os.path.join(BASE_DIR, "assets", "logo.png"),
    os.path.join(BUNDLE_DIR, "logo.jpg"),
    os.path.join(BUNDLE_DIR, "logo.png"),
    os.path.join(BUNDLE_DIR, "assets", "logo.jpg"),
    os.path.join(BUNDLE_DIR, "assets", "logo.png"),
]
LOGO_PATH = next((p for p in LOGO_CANDIDATES if os.path.exists(p)), None)
# ---> RUTA DEL ÍCONO <---
ICON_CANDIDATES = [
    os.path.join(BASE_DIR, "favicon.ico"),
    os.path.join(BASE_DIR, "assets", "favicon.ico"),
    os.path.join(BUNDLE_DIR, "favicon.ico"),
    os.path.join(BUNDLE_DIR, "assets", "favicon.ico"),
]
ICON_PATH = next((p for p in ICON_CANDIDATES if os.path.exists(p)), None)


def write_runtime_log(message: str) -> None:
    """Registra errores recuperables que no se ven en un EXE sin consola."""
    try:
        log_path = os.path.join(APP_DIR, "app_error.log")
        with open(log_path, "a", encoding="utf-8") as log_file:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            log_file.write(f"[{timestamp}] {message}\n")
    except Exception:
        pass


def write_pdf_performance(message: str) -> None:
    """Guarda mediciones del flujo PDF para detectar regresiones reales."""
    try:
        log_path = os.path.join(APP_DIR, "pdf_performance.log")
        with open(log_path, "a", encoding="utf-8") as log_file:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            log_file.write(f"[{timestamp}] {message}\n")
    except Exception:
        pass

UNIVERSAL_CATEGORIES = ["Medicamentos", "Materiales"]
ARS_CATEGORIES = ["Laboratorios", "Imágenes", "Procedimientos", "Honorarios"]
ALL_CATEGORIES = UNIVERSAL_CATEGORIES + ARS_CATEGORIES

# ---> DEFINICIÓN DE ROLES <---
ROLE_AUX = "auxiliar"
ROLE_ADMIN = "administrador"
ROLE_AUDIT = "facturador de auditoria"

CAT_EMOJIS = {
    "Medicamentos": "💊",
    "Materiales": "📦",
    "Laboratorios": "🩸",
    "Imágenes": "📷",
    "Procedimientos": "🩺",
    "Honorarios": "👨‍⚕️"
}

CAT_COLORS = {
    "Medicamentos": "#0277bd",      
    "Materiales": "#e65100",        
    "Laboratorios": "#c62828",      
    "Imágenes": "#6a1b9a",          
    "Procedimientos": "#00695c",    
    "Honorarios": "#e64a19"         
}

# =========================================================
# GENERADOR DE TEMA DINÁMICO (Claro / Oscuro)
# =========================================================
def get_stylesheet(is_dark=False):
    if is_dark:
        bg = "#121212"
        text = "#e0e0e0"
        input_bg = "#1e1e1e"
        border = "#333333"
        alt_bg = "#1a1a1a"
        sel_bg = "#1565c0"
        sel_text = "#ffffff"
        title_color = "#90caf9"
        total_bg = "#1b5e20"
        total_text = "#ffffff"
        header_bg = "#002171"
        nav_bg = "#0d47a1"
    else:
        bg = "#f4f6f8"
        text = "#333333"
        input_bg = "#ffffff"
        border = "#d1d9e6"
        alt_bg = "#f9fbfd"
        sel_bg = "#e3f2fd"
        sel_text = "#000000"
        title_color = "#1565c0"
        total_bg = "#e8f5e9"
        total_text = "#1b5e20"
        header_bg = "#1565c0"
        nav_bg = "#0d47a1"
        
    return f"""
    QMainWindow, QDialog {{ background-color: {bg}; color: {text}; }}
    QLabel {{ color: {text}; }}
    QCheckBox, QRadioButton {{ color: {text}; }}
    QRadioButton:disabled {{ color: {'#7F8C98' if is_dark else '#7A8794'}; }}
    QLineEdit, QComboBox, QSpinBox, QDoubleSpinBox, QDateEdit {{
        background: {input_bg}; border: 1px solid {border}; border-radius: 6px; padding: 6px; color: {text};
    }}
    QListWidget {{
        background: {input_bg}; border: 1px solid {border}; border-radius: 6px; padding: 6px;
        alternate-background-color: {alt_bg}; font-size: 10pt; color: {text};
    }}
    QListWidget::item {{ padding: 8px; border-bottom: 1px solid {border}; }}
    QListWidget::item:selected {{ background-color: {sel_bg}; color: {sel_text}; border-radius: 4px; }}
    
    QTableWidget {{
        background: {input_bg}; border: 1px solid {border}; border-radius: 6px; font-size: 10pt; color: {text};
        alternate-background-color: {alt_bg}; gridline-color: transparent;
    }}
    QTableWidget::item {{ padding: 5px; }}
    QTableWidget::item:selected {{ background-color: {sel_bg}; color: {sel_text}; }}
    
    QGroupBox {{ 
        background-color: {input_bg}; 
        border: 1px solid {border}; 
        border-radius: 8px; 
        margin-top: 16px;
        padding: 18px 10px 10px 10px;
        font-weight: bold; 
        color: {title_color};
    }}
    QGroupBox::title {{ 
        subcontrol-origin: margin; 
        subcontrol-position: top left; 
        left: 15px;
        padding: 0 6px; 
        background-color: {input_bg};
    }}
    
    QTabWidget::pane {{ border: 1px solid {border}; border-radius: 6px; background: {input_bg};}}
    QTabBar::tab {{ padding: 8px 12px; background: {alt_bg}; border: 1px solid {border}; border-bottom: none; border-top-left-radius: 6px; border-top-right-radius: 6px; color: {text};}}
    QTabBar::tab:selected {{ background: {input_bg}; font-weight: bold; color: {title_color};}}

    QTabWidget#DashboardResultsTabs::pane {{
        border: 1px solid {border}; border-radius: 7px; background: {input_bg};
    }}
    QTabWidget#DashboardResultsTabs QTabBar::tab {{
        min-width: 115px; padding: 9px 14px; font-weight: 800;
        background: {alt_bg}; color: {text}; border: 1px solid {border};
        border-bottom: 2px solid {border};
    }}
    QTabWidget#DashboardResultsTabs QTabBar::tab:hover {{
        background: {sel_bg}; color: {sel_text};
    }}
    QTabWidget#DashboardResultsTabs QTabBar::tab:selected {{
        background: {input_bg}; color: {title_color};
        border-bottom: 3px solid {title_color};
    }}

    QLabel#FilterDialogTitle {{ color: {title_color}; }}
    QPushButton#FilterSecondaryButton {{
        background: {alt_bg}; color: {text}; border: 1px solid {border};
        border-radius: 6px; padding: 8px 12px; font-weight: 600;
    }}
    QPushButton#FilterSecondaryButton:hover {{ background: {sel_bg}; color: {sel_text}; }}
    
    QToolButton {{ border: 1px solid {border}; border-radius: 6px; background: {input_bg}; padding: 4px; color: {text}; }}
    QToolButton:hover {{ background: {alt_bg}; }}
    
    QMenu {{ background-color: {input_bg}; border: 1px solid {border}; font-size: 11pt; color: {text}; }}
    QMenu::item {{ padding: 8px 24px; }}
    QMenu::item:selected {{ background-color: {sel_bg}; color: {sel_text}; }}
    
    QHeaderView::section {{ background-color: {bg}; font-weight: bold; padding: 5px; border: 1px solid {border}; color: {text}; }}

    #HeaderWidget {{ background-color: {header_bg}; }}
    #NavWidget {{ background-color: {nav_bg}; padding: 5px; }}
    #HeaderWidget QLabel {{ color: #ffffff; }}
    
    #TotalLabel {{
        font-size: 22pt; font-weight: 900; color: {total_text};
        background-color: {total_bg}; padding: 10px 20px;
        border-radius: 8px; border: 2px solid {total_text};
    }}
    #BottomBar {{ background-color: {input_bg}; border-top: 1px solid {border}; }}
    
    /* ---> ESTILOS MEJORADOS PARA EL CALENDARIO <--- */
    QCalendarWidget QWidget#qt_calendar_navigationbar {{ 
        background-color: {header_bg}; 
    }}
    QCalendarWidget QWidget#qt_calendar_navigationbar QToolButton {{ 
        color: #ffffff; 
        background-color: transparent; 
        font-weight: bold; 
        font-size: 11pt; 
        border: none; 
        padding: 4px; 
    }}
    QCalendarWidget QWidget#qt_calendar_navigationbar QToolButton:hover {{ 
        background-color: rgba(255, 255, 255, 0.2); 
        border-radius: 4px; 
    }}
    QCalendarWidget QTableView {{
        background-color: {input_bg};
        color: {text};
        selection-background-color: {sel_bg};
        selection-color: {sel_text};
        alternate-background-color: {alt_bg};
    }}
    QCalendarWidget QTableView QHeaderView::section {{
        background-color: {alt_bg};
        color: {text};
        border: none;
    }}
    QCalendarWidget QAbstractItemView:enabled {{ color: {text}; }}
    QCalendarWidget QAbstractItemView:disabled {{ color: #777777; }}
    QCalendarWidget QMenu {{ background-color: {input_bg}; color: {text}; border: 1px solid {border}; }}
    QCalendarWidget QSpinBox {{ background: {input_bg}; color: {text}; selection-background-color: {sel_bg}; selection-color: {sel_text}; }}
    
    QToolButton#SpinArrowBtn {{
        background-color: {input_bg};
        color: {text};
        border: 1px solid {border};
        border-radius: 4px;
        font-weight: bold;
        font-size: 11pt;
    }}
    QToolButton#SpinArrowBtn:hover {{ background-color: {alt_bg}; }}
    
    QInputDialog {{ background-color: {bg}; color: {text}; }}
    QInputDialog QLabel {{ color: {text}; }}
    QDialogButtonBox QPushButton {{ 
        background-color: {input_bg}; color: {text}; border: 1px solid {border}; border-radius: 6px; padding: 8px 12px; font-weight: 600; min-width: 80px;
    }}
    QDialogButtonBox QPushButton:hover {{ background-color: {alt_bg}; }}
    """ 

SCHEMA = """
CREATE TABLE IF NOT EXISTS ars(
  id SERIAL PRIMARY KEY,
  nombre TEXT UNIQUE NOT NULL,
  sala_emergencia REAL NOT NULL DEFAULT 0,
  is_active INTEGER NOT NULL DEFAULT 1
);
CREATE TABLE IF NOT EXISTS universal_items(
  id SERIAL PRIMARY KEY,
  categoria TEXT CHECK(categoria IN ('Medicamentos','Materiales')) NOT NULL,
  nombre TEXT NOT NULL,
  precio REAL NOT NULL,
  is_active INTEGER NOT NULL DEFAULT 1,
  UNIQUE(categoria, nombre)
);
CREATE TABLE IF NOT EXISTS ars_items(
  id SERIAL PRIMARY KEY,
  ars_id INTEGER NOT NULL REFERENCES ars(id) ON DELETE CASCADE,
  categoria TEXT CHECK(categoria IN ('Laboratorios','Imágenes','Procedimientos','Honorarios')) NOT NULL,
  nombre TEXT NOT NULL,
  precio REAL NOT NULL,
  is_active INTEGER NOT NULL DEFAULT 1,
  UNIQUE(ars_id, categoria, nombre)
);
CREATE TABLE IF NOT EXISTS users(
  id SERIAL PRIMARY KEY,
  full_name TEXT NOT NULL,
  username TEXT UNIQUE NOT NULL,
  role TEXT NOT NULL CHECK(role IN ('auxiliar','administrador','facturador de auditoria')) DEFAULT 'auxiliar',
  password_hash TEXT NOT NULL,
  password_salt TEXT NOT NULL,
  security_question TEXT NOT NULL,
  security_answer_hash TEXT NOT NULL,
  security_answer_salt TEXT NOT NULL,
  is_active INTEGER NOT NULL DEFAULT 1,
  created_at TEXT NOT NULL,
  last_login TEXT
);
CREATE TABLE IF NOT EXISTS recibos(
  id SERIAL PRIMARY KEY,
  numero INTEGER NOT NULL UNIQUE,
  nombre TEXT,
  fecha TEXT,
  dx TEXT,
  ars TEXT,
  tipo_cobertura TEXT NOT NULL DEFAULT 'ASEGURADO',
  sala REAL NOT NULL DEFAULT 0,
  total REAL NOT NULL DEFAULT 0,
  pdf_filename TEXT,
  username TEXT,
  created_at TEXT,
  is_backdated INTEGER NOT NULL DEFAULT 0,
  pdf_synced INTEGER NOT NULL DEFAULT 0,
  pdf_sync_error TEXT,
  is_deleted INTEGER NOT NULL DEFAULT 0,
  deleted_at TEXT
);
CREATE TABLE IF NOT EXISTS recibo_items(
  id SERIAL PRIMARY KEY,
  recibo_id INTEGER NOT NULL REFERENCES recibos(id) ON DELETE CASCADE,
  categoria TEXT NOT NULL,
  nombre TEXT NOT NULL,
  precio_unit REAL NOT NULL,
  cantidad INTEGER NOT NULL,
  total REAL NOT NULL,
  ars TEXT
);
CREATE TABLE IF NOT EXISTS action_history(
  id SERIAL PRIMARY KEY,
  username TEXT NOT NULL,
  action TEXT NOT NULL,
  details TEXT,
  created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS session_control(
  username TEXT PRIMARY KEY,
  force_logout_at TEXT,
  requested_by TEXT,
  reason TEXT
);
CREATE TABLE IF NOT EXISTS active_sessions(
  username TEXT PRIMARY KEY,
  session_id TEXT NOT NULL,
  login_at TEXT NOT NULL,
  last_seen TEXT NOT NULL,
  is_active INTEGER NOT NULL DEFAULT 1,
  logout_at TEXT
);
CREATE SEQUENCE IF NOT EXISTS recibos_numero_seq;
CREATE TABLE IF NOT EXISTS daily_reports(
  id SERIAL PRIMARY KEY,
  report_date TEXT NOT NULL UNIQUE,
  generated_at TEXT NOT NULL,
  generated_by TEXT NOT NULL,
  filepath TEXT NOT NULL,
  totals_json TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS report_history(
  id SERIAL PRIMARY KEY,
  report_type TEXT NOT NULL,
  start_date TEXT NOT NULL,
  end_date TEXT NOT NULL,
  generated_at TEXT NOT NULL,
  generated_by TEXT NOT NULL,
  filepath TEXT NOT NULL,
  totals_json TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS pdf_storage(
  filename TEXT PRIMARY KEY,
  file_data BYTEA NOT NULL
);
CREATE TABLE IF NOT EXISTS user_preferences(
username TEXT PRIMARY KEY,
   auto_add_guantes INTEGER NOT NULL DEFAULT 1,
   auto_print INTEGER NOT NULL DEFAULT 0,
   theme TEXT NOT NULL DEFAULT 'claro',
   auto_add_bajante_cateter INTEGER NOT NULL DEFAULT 1
);
"""

# Pool de conexiones global
db_pool = None

def init_pool():
    global db_pool
    if db_pool is None:
        db_pool = pool.ThreadedConnectionPool(1, 20, DB_URL)

class PostgresWrapper:
    def __init__(self, url=None):
        if db_pool is None:
            init_pool()
        self.con = db_pool.getconn()
        self.con.autocommit = False

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        try:
            if exc_type is None:
                self.con.commit()
            else:
                self.con.rollback()
        finally:
            db_pool.putconn(self.con)

    def execute(self, query, params=None):
        cur = self.con.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.execute(query, params)
        return cur

    def executescript(self, script):
        cur = self.con.cursor()
        cur.execute(script)

def db_connect():
    return PostgresWrapper()

def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def remove_accents(input_str):
    if not input_str: return ""
    s = str(input_str).lower()
    reemplazos = (("á", "a"), ("é", "e"), ("í", "i"), ("ó", "o"), ("ú", "u"), ("ü", "u"))
    for a, b in reemplazos:
        s = s.replace(a, b)
    return s

def filter_table_widget(table: QTableWidget, term: str):
    term_norm = remove_accents(term)
    for r in range(table.rowCount()):
        match = False
        for c in range(table.columnCount()):
            item = table.item(r, c)
            if item and term_norm in remove_accents(item.text()):
                match = True
                break
        table.setRowHidden(r, not match)

def safe_pdf_text(text):
    if not text: return ""
    return text.encode('latin1', 'ignore').decode('latin1').strip()

def _pbkdf2(text: str, salt: bytes) -> str:
    raw = hashlib.pbkdf2_hmac("sha256", text.encode("utf-8"), salt, 120_000)
    return base64.b64encode(raw).decode("ascii")

def hash_secret(text: str):
    salt = os.urandom(16)
    return _pbkdf2(text, salt), base64.b64encode(salt).decode("ascii")

def verify_secret(text: str, hashed: str, salt_b64: str) -> bool:
    salt = base64.b64decode(salt_b64.encode("ascii"))
    probe = _pbkdf2(text, salt)
    return hmac.compare_digest(probe, hashed)

def log_action(username: str, action: str, details=None):
    username = str(username or "Sistema").strip()
    action = str(action or "Acción sin descripción").strip()
    details = "" if details is None else str(details).strip()
    with db_connect() as con:
        con.execute(
            "INSERT INTO action_history(username, action, details, created_at) VALUES(%s,%s,%s,%s)",
            (username, action, details, now_str())
        )

def db_init():
    os.makedirs(REPORTS_DIR, exist_ok=True)
    os.makedirs(PDFS_DIR, exist_ok=True)

    with db_connect() as con:
        con.executescript(SCHEMA)

        try:
            con.execute("ALTER TABLE ars_items DROP CONSTRAINT IF EXISTS ars_items_categoria_check")
            con.execute("ALTER TABLE ars_items ADD CONSTRAINT ars_items_categoria_check CHECK(categoria IN ('Laboratorios','Imágenes','Procedimientos','Honorarios'))")
        except Exception:
            pass

        try:
            con.execute("ALTER TABLE users DROP CONSTRAINT IF EXISTS users_role_check")
            con.execute("ALTER TABLE users ADD CONSTRAINT users_role_check CHECK(role IN ('auxiliar','administrador','facturador de auditoria'))")
        except Exception:
            pass

        cur = con.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'recibos'")
        recibos_cols = {row[0] for row in cur.fetchall()}
        if "username" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN username TEXT")
        if "created_at" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN created_at TEXT")
        if "is_backdated" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN is_backdated INTEGER NOT NULL DEFAULT 0")
        if "pdf_synced" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN pdf_synced INTEGER NOT NULL DEFAULT 0")
        if "pdf_sync_error" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN pdf_sync_error TEXT")
        if "is_deleted" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN is_deleted INTEGER NOT NULL DEFAULT 0")
        if "deleted_at" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN deleted_at TEXT")
        if "tipo_cobertura" not in recibos_cols:
            con.execute("ALTER TABLE recibos ADD COLUMN tipo_cobertura TEXT NOT NULL DEFAULT 'ASEGURADO'")
            con.execute(
                "UPDATE recibos SET tipo_cobertura=CASE WHEN COALESCE(ars, '')='' "
                "THEN 'NO_ASEGURADO' ELSE 'ASEGURADO' END"
            )

        con.execute(
            """UPDATE recibos r SET pdf_synced=1, pdf_sync_error=NULL
               WHERE r.pdf_synced=0
                 AND EXISTS (
                     SELECT 1 FROM pdf_storage p
                     WHERE p.filename=r.pdf_filename
                 )"""
        )

        # Alinea la secuencia una sola vez sin retrocederla ni duplicar números.
        con.execute("CREATE SEQUENCE IF NOT EXISTS recibos_numero_seq")
        max_numero = int(
            con.execute("SELECT COALESCE(MAX(numero), 0) FROM recibos").fetchone()[0]
        )
        sequence_row = con.execute(
            "SELECT last_value, is_called FROM recibos_numero_seq"
        ).fetchone()
        sequence_value = int(sequence_row[0]) if bool(sequence_row[1]) else 0
        sequence_target = max(max_numero, sequence_value)
        con.execute(
            "SELECT setval('recibos_numero_seq', %s, %s)",
            (max(sequence_target, 1), sequence_target > 0),
        )

        cur = con.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'ars'")
        ars_cols = {row[0] for row in cur.fetchall()}
        if "is_active" not in ars_cols:
            con.execute("ALTER TABLE ars ADD COLUMN is_active INTEGER NOT NULL DEFAULT 1")

        cur = con.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'universal_items'")
        univ_cols = {row[0] for row in cur.fetchall()}
        if "is_active" not in univ_cols:
            con.execute("ALTER TABLE universal_items ADD COLUMN is_active INTEGER NOT NULL DEFAULT 1")

        cur = con.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'ars_items'")
        ars_items_cols = {row[0] for row in cur.fetchall()}
        if "is_active" not in ars_items_cols:
            con.execute("ALTER TABLE ars_items ADD COLUMN is_active INTEGER NOT NULL DEFAULT 1")

        cur = con.execute("SELECT column_name FROM information_schema.columns WHERE table_name = 'user_preferences'")
        prefs_cols = {row[0] for row in cur.fetchall()}
        if "auto_add_bajante_cateter" not in prefs_cols:
            con.execute("ALTER TABLE user_preferences ADD COLUMN auto_add_bajante_cateter INTEGER NOT NULL DEFAULT 1")

    with db_connect() as con:
        cur = con.execute("SELECT COUNT(*) AS c FROM users WHERE role=%s", (ROLE_ADMIN,))
        has_admin = int(cur.fetchone()["c"]) > 0
        if not has_admin:
            pwd_hash, pwd_salt = hash_secret("admin123")
            ans_hash, ans_salt = hash_secret("admin")
            con.execute(
                """
                INSERT INTO users(
                    full_name, username, role, password_hash, password_salt,
                    security_question, security_answer_hash, security_answer_salt,
                    is_active, created_at
                ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                """,
                ("Administrador del sistema", "admin", ROLE_ADMIN, pwd_hash, pwd_salt, "Respuesta temporal", ans_hash, ans_salt, 1, now_str())
            )

def edit_universal_item(old_nombre: str, new_nombre: str, new_precio: float, categoria: str):
    with db_connect() as con:
        con.execute("UPDATE universal_items SET nombre=%s, precio=%s WHERE categoria=%s AND nombre=%s", (new_nombre.strip(), float(new_precio), categoria, old_nombre))

def edit_ars_item(ars_nombre: str, old_nombre: str, new_nombre: str, new_precio: float, categoria: str):
    with db_connect() as con:
        con.execute("UPDATE ars_items SET nombre=%s, precio=%s WHERE categoria=%s AND nombre=%s AND ars_id=(SELECT id FROM ars WHERE nombre=%s)", (new_nombre.strip(), float(new_precio), categoria, old_nombre, ars_nombre))

def move_universal_item(nombre: str, old_cat: str, new_cat: str):
    with db_connect() as con:
        con.execute("UPDATE universal_items SET categoria=%s WHERE categoria=%s AND nombre=%s", (new_cat, old_cat, nombre))

def move_ars_item(ars_nombre: str, nombre: str, old_cat: str, new_cat: str):
    with db_connect() as con:
        con.execute("UPDATE ars_items SET categoria=%s WHERE categoria=%s AND nombre=%s AND ars_id=(SELECT id FROM ars WHERE nombre=%s)", (new_cat, old_cat, nombre, ars_nombre))

def create_user(full_name: str, username: str, password: str, question: str, answer: str, role: str = ROLE_AUX):
    full_name = (full_name or "").strip()
    username = (username or "").strip()
    question = (question or "").strip()
    answer = (answer or "").strip()
    if not full_name or not username or not password or not question or not answer: raise ValueError("Todos los campos son obligatorios.")
    if len(password) < 4: raise ValueError("La contraseña debe tener al menos 4 caracteres.")
    pwd_hash, pwd_salt = hash_secret(password)
    ans_hash, ans_salt = hash_secret(answer.casefold())
    with db_connect() as con:
        con.execute("""INSERT INTO users(full_name, username, role, password_hash, password_salt, security_question, security_answer_hash, security_answer_salt, is_active, created_at) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""", (full_name, username, role, pwd_hash, pwd_salt, question, ans_hash, ans_salt, 1, now_str()))

def authenticate_user(username: str, password: str):
    with db_connect() as con:
        cur = con.execute("SELECT * FROM users WHERE username=%s", ((username or "").strip(),))
        row = cur.fetchone()
        if not row: return None
        if int(row["is_active"]) != 1: return None
        if not verify_secret(password or "", row["password_hash"], row["password_salt"]): return None
        con.execute("UPDATE users SET last_login=%s WHERE id=%s", (now_str(), row["id"]))
        return dict(row)

def get_user(username: str):
    with db_connect() as con:
        cur = con.execute("SELECT * FROM users WHERE username=%s", ((username or "").strip(),))
        row = cur.fetchone()
        return dict(row) if row else None

def list_users():
    with db_connect() as con:
        cur = con.execute("SELECT id, full_name, username, role, is_active, created_at, last_login FROM users ORDER BY role DESC, username")
        return [dict(r) for r in cur.fetchall()]

def list_usernames():
    with db_connect() as con:
        cur = con.execute("SELECT username FROM users WHERE is_active=1 ORDER BY username")
        return [r["username"] for r in cur.fetchall()]

def update_user_role(username: str, role: str):
    with db_connect() as con: con.execute("UPDATE users SET role=%s WHERE username=%s", (role, username))

def set_user_active(username: str, is_active: bool):
    with db_connect() as con: con.execute("UPDATE users SET is_active=%s WHERE username=%s", (1 if is_active else 0, username))

def request_remote_logout(username: str, requested_by: str, reason: str = "Sesión cerrada"):
    username = (username or "").strip()
    requested_by = (requested_by or "Sistema").strip()
    if not username:
        raise ValueError("Usuario inválido para cerrar sesión.")
    with db_connect() as con:
        con.execute(
            """
            INSERT INTO session_control(username, force_logout_at, requested_by, reason)
            VALUES(%s,%s,%s,%s)
            ON CONFLICT(username) DO UPDATE SET
                force_logout_at=EXCLUDED.force_logout_at,
                requested_by=EXCLUDED.requested_by,
                reason=EXCLUDED.reason
            """,
            (username, now_str(), requested_by, reason)
        )

def get_remote_logout_signal(username: str):
    username = (username or "").strip()
    if not username:
        return None
    with db_connect() as con:
        cur = con.execute("SELECT force_logout_at, requested_by, reason FROM session_control WHERE username=%s", (username,))
        row = cur.fetchone()
        return dict(row) if row else None

def make_session_id() -> str:
    return base64.urlsafe_b64encode(os.urandom(18)).decode("ascii").rstrip("=")

def register_active_session(username: str, session_id: str):
    username = (username or "").strip()
    session_id = (session_id or "").strip()
    if not username or not session_id:
        return
    stamp = now_str()
    with db_connect() as con:
        con.execute(
            """
            INSERT INTO active_sessions(username, session_id, login_at, last_seen, is_active, logout_at)
            VALUES(%s,%s,%s,%s,1,NULL)
            ON CONFLICT(username) DO UPDATE SET
                session_id=EXCLUDED.session_id,
                login_at=EXCLUDED.login_at,
                last_seen=EXCLUDED.last_seen,
                is_active=1,
                logout_at=NULL
            """,
            (username, session_id, stamp, stamp)
        )

def heartbeat_active_session(username: str, session_id: str):
    username = (username or "").strip()
    session_id = (session_id or "").strip()
    if not username or not session_id:
        return
    with db_connect() as con:
        con.execute(
            "UPDATE active_sessions SET last_seen=%s WHERE username=%s AND session_id=%s AND is_active=1",
            (now_str(), username, session_id)
        )

def end_active_session(username: str, session_id: str = ""):
    username = (username or "").strip()
    session_id = (session_id or "").strip()
    if not username:
        return
    with db_connect() as con:
        if session_id:
            con.execute(
                "UPDATE active_sessions SET is_active=0, logout_at=%s WHERE username=%s AND session_id=%s",
                (now_str(), username, session_id)
            )
        else:
            con.execute(
                "UPDATE active_sessions SET is_active=0, logout_at=%s WHERE username=%s",
                (now_str(), username)
            )

def get_active_sessions_map(max_age_seconds: int = 90) -> dict:
    active = {}
    cutoff = datetime.now() - timedelta(seconds=int(max_age_seconds))
    with db_connect() as con:
        cur = con.execute(
            "SELECT username, login_at, last_seen FROM active_sessions WHERE is_active=1"
        )
        rows = [dict(r) for r in cur.fetchall()]

    stale_users = []
    for row in rows:
        username = row.get("username") or ""
        last_seen_raw = row.get("last_seen") or ""
        try:
            last_seen_dt = datetime.strptime(str(last_seen_raw), "%Y-%m-%d %H:%M:%S")
        except Exception:
            last_seen_dt = datetime.min

        if last_seen_dt >= cutoff:
            active[username] = {
                "login_at": row.get("login_at") or "",
                "last_seen": last_seen_raw,
            }
        else:
            stale_users.append(username)

    if stale_users:
        try:
            with db_connect() as con:
                for username in stale_users:
                    con.execute(
                        "UPDATE active_sessions SET is_active=0, logout_at=%s WHERE username=%s AND is_active=1",
                        (now_str(), username)
                    )
        except Exception:
            pass

    return active

def admin_reset_password(username: str, new_password: str):
    pwd_hash, pwd_salt = hash_secret(new_password)
    with db_connect() as con: con.execute("UPDATE users SET password_hash=%s, password_salt=%s WHERE username=%s", (pwd_hash, pwd_salt, username))

def reset_password_by_security(username: str, answer: str, new_password: str) -> bool:
    with db_connect() as con:
        cur = con.execute("SELECT security_answer_hash, security_answer_salt FROM users WHERE username=%s", (username,))
        row = cur.fetchone()
        if not row: return False
        if not verify_secret((answer or "").strip().casefold(), row["security_answer_hash"], row["security_answer_salt"]): return False
        pwd_hash, pwd_salt = hash_secret(new_password)
        con.execute("UPDATE users SET password_hash=%s, password_salt=%s WHERE username=%s", (pwd_hash, pwd_salt, username))
        return True

def delete_user_db(username: str):
    with db_connect() as con:
        con.execute("DELETE FROM users WHERE username=%s", (username,))

def get_security_question(username: str):
    with db_connect() as con:
        cur = con.execute("SELECT security_question FROM users WHERE username=%s", (username,))
        row = cur.fetchone()
        return row["security_question"] if row else ""

def ars_list():
    with db_connect() as con:
        cur = con.execute("SELECT nombre FROM ars WHERE is_active=1 ORDER BY nombre")
        return [r["nombre"] for r in cur.fetchall()]

def get_emergency_price(nombre_ars: str) -> float:
    with db_connect() as con:
        cur = con.execute("SELECT sala_emergencia FROM ars WHERE nombre=%s AND is_active=1", (nombre_ars,))
        row = cur.fetchone()
        return float(row["sala_emergencia"]) if row else 0.0

def set_emergency_price(nombre_ars: str, precio: float):
    with db_connect() as con: con.execute("INSERT INTO ars(nombre, sala_emergencia, is_active) VALUES(%s,%s, 1) ON CONFLICT(nombre) DO UPDATE SET sala_emergencia=EXCLUDED.sala_emergencia, is_active=1", (nombre_ars, float(precio)))

def upsert_ars(nombre: str, sala: float = 0.0):
    set_emergency_price(nombre.strip(), sala)

def delete_ars(nombre: str):
    with db_connect() as con: con.execute("UPDATE ars SET is_active=0 WHERE nombre=%s", (nombre,))

def get_universal(categoria: str) -> dict:
    with db_connect() as con:
        cur = con.execute("SELECT nombre, precio FROM universal_items WHERE categoria=%s AND is_active=1 ORDER BY nombre", (categoria,))
        return {r["nombre"]: float(r["precio"]) for r in cur.fetchall()}

def upsert_universal(categoria: str, nombre: str, precio: float):
    with db_connect() as con: con.execute("INSERT INTO universal_items(categoria, nombre, precio, is_active) VALUES(%s,%s,%s, 1) ON CONFLICT(categoria, nombre) DO UPDATE SET precio=EXCLUDED.precio, is_active=1", (categoria, nombre.strip(), float(precio)))

def delete_universal(categoria: str, nombre: str):
    with db_connect() as con: con.execute("UPDATE universal_items SET is_active=0 WHERE categoria=%s AND nombre=%s", (categoria, nombre))

def _ensure_ars(nombre_ars: str):
    with db_connect() as con: con.execute("INSERT INTO ars(nombre, sala_emergencia, is_active) VALUES(%s, COALESCE((SELECT sala_emergencia FROM ars WHERE nombre=%s), 0), 1) ON CONFLICT(nombre) DO UPDATE SET is_active=1", (nombre_ars, nombre_ars))

def get_ars_items(categoria: str, nombre_ars: str) -> dict:
    with db_connect() as con:
        cur = con.execute("""SELECT ai.nombre, ai.precio FROM ars_items ai JOIN ars a ON a.id = ai.ars_id WHERE ai.categoria=%s AND a.nombre=%s AND ai.is_active=1 ORDER BY ai.nombre""", (categoria, nombre_ars))
        return {r["nombre"]: float(r["precio"]) for r in cur.fetchall()}

def upsert_ars_item(categoria: str, nombre_ars: str, nombre: str, precio: float):
    _ensure_ars(nombre_ars)
    with db_connect() as con: con.execute("""INSERT INTO ars_items(ars_id, categoria, nombre, precio, is_active) SELECT a.id, %s, %s, %s, 1 FROM ars a WHERE a.nombre=%s ON CONFLICT(ars_id, categoria, nombre) DO UPDATE SET precio=EXCLUDED.precio, is_active=1""", (categoria, nombre.strip(), float(precio), nombre_ars))

def delete_ars_item(categoria: str, nombre_ars: str, nombre: str):
    with db_connect() as con: con.execute("UPDATE ars_items SET is_active=0 WHERE categoria=%s AND nombre=%s AND ars_id=(SELECT id FROM ars WHERE nombre=%s)", (categoria, nombre, nombre_ars))

def get_next_recibo_number() -> int:
    with db_connect() as con:
        cur = con.execute("SELECT nextval('recibos_numero_seq') AS nextnum")
        return int(cur.fetchone()["nextnum"])

def add_recibo(numero: int, nombre: str, fecha: str, dx: str, ars: str, sala: float, total: float, pdf_filename: str, username: str, is_backdated: int = 0, created_at: str = None):
    # fecha = fecha seleccionada del servicio/recibo.
    # created_at = fecha real en que se generó el recibo.
    # Los reportes se calculan por created_at, para que un recibo atrasado hecho hoy caiga en el reporte de hoy.
    created_at = created_at or now_str()
    with db_connect() as con:
        cur = con.execute("""INSERT INTO recibos(numero, nombre, fecha, dx, ars, sala, total, pdf_filename, username, created_at, is_backdated) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id""", (int(numero), nombre or "", fecha or "", dx or "", ars or "", float(sala), float(total), pdf_filename or "", username or "", created_at, is_backdated))
        return int(cur.fetchone()[0])

def add_recibo_item(recibo_id: int, categoria: str, nombre: str, precio_unit: float, cantidad: int, total: float, ars: str = ""):
    with db_connect() as con: con.execute("INSERT INTO recibo_items(recibo_id, categoria, nombre, precio_unit, cantidad, total, ars) VALUES(%s,%s,%s,%s,%s,%s,%s)", (recibo_id, categoria, nombre, float(precio_unit), int(cantidad), float(total), ars or ""))

def update_recibo_db(recibo_id: int, nombre: str, fecha: str, dx: str, ars: str, sala: float, total: float, pdf_filename: str, username: str, is_backdated: int = 0):
    with db_connect() as con:
        # Mantener el username original al editar (no actualizar)
        con.execute("""UPDATE recibos SET nombre=%s, fecha=%s, dx=%s, ars=%s, sala=%s, total=%s, pdf_filename=%s, is_backdated=%s WHERE id=%s""", (nombre or "", fecha or "", dx or "", ars or "", float(sala), float(total), pdf_filename or "", is_backdated, recibo_id))
        con.execute("DELETE FROM recibo_items WHERE recibo_id=%s", (recibo_id,))


def save_receipt_with_items(
    recibo_id,
    numero,
    nombre,
    fecha,
    dx,
    ars,
    sala,
    total,
    pdf_filename,
    username,
    is_backdated,
    created_at,
    grouped,
    coverage="ASEGURADO",
):
    """Guarda cabecera, ítems e historial con un solo commit a PostgreSQL."""
    item_rows = []
    for categoria, items in grouped:
        item_ars = ars if categoria in ARS_CATEGORIES else ""
        for item_name, precio, cantidad, item_total, _ in items:
            item_rows.append(
                (
                    categoria,
                    item_name,
                    float(precio),
                    int(cantidad),
                    float(item_total),
                    item_ars,
                )
            )

    editing = recibo_id is not None
    with db_connect() as con:
        if editing:
            con.execute(
                """UPDATE recibos
                   SET nombre=%s, fecha=%s, dx=%s, ars=%s, tipo_cobertura=%s, sala=%s, total=%s,
                       pdf_filename=%s, is_backdated=%s, pdf_synced=0,
                       pdf_sync_error=NULL
                   WHERE id=%s""",
                (
                    nombre or "", fecha or "", dx or "", ars or "", coverage,
                    float(sala), float(total), pdf_filename or "",
                    int(is_backdated), int(recibo_id),
                ),
            )
            con.execute("DELETE FROM recibo_items WHERE recibo_id=%s", (int(recibo_id),))
            saved_id = int(recibo_id)
            action = "Edición de factura PDF"
            details = f"Recibo {numero} - Total ${float(total):,.2f}"
        else:
            cur = con.execute(
                """INSERT INTO recibos(
                       numero, nombre, fecha, dx, ars, tipo_cobertura, sala, total, pdf_filename,
                       username, created_at, is_backdated, pdf_synced
                   ) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,0)
                   RETURNING id""",
                (
                    int(numero), nombre or "", fecha or "", dx or "", ars or "", coverage,
                    float(sala), float(total), pdf_filename or "", username or "",
                    created_at or now_str(), int(is_backdated),
                ),
            )
            saved_id = int(cur.fetchone()[0])
            action = "Generar factura PDF"
            details = (
                f"Recibo {numero} - Total ${float(total):,.2f} "
                f"(Pasado: {bool(is_backdated)})"
            )

        if item_rows:
            cursor = con.con.cursor()
            psycopg2.extras.execute_values(
                cursor,
                """INSERT INTO recibo_items(
                       recibo_id, categoria, nombre, precio_unit, cantidad, total, ars
                   ) VALUES %s""",
                [(saved_id, *row) for row in item_rows],
                page_size=500,
            )

        con.execute(
            "INSERT INTO action_history(username, action, details, created_at) VALUES(%s,%s,%s,%s)",
            (str(username or "Sistema"), action, details, now_str()),
        )

    return saved_id

def delete_recibo(recibo_id: int):
    with db_connect() as con:
        con.execute("UPDATE recibos SET is_deleted=1, deleted_at=%s WHERE id=%s", (now_str(), int(recibo_id)))

def purge_old_deleted_receipts(days: int = 30):
    cutoff = datetime.now() - timedelta(days=int(days))
    cutoff_str = cutoff.strftime("%Y-%m-%d %H:%M:%S")
    with db_connect() as con:
        cur = con.execute("SELECT id, pdf_filename FROM recibos WHERE is_deleted=1 AND deleted_at IS NOT NULL AND deleted_at < %s", (cutoff_str,))
        rows = [dict(r) for r in cur.fetchall()]
        con.execute("DELETE FROM recibos WHERE is_deleted=1 AND deleted_at IS NOT NULL AND deleted_at < %s", (cutoff_str,))
    return rows

def restore_recibo(recibo_id: int):
    with db_connect() as con:
        con.execute("UPDATE recibos SET is_deleted=0, deleted_at=NULL WHERE id=%s", (int(recibo_id),))

def permanently_delete_recibo(recibo_id: int):
    with db_connect() as con:
        con.execute("DELETE FROM recibo_items WHERE recibo_id=%s", (int(recibo_id),))
        con.execute("DELETE FROM recibos WHERE id=%s", (int(recibo_id),))

def list_deleted_receipts(limit: int = 50, offset: int = 0):
    with db_connect() as con:
        cur = con.execute("""SELECT id, numero, nombre, fecha, ars, total, username, created_at, deleted_at FROM recibos WHERE is_deleted=1 ORDER BY deleted_at DESC LIMIT %s OFFSET %s""", (int(limit), int(offset)))
        return [dict(r) for r in cur.fetchall()]

def get_recent_history(limit: int = 30):
    with db_connect() as con:
        cur = con.execute("SELECT username, action, details, created_at FROM action_history ORDER BY id DESC LIMIT %s", (int(limit),))
        return [dict(r) for r in cur.fetchall()]

def list_receipts_history(limit: int = 500, offset: int = 0):
    with db_connect() as con:
        cur = con.execute("""SELECT id, numero, nombre, fecha, ars, total, username, created_at, pdf_filename, is_backdated FROM recibos WHERE is_deleted=0 ORDER BY id DESC LIMIT %s OFFSET %s""", (int(limit), int(offset)))
        return [dict(r) for r in cur.fetchall()]

# ---> LÓGICA DE FILTRADO (ARS Y USUARIOS) <---
def _normalize_report_filter(value, empty_text):
    if isinstance(value, dict):
        mode = "exclude" if value.get("mode") in ("exclude", "excluir") else "include"
        values = [str(item).strip() for item in value.get("values", []) if str(item).strip()]
        return {"mode": mode, "values": values}
    text = str(value or "").strip()
    if not text or text == empty_text:
        return {"mode": "include", "values": []}
    return {"mode": "include", "values": [text]}


def _report_filter_clause(column, selection):
    values = selection.get("values", [])
    if not values:
        return "", None
    if selection.get("mode") == "exclude":
        return f" AND NOT (COALESCE({column}, '') = ANY(%s))", values
    return f" AND {column} = ANY(%s)", values


def _report_filter_text(selection, empty_text, label):
    values = selection.get("values", [])
    if not values:
        return empty_text
    action = "Excluir" if selection.get("mode") == "exclude" else "Incluir"
    return f"{label} · {action}: {', '.join(values)}"


def _report_filter_is_all(selection):
    return not selection.get("values")


def get_receipt_stats_between(start_date: str, end_date: str, is_backdated: int = 0, ars_filter="Todas las ARS", user_filter="Todos los Usuarios"):
    totals = {cat: 0.0 for cat in ALL_CATEGORIES}

    # Los reportes se calculan por la FECHA REAL DE CREACIÓN del recibo (created_at).
    # Solo se incluyen recibos que tengan created_at definido.
    created_date_expr = "r.created_at::timestamp::date"
    include_all_histories = is_backdated is None or str(is_backdated) == "-1"

    with db_connect() as con:
        base_query_items = (
            "SELECT ri.categoria, COALESCE(SUM(ri.total), 0) "
            "FROM recibo_items ri "
            "JOIN recibos r ON r.id = ri.recibo_id "
            f"WHERE {created_date_expr} BETWEEN %s::date AND %s::date "
            "AND r.is_deleted=0 "
        )
        base_query_sala = (
            "SELECT COALESCE(SUM(sala), 0), COALESCE(SUM(total), 0) "
            "FROM recibos r "
            f"WHERE {created_date_expr} BETWEEN %s::date AND %s::date "
            "AND r.is_deleted=0 "
        )

        count_query_ars = (
            "SELECT ars, COUNT(id) "
            "FROM recibos r "
            f"WHERE {created_date_expr} BETWEEN %s::date AND %s::date "
            "AND r.is_deleted=0 "
        )
        count_query_user = (
            "SELECT username, COUNT(id) "
            "FROM recibos r "
            f"WHERE {created_date_expr} BETWEEN %s::date AND %s::date "
            "AND r.is_deleted=0 "
        )

        params = [start_date, end_date]

        if not include_all_histories:
            base_query_items += " AND r.is_backdated = %s"
            base_query_sala += " AND r.is_backdated = %s"
            count_query_ars += " AND r.is_backdated = %s"
            count_query_user += " AND r.is_backdated = %s"
            params.append(int(is_backdated))

        ars_selection = _normalize_report_filter(ars_filter, "Todas las ARS")
        user_selection = _normalize_report_filter(user_filter, "Todos los Usuarios")
        for column, selection in (("r.ars", ars_selection), ("r.username", user_selection)):
            clause, value = _report_filter_clause(column, selection)
            if clause:
                base_query_items += clause
                base_query_sala += clause
                count_query_ars += clause
                count_query_user += clause
                params.append(value)

        base_query_items += " GROUP BY ri.categoria"
        count_query_ars += " GROUP BY ars ORDER BY ars"
        count_query_user += " GROUP BY username ORDER BY username"

        cur = con.execute(base_query_items, tuple(params))
        for r in cur.fetchall():
            if r[0] in totals:
                totals[r[0]] = float(r[1])

        cur = con.execute(base_query_sala, tuple(params))
        r = cur.fetchone()
        totals["Sala Emergencia"] = float(r[0] or 0.0)
        totals["Total General"] = float(r[1] or 0.0)

        ars_counts = {}
        total_recibos = 0
        cur = con.execute(count_query_ars, tuple(params))
        for row in cur.fetchall():
            a_name = row[0] if row[0] else "Sin ARS"
            ars_counts[a_name] = int(row[1])
            total_recibos += int(row[1])

        user_counts = {}
        cur = con.execute(count_query_user, tuple(params))
        for row in cur.fetchall():
            user_counts[row[0]] = int(row[1])

        totals["_ars_counts"] = ars_counts
        totals["_user_counts"] = user_counts
        totals["_total_recibos"] = total_recibos
        totals["_include_all_histories"] = bool(include_all_histories)
        totals["_filters"] = {"ars": ars_selection, "users": user_selection}
        totals["_filter_summary"] = (
            f"{_report_filter_text(ars_selection, 'Todas las ARS', 'ARS')} · "
            f"{_report_filter_text(user_selection, 'Todos los facturadores', 'Facturadores')}"
        )

    return totals

def get_receipt_stats_by_date(report_date: str, is_backdated: int = 0, ars_filter="Todas las ARS", user_filter="Todos los Usuarios"):
    return get_receipt_stats_between(report_date, report_date, is_backdated, ars_filter, user_filter)


def _legacy_get_dashboard_statistics(
    start_date: str,
    end_date: str,
    ars_filter: str = "Todas las ARS",
    user_filter: str = "Todos los Usuarios",
    breakdown: str = "ars",
):
    """Datos consolidados del panel usando la fecha real de generación."""
    start_dt = datetime.strptime(start_date, "%Y-%m-%d")
    end_dt = datetime.strptime(end_date, "%Y-%m-%d")
    period_days = max(1, (end_dt - start_dt).days + 1)
    previous_end = start_dt - timedelta(days=1)
    previous_start = previous_end - timedelta(days=period_days - 1)

    def filters(alias: str, from_date: str, to_date: str):
        prefix = f"{alias}." if alias else ""
        clauses = [
            f"{prefix}created_at IS NOT NULL",
            f"{prefix}created_at::timestamp::date BETWEEN %s::date AND %s::date",
            f"{prefix}is_deleted=0",
        ]
        params = [from_date, to_date]
        if ars_filter != "Todas las ARS":
            clauses.append(f"{prefix}ars=%s")
            params.append(ars_filter)
        if user_filter != "Todos los Usuarios":
            clauses.append(f"{prefix}username=%s")
            params.append(user_filter)
        return " AND ".join(clauses), params

    current_where, current_params = filters("r", start_date, end_date)
    previous_where, previous_params = filters(
        "r",
        previous_start.strftime("%Y-%m-%d"),
        previous_end.strftime("%Y-%m-%d"),
    )
    dimension = "r.username" if breakdown == "username" else "r.ars"

    with db_connect() as con:
        row = con.execute(
            f"""SELECT COUNT(*) AS receipts,
                       COALESCE(SUM(r.total), 0) AS total,
                       COALESCE(AVG(r.total), 0) AS average,
                       COALESCE(SUM(r.sala), 0) AS room
                FROM recibos r WHERE {current_where}""",
            tuple(current_params),
        ).fetchone()
        summary = {
            "receipts": int(row[0] or 0),
            "total": float(row[1] or 0),
            "average": float(row[2] or 0),
            "room": float(row[3] or 0),
        }

        previous = con.execute(
            f"""SELECT COUNT(*), COALESCE(SUM(r.total), 0)
                FROM recibos r WHERE {previous_where}""",
            tuple(previous_params),
        ).fetchone()
        previous_summary = {
            "receipts": int(previous[0] or 0),
            "total": float(previous[1] or 0),
        }

        monthly_rows = con.execute(
            f"""SELECT TO_CHAR(DATE_TRUNC('month', r.created_at::timestamp), 'YYYY-MM') AS month,
                       COUNT(*) AS receipts, COALESCE(SUM(r.total), 0) AS total
                FROM recibos r
                WHERE {current_where}
                GROUP BY DATE_TRUNC('month', r.created_at::timestamp)
                ORDER BY DATE_TRUNC('month', r.created_at::timestamp)""",
            tuple(current_params),
        ).fetchall()
        monthly = [
            {"label": str(row[0]), "receipts": int(row[1]), "total": float(row[2])}
            for row in monthly_rows
        ]

        category_rows = con.execute(
            f"""SELECT ri.categoria, COALESCE(SUM(ri.total), 0)
                FROM recibo_items ri
                JOIN recibos r ON r.id=ri.recibo_id
                WHERE {current_where}
                GROUP BY ri.categoria
                ORDER BY SUM(ri.total) DESC""",
            tuple(current_params),
        ).fetchall()
        categories = [(str(row[0]), float(row[1])) for row in category_rows]
        if summary["room"] > 0:
            categories.append(("Sala Emergencia", summary["room"]))

        breakdown_rows = con.execute(
            f"""SELECT COALESCE(NULLIF({dimension}, ''), 'Sin dato') AS label,
                       COUNT(*) AS receipts, COALESCE(SUM(r.total), 0) AS total
                FROM recibos r
                WHERE {current_where}
                GROUP BY {dimension}
                ORDER BY SUM(r.total) DESC
                LIMIT 10""",
            tuple(current_params),
        ).fetchall()
        breakdown_data = [
            {"label": str(row[0]), "receipts": int(row[1]), "total": float(row[2])}
            for row in breakdown_rows
        ]

    return {
        "start_date": start_date,
        "end_date": end_date,
        "summary": summary,
        "previous": previous_summary,
        "monthly": monthly,
        "categories": categories,
        "breakdown": breakdown_data,
        "breakdown_type": breakdown,
    }


def get_dashboard_statistics(
    start_date: str,
    end_date: str,
    ars_filter: str = "Todas las ARS",
    user_filter: str = "Todos los Usuarios",
    medication: str = "Todos los medicamentos",
    category: str = "Todas las categorías",
    trend_granularity: str = "day",
    coverage: str = "Todas",
    compare_previous: bool = False,
    previous_start: str = "",
    previous_end: str = "",
):
    return PanelDataService(db_connect).load(
        start_date=start_date,
        end_date=end_date,
        ars_filter=ars_filter,
        user_filter=user_filter,
        medication=medication,
        category=category,
        trend_granularity=trend_granularity,
        coverage=coverage,
        compare_previous=compare_previous,
        previous_start=previous_start,
        previous_end=previous_end,
    )

def list_receipts_for_report(start_date: str, end_date: str, is_backdated: int = 0, ars_filter="Todas las ARS", user_filter="Todos los Usuarios"):
    """Devuelve los recibos incluidos en un reporte usando la fecha real de generación (created_at).

    Mantiene ambas fechas:
    - fecha: fecha seleccionada/servicio.
    - created_at: fecha y hora en que el recibo fue generado.
    Solo se incluyen recibos con created_at definido.
    """
    created_date_expr = "r.created_at::timestamp::date"
    include_all_histories = is_backdated is None or str(is_backdated) == "-1"
    query = (
        "SELECT r.numero, r.nombre, r.fecha, r.created_at AS created_at, "
        "r.ars, r.total, r.username "
        "FROM recibos r "
        f"WHERE {created_date_expr} BETWEEN %s::date AND %s::date "
        "AND r.is_deleted=0 "
    )
    params = [start_date, end_date]

    if not include_all_histories:
        query += " AND r.is_backdated = %s"
        params.append(int(is_backdated))

    ars_selection = _normalize_report_filter(ars_filter, "Todas las ARS")
    user_selection = _normalize_report_filter(user_filter, "Todos los Usuarios")
    for column, selection in (("r.ars", ars_selection), ("r.username", user_selection)):
        clause, value = _report_filter_clause(column, selection)
        if clause:
            query += clause
            params.append(value)

    query += f" ORDER BY {created_date_expr}, r.numero"

    with db_connect() as con:
        cur = con.execute(query, tuple(params))
        return [dict(r) for r in cur.fetchall()]

def report_exists(report_date: str) -> bool:
    with db_connect() as con:
        cur = con.execute("SELECT 1 FROM daily_reports WHERE report_date=%s", (report_date,))
        return cur.fetchone() is not None

def save_daily_report_record(report_date: str, filepath: str, totals: dict, generated_by: str):
    with db_connect() as con: con.execute("INSERT INTO daily_reports(report_date, generated_at, generated_by, filepath, totals_json) VALUES(%s,%s,%s,%s,%s) ON CONFLICT(report_date) DO UPDATE SET generated_at=EXCLUDED.generated_at, generated_by=EXCLUDED.generated_by, filepath=EXCLUDED.filepath, totals_json=EXCLUDED.totals_json", (report_date, now_str(), generated_by, filepath, json.dumps(totals, ensure_ascii=False)))

def save_report_history(report_type: str, start_date: str, end_date: str, filepath: str, totals: dict, generated_by: str):
    with db_connect() as con: con.execute("INSERT INTO report_history(report_type, start_date, end_date, generated_at, generated_by, filepath, totals_json) VALUES(%s,%s,%s,%s,%s,%s,%s)", (report_type, start_date, end_date, now_str(), generated_by, filepath, json.dumps(totals, ensure_ascii=False)))

def list_daily_reports(limit: int = 200):
    with db_connect() as con:
        cur = con.execute(
            "SELECT id AS record_id, 'daily_reports' AS source_table, 'Diario' AS report_type, "
            "report_date AS start_date, report_date AS end_date, generated_at, generated_by, filepath, totals_json "
            "FROM daily_reports ORDER BY report_date DESC LIMIT %s",
            (int(limit),)
        )
        return [dict(r) for r in cur.fetchall()]

def list_report_history(limit: int = 300):
    with db_connect() as con:
        cur = con.execute(
            "SELECT id AS record_id, 'report_history' AS source_table, report_type, start_date, end_date, "
            "generated_at, generated_by, filepath, totals_json "
            "FROM report_history ORDER BY id DESC LIMIT %s",
            (int(limit),)
        )
        rows = [dict(r) for r in cur.fetchall()]
    rows.extend(list_daily_reports(limit))
    rows.sort(key=lambda r: (r.get("generated_at") or ""), reverse=True)
    return rows[:limit]

def delete_report_record(source_table: str, record_id: int, filepath: str = ""):
    allowed_tables = {"daily_reports", "report_history"}
    if source_table not in allowed_tables:
        raise ValueError("Tipo de reporte inválido.")

    stored_filepath = filepath or ""
    with db_connect() as con:
        cur = con.execute(f"SELECT filepath FROM {source_table} WHERE id=%s", (int(record_id),))
        row = cur.fetchone()
        if row and row["filepath"]:
            stored_filepath = row["filepath"]

        con.execute(f"DELETE FROM {source_table} WHERE id=%s", (int(record_id),))

        filename = os.path.basename(stored_filepath or "")
        if filename:
            con.execute("DELETE FROM pdf_storage WHERE filename=%s", (filename,))

    filename = os.path.basename(stored_filepath or filepath or "")
    for candidate in {
        stored_filepath,
        stable_storage_path(REPORTS_DIR, filename),
    }:
        if candidate and os.path.exists(candidate):
            try:
                os.remove(candidate)
            except Exception:
                pass

def get_user_preferences(username: str) -> dict:
    with db_connect() as con:
        cur = con.execute("SELECT auto_add_guantes, auto_print, auto_add_bajante_cateter FROM user_preferences WHERE username=%s", (username,))
        row = cur.fetchone()
        if row:
            return {
                "auto_add_guantes": bool(row["auto_add_guantes"]),
                "auto_print": bool(row["auto_print"]),
                "auto_add_bajante_cateter": bool(row["auto_add_bajante_cateter"]),
            }
        return {"auto_add_guantes": True, "auto_print": False, "auto_add_bajante_cateter": True}

def upsert_user_preferences(username: str, auto_add_guantes: bool = None, auto_print: bool = None, auto_add_bajante_cateter: bool = None):
    with db_connect() as con:
        cur = con.execute("SELECT username FROM user_preferences WHERE username=%s", (username,))
        exists = cur.fetchone() is not None
        if not exists:
            con.execute(
                "INSERT INTO user_preferences(username, auto_add_guantes, auto_print, auto_add_bajante_cateter) VALUES(%s,%s,%s,%s)",
                (username, 1 if auto_add_guantes is None else (1 if auto_add_guantes else 0), 0 if auto_print is None else (1 if auto_print else 0), 1 if auto_add_bajante_cateter is None else (1 if auto_add_bajante_cateter else 0)),
            )
        else:
            sets = []
            params = []
            if auto_add_guantes is not None:
                sets.append("auto_add_guantes=%s")
                params.append(1 if auto_add_guantes else 0)
            if auto_print is not None:
                sets.append("auto_print=%s")
                params.append(1 if auto_print else 0)
            if auto_add_bajante_cateter is not None:
                sets.append("auto_add_bajante_cateter=%s")
                params.append(1 if auto_add_bajante_cateter else 0)
            if sets:
                params.append(username)
                con.execute(f"UPDATE user_preferences SET {', '.join(sets)} WHERE username=%s", tuple(params))

def get_user_receipt_counts() -> dict:
    with db_connect() as con:
        cur = con.execute("SELECT username, COUNT(id) AS c FROM recibos GROUP BY username")
        return {row["username"]: int(row["c"]) for row in cur.fetchall()}

def get_recibo_data(recibo_id: int):
    with db_connect() as con:
        cur = con.execute("SELECT * FROM recibos WHERE id=%s", (recibo_id,))
        recibo = dict(cur.fetchone())
        cur = con.execute("SELECT * FROM recibo_items WHERE recibo_id=%s", (recibo_id,))
        recibo["items"] = [dict(r) for r in cur.fetchall()]
        return recibo

def _to_float(raw: str):
    if raw is None: return None
    v = re.sub(r'[^0-9,.\-]', '', raw)
    if not v: return None
    if ',' in v and '.' in v:
        if v.rfind('.') > v.rfind(','): v = v.replace(',', '')
        else: v = v.replace('.', '').replace(',', '.')
    elif ',' in v:
        left, right = v.rsplit(',', 1)
        if right.isdigit() and 1 <= len(right) <= 2: v = left.replace(',', '') + '.' + right
        else: v = v.replace(',', '')
    try: return float(v)
    except Exception: return None

def _parse_price(text: str):
    if text is None: return None
    if '$' in text:
        tail = text.rsplit('$', 1)[-1]
        return _to_float(tail)
    nums = re.findall(r'(-?\d[\d.,]*)', text)
    if not nums: return None
    return _to_float(nums[-1])

def _norm_cat(txt: str):
    if not txt: return None
    t = txt.strip().casefold()
    if 'medic' in t: return 'Medicamentos'
    if 'mater' in t: return 'Materiales'
    return None

def _norm_ars_cat(txt: str):
    if not txt: return None
    t = txt.strip().casefold()
    t = t.replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
    if 'laborat' in t: return 'Laboratorios'
    if 'imagen' in t: return 'Imágenes'
    if 'proced' in t: return 'Procedimientos'
    if 'honorario' in t: return 'Honorarios'
    return None

def apply_price_rule(category: str, base_price: float) -> float:
    return round(float(base_price), 2)

def get_effective_price(category: str, stored_price: float) -> float:
    stored_price = float(stored_price)
    if category == 'Medicamentos': return round(stored_price * 1.20, 2)
    return round(stored_price, 2)

def set_button_role(button: QPushButton, role: str):
    palette = {
        'success': ('#2e7d32', '#1b5e20'),
        'danger': ('#c62828', '#8e0000'),
        'warning': ('#ef6c00', '#c25e00'),
        'info': ('#1565c0', '#0d47a1'),
        'report': ('#6a1b9a', '#4a148c'),
        'neutral': ('#e7edf3', '#d5e0ea'),
    }
    base, hover = palette.get(role, ('#1e88e5', '#1565c0'))
    border = "1px solid #c4d0dc" if role == 'neutral' else "none"
    button.setStyleSheet(
        f"QPushButton {{ background-color: {base}; color: {'#263238' if role == 'neutral' else 'white'}; border: {border}; border-radius: 6px; padding: 8px 12px; font-weight: 600; }}"
        f"QPushButton:hover {{ background-color: {hover}; }}"
        f"QPushButton:disabled {{ background-color: #b0bec5; color: #757575; }}"
    )


class NoWheelTabBar(QTabBar):
    """Evita cambiar de sección accidentalmente con la rueda del mouse."""

    def wheelEvent(self, event):
        event.ignore()


class MultiSelectFilter(QToolButton):
    """Abre un diálogo transaccional: los cambios solo se guardan al aplicar."""

    selectionChanged = Signal()

    def __init__(self, values, empty_text="Todos", item_name="elemento", feminine=False, parent=None):
        super().__init__(parent)
        self.empty_text = empty_text
        self.item_name = item_name
        self.feminine = feminine
        self._mode = "include"
        self._values = [str(value) for value in values]
        self._selected = set()
        self.setToolButtonStyle(Qt.ToolButtonTextOnly)
        self.setArrowType(Qt.DownArrow)
        self.clicked.connect(self._open_dialog)
        self._refresh_summary(emit=False)

    def mode(self):
        return self._mode

    def set_mode(self, mode):
        self._mode = "exclude" if mode in ("exclude", "excluir") else "include"
        self._refresh_summary()

    def filter_data(self):
        return {"mode": self._mode, "values": self.selected_values()}

    def selected_values(self):
        return [value for value in self._values if value in self._selected]

    def clear_selection(self):
        self._selected.clear()
        self._refresh_summary()

    def select_all(self):
        self._selected = set(self._values)
        self._refresh_summary()

    def _refresh_summary(self, emit=True):
        values = self.selected_values()
        if not values:
            self.setText(self.empty_text)
        else:
            noun = self.item_name
            if len(values) != 1 and noun != "ARS":
                noun += "es" if noun.endswith("r") else "s"
            suffix = (
                ("excluida" if self.feminine else "excluido")
                if self._mode == "exclude"
                else ("incluida" if self.feminine else "incluido")
            )
            if len(values) != 1:
                suffix += "s"
            self.setText(f"{len(values)} {noun} {suffix}")
        mode_text = "Excluir" if self._mode == "exclude" else "Incluir"
        self.setToolTip(
            f"{mode_text}: {', '.join(values)}" if values else self.empty_text
        )
        if emit:
            self.selectionChanged.emit()

    def _open_dialog(self):
        dialog = QDialog(self)
        dialog.setObjectName("MultiSelectFilterDialog")
        dialog.setWindowTitle(f"Filtrar {self.item_name}")
        dialog.setMinimumSize(500, 560)
        root = QVBoxLayout(dialog)
        root.setContentsMargins(18, 16, 18, 16)
        root.setSpacing(11)

        title = QLabel(f"Filtrar {self.item_name}")
        title.setObjectName("FilterDialogTitle")
        title.setStyleSheet("font-size: 16pt; font-weight: 900;")
        root.addWidget(title)

        root.addWidget(QLabel("Modo del filtro:"))
        mode_row = QHBoxLayout()
        include_radio = QRadioButton("Incluir seleccionadas")
        exclude_radio = QRadioButton("Excluir seleccionadas")
        include_radio.setChecked(self._mode == "include")
        exclude_radio.setChecked(self._mode == "exclude")
        mode_row.addWidget(include_radio)
        mode_row.addWidget(exclude_radio)
        mode_row.addStretch(1)
        root.addLayout(mode_row)

        search = QLineEdit()
        search.setPlaceholderText(f"Buscar {self.item_name}...")
        root.addWidget(search)

        choices = QListWidget()
        for value in self._values:
            item = QListWidgetItem(value)
            item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
            item.setCheckState(Qt.Checked if value in self._selected else Qt.Unchecked)
            choices.addItem(item)
        root.addWidget(choices, 1)

        def filter_choices(text):
            term = (text or "").strip().casefold()
            for index in range(choices.count()):
                item = choices.item(index)
                item.setHidden(bool(term and term not in item.text().casefold()))

        def set_all(state):
            for index in range(choices.count()):
                choices.item(index).setCheckState(state)

        search.textChanged.connect(filter_choices)

        buttons = QHBoxLayout()
        select_all = QPushButton("Seleccionar todas")
        clear = QPushButton("Limpiar")
        cancel = QPushButton("Cancelar")
        apply_button = QPushButton("Aplicar filtro")
        select_all.clicked.connect(lambda: set_all(Qt.Checked))
        clear.clicked.connect(lambda: set_all(Qt.Unchecked))
        cancel.clicked.connect(dialog.reject)
        apply_button.clicked.connect(dialog.accept)
        for secondary_button in (select_all, clear, cancel):
            secondary_button.setObjectName("FilterSecondaryButton")
        set_button_role(apply_button, "report")
        buttons.addWidget(select_all)
        buttons.addWidget(clear)
        buttons.addStretch(1)
        buttons.addWidget(cancel)
        buttons.addWidget(apply_button)
        root.addLayout(buttons)

        search.setFocus()
        if dialog.exec() != QDialog.Accepted:
            return
        self._mode = "exclude" if exclude_radio.isChecked() else "include"
        self._selected = {
            choices.item(index).text()
            for index in range(choices.count())
            if choices.item(index).checkState() == Qt.Checked
        }
        self._refresh_summary()


class MonthPickerButton(QToolButton):
    monthChanged = Signal(int)
    MONTHS = [
        "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
        "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
    ]

    def __init__(self, month=None, parent=None):
        super().__init__(parent)
        self._month = int(month or QDate.currentDate().month())
        self.setPopupMode(QToolButton.InstantPopup)
        menu = QMenu(self)
        for index, label in enumerate(self.MONTHS, 1):
            action = menu.addAction(label)
            action.triggered.connect(lambda _checked=False, value=index: self.set_month(value))
        self.setMenu(menu)
        self._refresh_text()

    def month(self):
        return self._month

    def set_month(self, month):
        month = max(1, min(12, int(month)))
        if month == self._month:
            return
        self._month = month
        self._refresh_text()
        self.monthChanged.emit(month)

    def _refresh_text(self):
        self.setText(self.MONTHS[self._month - 1])


def resolve_period_range(
    period_type,
    year=None,
    period_value=None,
    custom_start=None,
    custom_end=None,
    anchor_date=None,
):
    """Resuelve un período calendario y su período anterior de forma única."""
    period_type = str(period_type or "Mensual")
    today = QDate.currentDate()
    year = int(year or today.year())
    anchor = anchor_date if isinstance(anchor_date, QDate) else today
    start = end = None
    number = None

    if period_type == "Diario":
        start = end = anchor
        label = f"Día {anchor.toString('dd/MM/yyyy')}"
        previous_start = previous_end = anchor.addDays(-1)
        previous_label = previous_start.toString("dd/MM/yyyy")
    elif period_type == "Semanal":
        maximum_week = QDate(year, 12, 28).weekNumber()[0]
        number = max(1, min(maximum_week, int(period_value or anchor.weekNumber()[0])))
        first_week_monday = QDate(year, 1, 4)
        first_week_monday = first_week_monday.addDays(-(first_week_monday.dayOfWeek() - 1))
        start = first_week_monday.addDays((number - 1) * 7)
        end = start.addDays(6)
        label = f"Semana {number} de {year}"
        previous_start, previous_end = start.addDays(-7), end.addDays(-7)
        previous_week, previous_year = previous_start.weekNumber()
        previous_label = f"Semana {previous_week} de {previous_year}"
    elif period_type == "Mensual":
        number = max(1, min(12, int(period_value or today.month())))
        start = QDate(year, number, 1)
        end = start.addMonths(1).addDays(-1)
        label = f"{MonthPickerButton.MONTHS[number - 1]} de {year}"
        previous_start, previous_end = start.addMonths(-1), start.addDays(-1)
        previous_label = (
            f"{MonthPickerButton.MONTHS[previous_start.month() - 1]} de {previous_start.year()}"
        )
    elif period_type == "Trimestral":
        number = max(1, min(4, int(period_value or ((today.month() - 1) // 3 + 1))))
        start = QDate(year, (number - 1) * 3 + 1, 1)
        end = start.addMonths(3).addDays(-1)
        names = ["Primer", "Segundo", "Tercer", "Cuarto"]
        label = f"{names[number - 1]} trimestre de {year}"
        previous_start, previous_end = start.addMonths(-3), start.addDays(-1)
        previous_number = (previous_start.month() - 1) // 3 + 1
        previous_label = f"{names[previous_number - 1]} trimestre de {previous_start.year()}"
    elif period_type == "Semestral":
        number = max(1, min(2, int(period_value or (1 if today.month() <= 6 else 2))))
        start = QDate(year, 1 if number == 1 else 7, 1)
        end = start.addMonths(6).addDays(-1)
        names = ["Primer", "Segundo"]
        label = f"{names[number - 1]} semestre de {year}"
        previous_start, previous_end = start.addMonths(-6), start.addDays(-1)
        previous_number = 1 if previous_start.month() == 1 else 2
        previous_label = f"{names[previous_number - 1]} semestre de {previous_start.year()}"
    elif period_type == "Anual":
        start, end = QDate(year, 1, 1), QDate(year, 12, 31)
        label = f"Año {year}"
        previous_start, previous_end = QDate(year - 1, 1, 1), QDate(year - 1, 12, 31)
        previous_label = f"Año {year - 1}"
    else:
        start = custom_start if isinstance(custom_start, QDate) else today
        end = custom_end if isinstance(custom_end, QDate) else start
        label = "Período personalizado"
        previous_end = start.addDays(-1)
        previous_start = previous_end.addDays(-start.daysTo(end))
        previous_label = (
            f"Período anterior del {previous_start.toString('dd/MM/yyyy')} "
            f"al {previous_end.toString('dd/MM/yyyy')}"
        )

    return {
        "period_type": period_type,
        "period_year": year if period_type in (
            "Semanal", "Mensual", "Trimestral", "Semestral", "Anual"
        ) else start.year(),
        "period_number": number,
        "period_label": label,
        "start_date": start.toString("yyyy-MM-dd"),
        "end_date": end.toString("yyyy-MM-dd"),
        "comparison_label": previous_label,
        "comparison_start": previous_start.toString("yyyy-MM-dd"),
        "comparison_end": previous_end.toString("yyyy-MM-dd"),
    }


class PeriodSelectorWidget(QWidget):
    """Selector reutilizable para panel, reportes y comparaciones."""

    periodChanged = Signal(dict)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._updating = False
        self._layout_mode = None
        self.layout_grid = QGridLayout(self)
        self.layout_grid.setContentsMargins(0, 0, 0, 0)
        self.layout_grid.setHorizontalSpacing(12)
        self.layout_grid.setVerticalSpacing(7)

        self.period_type = QComboBox()
        self.period_type.addItems(
            ["Diario", "Semanal", "Mensual", "Trimestral", "Semestral", "Anual", "Personalizado"]
        )
        self.period_type.setCurrentText("Mensual")
        self.year = QSpinBox()
        self.year.setRange(2000, 2100)
        self.year.setValue(QDate.currentDate().year())
        self.month = QComboBox()
        self.month.addItems(MonthPickerButton.MONTHS)
        self.month.setCurrentIndex(QDate.currentDate().month() - 1)
        self.week = QComboBox()
        self.quarter = QComboBox()
        self.quarter.addItems([
            "Primer trimestre - enero a marzo",
            "Segundo trimestre - abril a junio",
            "Tercer trimestre - julio a septiembre",
            "Cuarto trimestre - octubre a diciembre",
        ])
        self.quarter.setCurrentIndex((QDate.currentDate().month() - 1) // 3)
        self.semester = QComboBox()
        self.semester.addItems([
            "Primer semestre - enero a junio",
            "Segundo semestre - julio a diciembre",
        ])
        self.semester.setCurrentIndex(0 if QDate.currentDate().month() <= 6 else 1)
        self._refresh_week_options()
        for combo in (self.period_type, self.week, self.month, self.quarter, self.semester):
            combo.setSizeAdjustPolicy(QComboBox.AdjustToMinimumContentsLengthWithIcon)
            combo.setMinimumContentsLength(12)
        self.anchor_date = QDateEdit(QDate.currentDate())
        self.custom_from = QDateEdit(QDate.currentDate())
        self.custom_to = QDateEdit(QDate.currentDate())
        for field in (self.anchor_date, self.custom_from, self.custom_to):
            field.setCalendarPopup(True)
            field.setDisplayFormat("dd/MM/yyyy")

        self.summary = QLabel()
        self.summary.setWordWrap(True)
        self.summary.setStyleSheet(
            "background: #F3F7FC; color: #36516E; border-radius: 7px; "
            "padding: 8px 11px; font-weight: 700;"
        )
        self._fields = [
            (QLabel("Tipo de período:"), self.period_type),
            (QLabel("Semana:"), self.week),
            (QLabel("Mes:"), self.month),
            (QLabel("Trimestre:"), self.quarter),
            (QLabel("Semestre:"), self.semester),
            (QLabel("Año:"), self.year),
            (QLabel("Fecha de referencia:"), self.anchor_date),
            (QLabel("Desde:"), self.custom_from),
            (QLabel("Hasta:"), self.custom_to),
        ]
        for _label, field in self._fields:
            field.setMinimumWidth(150)
            field.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        self.period_type.currentTextChanged.connect(self._changed)
        self.year.valueChanged.connect(self._year_changed)
        self.week.currentIndexChanged.connect(self._changed)
        self.month.currentIndexChanged.connect(self._changed)
        self.quarter.currentIndexChanged.connect(self._changed)
        self.semester.currentIndexChanged.connect(self._changed)
        self.anchor_date.dateChanged.connect(self._changed)
        self.custom_from.dateChanged.connect(self._changed)
        self.custom_to.dateChanged.connect(self._changed)
        self._changed()

    def definition(self):
        period_type = self.period_type.currentText()
        value = None
        if period_type == "Mensual":
            value = self.month.currentIndex() + 1
        elif period_type == "Semanal":
            value = self.week.currentData() or self.week.currentIndex() + 1
        elif period_type == "Trimestral":
            value = self.quarter.currentIndex() + 1
        elif period_type == "Semestral":
            value = self.semester.currentIndex() + 1
        return resolve_period_range(
            period_type,
            year=self.year.value(),
            period_value=value,
            custom_start=self.custom_from.date(),
            custom_end=self.custom_to.date(),
            anchor_date=self.anchor_date.date(),
        )

    def _changed(self, *_args):
        if self._updating:
            return
        period_type = self.period_type.currentText()
        visible_widgets = {self.period_type}
        if period_type == "Diario":
            visible_widgets.add(self.anchor_date)
        elif period_type == "Semanal":
            visible_widgets.update((self.week, self.year))
        elif period_type == "Mensual":
            visible_widgets.update((self.year, self.month))
        elif period_type == "Trimestral":
            visible_widgets.update((self.year, self.quarter))
        elif period_type == "Semestral":
            visible_widgets.update((self.year, self.semester))
        elif period_type == "Anual":
            visible_widgets.add(self.year)
        else:
            visible_widgets.update((self.custom_from, self.custom_to))

        visible_fields = [(label, field) for label, field in self._fields if field in visible_widgets]
        self._visible_fields = visible_fields
        self._arrange_fields()
        definition = self.definition()
        display_start = QDate.fromString(definition["start_date"], "yyyy-MM-dd").toString("dd/MM/yyyy")
        display_end = QDate.fromString(definition["end_date"], "yyyy-MM-dd").toString("dd/MM/yyyy")
        self.summary.setText(
            f"{definition['period_label']}   ·   "
            f"Fechas aplicadas: {display_start} al {display_end}"
        )
        self.periodChanged.emit(definition)

    def _year_changed(self, *_args):
        self._refresh_week_options()
        self._changed()

    def _refresh_week_options(self):
        year = self.year.value()
        previous_week = self.week.currentData() if self.week.count() else None
        current_date = QDate.currentDate()
        preferred = (
            current_date.weekNumber()[0] if year == current_date.weekNumber()[1]
            else int(previous_week or 1)
        )
        maximum_week = QDate(year, 12, 28).weekNumber()[0]
        first_monday = QDate(year, 1, 4)
        first_monday = first_monday.addDays(-(first_monday.dayOfWeek() - 1))
        self.week.blockSignals(True)
        self.week.clear()
        for week_number in range(1, maximum_week + 1):
            start = first_monday.addDays((week_number - 1) * 7)
            end = start.addDays(6)
            start_month = MonthPickerButton.MONTHS[start.month() - 1].lower()
            end_month = MonthPickerButton.MONTHS[end.month() - 1].lower()
            if start.month() == end.month():
                dates = f"{start.day()} al {end.day()} de {end_month}"
            else:
                dates = f"{start.day()} de {start_month} al {end.day()} de {end_month}"
            self.week.addItem(f"Semana {week_number} - {dates}", week_number)
        self.week.setCurrentIndex(max(0, min(maximum_week, preferred) - 1))
        self.week.blockSignals(False)

    def _arrange_fields(self):
        visible_fields = getattr(self, "_visible_fields", [])
        for label, field in self._fields:
            self.layout_grid.removeWidget(label)
            self.layout_grid.removeWidget(field)
            visible = any(field is current for _text, current in visible_fields)
            label.setVisible(visible)
            field.setVisible(visible)
        self.layout_grid.removeWidget(self.summary)
        for column in range(4):
            self.layout_grid.setColumnStretch(column, 0)

        width = self.width()
        if width < 480:
            mode = "stacked"
        elif width < 700 and len(visible_fields) == 3:
            mode = "two_rows"
        else:
            mode = "horizontal"
        self._layout_mode = mode

        if mode == "stacked":
            row = 0
            for label, field in visible_fields:
                self.layout_grid.addWidget(label, row, 0)
                self.layout_grid.addWidget(field, row + 1, 0)
                row += 2
            summary_row, summary_span = row, 1
            self.layout_grid.setColumnStretch(0, 1)
        elif mode == "two_rows":
            for column, (label, field) in enumerate(visible_fields[:2]):
                self.layout_grid.addWidget(label, 0, column)
                self.layout_grid.addWidget(field, 1, column)
                self.layout_grid.setColumnStretch(column, 1)
            label, field = visible_fields[2]
            self.layout_grid.addWidget(label, 2, 0, 1, 2)
            self.layout_grid.addWidget(field, 3, 0, 1, 2)
            summary_row, summary_span = 4, 2
        else:
            for column, (label, field) in enumerate(visible_fields):
                self.layout_grid.addWidget(label, 0, column)
                self.layout_grid.addWidget(field, 1, column)
                self.layout_grid.setColumnStretch(column, 1)
            summary_row, summary_span = 2, max(1, len(visible_fields))
        self.layout_grid.addWidget(self.summary, summary_row, 0, 1, summary_span)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        previous_mode = self._layout_mode
        width = event.size().width()
        new_mode = "stacked" if width < 480 else (
            "two_rows" if width < 700 and len(getattr(self, "_visible_fields", [])) == 3
            else "horizontal"
        )
        if new_mode != previous_mode:
            self._arrange_fields()


class ComparisonPdfDialog(QDialog):
    """Vista previa, guardado e impresión del PDF comparativo ya generado."""

    def __init__(self, pdf_path, parent=None, dialog_title="Reporte comparativo listo", detail_text=None):
        super().__init__(parent)
        self.pdf_path = pdf_path
        self.setWindowTitle("Vista previa de la comparación")
        self.setMinimumWidth(620)
        root = QVBoxLayout(self)
        title = QLabel(dialog_title)
        title.setStyleSheet("font-size: 16pt; font-weight: 900; color: #123F83;")
        detail = QLabel(detail_text or (
            "El documento se generó con la misma información que está visible en el panel. "
            "Revísalo antes de imprimir o guarda una copia en otra ubicación."
        ))
        detail.setWordWrap(True)
        detail.setStyleSheet(
            "background: #F3F7FC; color: #36516E; padding: 12px; border-radius: 8px;"
        )
        filename = QLabel(os.path.basename(pdf_path))
        filename.setWordWrap(True)
        filename.setStyleSheet("font-weight: 800; color: #31475F; padding: 6px 2px;")
        root.addWidget(title)
        root.addWidget(detail)
        root.addWidget(filename)
        buttons = QHBoxLayout()
        preview = QPushButton("Vista previa")
        save = QPushButton("Guardar PDF")
        print_button = QPushButton("Imprimir")
        close = QPushButton("Cerrar")
        set_button_role(preview, "report")
        set_button_role(save, "success")
        set_button_role(print_button, "report")
        set_button_role(close, "neutral")
        buttons.addWidget(preview)
        buttons.addWidget(save)
        buttons.addWidget(print_button)
        buttons.addStretch(1)
        buttons.addWidget(close)
        root.addLayout(buttons)
        preview.clicked.connect(self.open_preview)
        save.clicked.connect(self.save_copy)
        print_button.clicked.connect(self.print_pdf)
        close.clicked.connect(self.accept)

    def _printer(self):
        printer = QPrinter(QPrinter.HighResolution)
        printer.setPageOrientation(QPageLayout.Landscape)
        return printer

    def _paint_pdf(self, printer):
        document = QPdfDocument(self)
        document.load(self.pdf_path)
        if document.pageCount() <= 0:
            return
        painter = QPainter(printer)
        try:
            for page_index in range(document.pageCount()):
                if page_index:
                    printer.newPage()
                page_rect = printer.pageRect(QPrinter.Unit.DevicePixel)
                page_size = QSize(int(page_rect.width()), int(page_rect.height()))
                target = QSize(page_size)
                target.scale(QSize(2400, 2400), Qt.KeepAspectRatio)
                image = document.render(page_index, target)
                scaled = image.size().scaled(page_size, Qt.KeepAspectRatio)
                x = page_rect.x() + (page_rect.width() - scaled.width()) // 2
                y = page_rect.y() + (page_rect.height() - scaled.height()) // 2
                painter.drawImage(QRectF(x, y, scaled.width(), scaled.height()), image)
        finally:
            painter.end()

    def open_preview(self):
        printer = self._printer()
        dialog = QPrintPreviewDialog(printer, self)
        dialog.setWindowTitle("Vista previa - Reporte comparativo")
        dialog.resize(1100, 760)
        dialog.paintRequested.connect(self._paint_pdf)
        dialog.exec()

    def print_pdf(self):
        printer = self._printer()
        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QDialog.Accepted:
            self._paint_pdf(printer)

    def save_copy(self):
        default_name = os.path.basename(self.pdf_path)
        destination, _ = QFileDialog.getSaveFileName(
            self, "Guardar reporte comparativo", default_name, "PDF (*.pdf)"
        )
        if not destination:
            return
        if not destination.lower().endswith(".pdf"):
            destination += ".pdf"
        try:
            shutil.copy2(self.pdf_path, destination)
            FloatingToast("Copia del reporte guardada", self).show()
        except Exception as exc:
            QMessageBox.critical(self, "Guardar PDF", f"No se pudo guardar la copia:\n{exc}")

def _clean_name(s: str):
    s = (s or '').strip()
    if '$' in s: s = s.rsplit('$', 1)[0]
    s = re.sub(r'[\s\-\–\—\|\:]+$', '', s).strip()
    s = re.sub(r'\s+', ' ', s)
    return s


CATALOG_SPELLING_CORRECTIONS = {
    "ACIDO": "ÁCIDO",
    "ACETILISTEINA": "ACETILCISTEÍNA",
    "ACETLCISTEINA": "ACETILCISTEÍNA",
    "AMITROTILINA": "AMITRIPTILINA",
    "AMOXICIINA": "AMOXICILINA",
    "CALVULAMICO": "CLAVULÁNICO",
    "CANULA": "CÁNULA",
    "CATETER": "CATÉTER",
    "CARVELIDOL": "CARVEDILOL",
    "CEFEPINE": "CEFEPIMA",
    "CEFTRIAZONA": "CEFTRIAXONA",
    "CIPROFLAXACINA": "CIPROFLOXACINA",
    "DEXAMETOZONA": "DEXAMETASONA",
    "DIAZEPAN": "DIAZEPAM",
    "DICLOXOCILINA": "DICLOXACILINA",
    "EPINEFINA": "EPINEFRINA",
    "ESTERIL": "ESTÉRIL",
    "ERITROPROYECTINA": "ERITROPOYETINA",
    "FENTANYLO": "FENTANILO",
    "GANMAGLOBULINA": "GAMMAGLOBULINA",
    "IBERSARTAN": "IRBESARTÁN",
    "IBUPREFENO": "IBUPROFENO",
    "IMIPENEN": "IMIPENEM",
    "KETEROLACO": "KETOROLACO",
    "LEVETITACETAM": "LEVETIRACETAM",
    "MEROPENEN": "MEROPENEM",
    "MIDAZOLAN": "MIDAZOLAM",
    "OMEPREZOL": "OMEPRAZOL",
    "OXIGENO": "OXÍGENO",
    "PEDIATRICO": "PEDIÁTRICO",
    "QUIROFANO": "QUIRÓFANO",
    "RISPERIDENA": "RISPERIDONA",
    "SOLUCION DESTROSA": "SOLUCIÓN DEXTROSA",
    "SOLUCION": "SOLUCIÓN",
    "SULFATO DE MAGENSIO": "SULFATO DE MAGNESIO",
    "SURFARCTANTE": "SURFACTANTE",
    "SEVOFLUXANO": "SEVOFLURANO",
    "TRIMETROPIN": "TRIMETOPRIM",
    "VIAS": "VÍAS",
}


def normalize_catalog_name(raw_name: str) -> str:
    """Normaliza nombres importados sin alterar dosis, tamaños ni marcas."""
    name = str(raw_name or "").strip().upper()
    name = re.sub(r"^[\s•·●▪◦*\-–—]+", "", name)
    name = name.replace("“", '"').replace("”", '"').replace("''", '""')
    name = re.sub(r"\s+", " ", name)
    name = re.sub(r"\(\s+", "(", name)
    name = re.sub(r"\s+\)", ")", name)
    name = re.sub(r"\s*/\s*", "/", name)
    for incorrect, correct in CATALOG_SPELLING_CORRECTIONS.items():
        name = re.sub(rf"\b{re.escape(incorrect)}\b", correct, name)
    return name.strip(" -–—,.;")


def catalog_identity_key(raw_name: str) -> str:
    """Compara ítems ignorando viñetas, acentos y espacios de formato."""
    normalized = remove_accents(normalize_catalog_name(raw_name)).casefold()
    return re.sub(r"[^a-z0-9]+", "", normalized)


def _ditto_base(previous_name: str) -> str:
    """Obtiene el nombre base para líneas Word que usan comillas de repetición."""
    base = normalize_catalog_name(previous_name)
    base = re.sub(r"\s+(?:#\s*)?\d+(?:[.,]\d+)?\s*$", "", base)
    return base.strip()


def _parse_word_price_token(token: str):
    """Interpreta separadores dominicanos sin confundir 3.400 con 3.40."""
    value = str(token or "").strip().replace(" ", "")
    if re.fullmatch(r"\d{1,3}(?:\.\d{3})+", value):
        value = value.replace(".", "")
    return _parse_price(value)


def parse_universal_catalog_line(raw_text: str, previous_name: str = "") -> list[tuple[str, float]]:
    """Convierte una línea Word en uno o varios ítems independientes."""
    text = str(raw_text or "").strip()
    if not text:
        return []
    text = text.replace("RD$", "$")
    ditto = bool(re.match(r'^\s*(?:""|″|〃)', text))
    if ditto:
        text = re.sub(r'^\s*(?:""|″|〃)\s*', f"{_ditto_base(previous_name)} ", text, count=1)

    prices = list(re.finditer(r"\$\s*([0-9][0-9.,]*)", text))
    if not prices:
        return []

    upper = remove_accents(text).upper()
    is_multi_sonda = "SONDA" in upper and "VIA" in upper and len(prices) > 1
    if is_multi_sonda:
        first_prefix = text[:prices[0].start()]
        base_match = re.match(r"^(.*?\b(?:2|3)\s*V[IÍ]AS?)\b(.*)$", first_prefix, re.I)
        base = normalize_catalog_name(base_match.group(1) if base_match else first_prefix)
        variants = []
        prior_end = len(base_match.group(1)) if base_match else 0
        for index, match in enumerate(prices):
            segment_start = prior_end if index == 0 else prices[index - 1].end()
            variant_text = text[segment_start:match.start()]
            size = re.findall(r"#?\s*(\d+(?:[.,]\d+)?)", variant_text)
            if not size:
                continue
            price = _parse_word_price_token(match.group(1))
            if price is not None:
                variants.append((f"{base} (#{size[-1]})", float(price)))
        return variants

    # Explicaciones como "equivale a $ 2.60" no son el precio del ítem;
    # el precio final de la línea es el que se importa.
    match = prices[-1]
    price = _parse_word_price_token(match.group(1))
    if price is None:
        return []
    before = text[:match.start()].strip()
    after = text[match.end():].strip(" ,;.-")
    if after:
        before = f"{before} {after}"
    name = normalize_catalog_name(_clean_name(before))
    return [(name, float(price))] if name else []

def open_file_path(path: str) -> bool:
    try:
        if sys.platform.startswith("win"): os.startfile(path)
        elif sys.platform == "darwin": os.system(f'open "{path}"')
        else: os.system(f'xdg-open "{path}"')
        return True
    except Exception:
        return False

def stable_storage_path(folder: str, stored_path_or_filename: str) -> str:
    """Devuelve una ruta estable en la carpeta actual de la app usando solo el nombre del archivo.
    Esto evita errores cuando un PDF fue creado desde una carpeta temporal de PyInstaller (_MEI...).
    """
    filename = os.path.basename(stored_path_or_filename or "")
    if not filename:
        return ""
    return os.path.join(folder, filename)

_MONTHS_ES = r"(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|setiembre|octubre|noviembre|diciembre)"
_MONTH_LINE_RE = re.compile(rf"^\s*\(?\s*{_MONTHS_ES}\s*\)?(?:\s*(?:de)?\s*\d{{4}})?\s*$", re.IGNORECASE)

def _looks_like_month_or_month_year(txt: str) -> bool:
    return bool(txt and _MONTH_LINE_RE.match(txt.strip()))

def user_can_manage_catalog(user: dict) -> bool:
    return bool(user) and user.get("role") in (ROLE_AUX, ROLE_ADMIN, ROLE_AUDIT)

def normalize_role(role) -> str:
    normalized = remove_accents(str(role or "")).strip().lower()
    if normalized in {"admin", "administrator", "administrador"}:
        return ROLE_ADMIN
    return normalized

def is_administrator(user: dict) -> bool:
    return bool(user) and normalize_role(user.get("role")) == ROLE_ADMIN

def user_is_admin(user: dict) -> bool:
    return is_administrator(user)

def user_can_manage_sessions(user: dict) -> bool:
    return is_administrator(user)

def user_can_delete_receipts(user: dict) -> bool:
    return bool(user) and user.get("role") in (ROLE_ADMIN, ROLE_AUDIT)

def wrap_text(text: str, width: int = 76):
    if not text: return [""]
    words = text.split(); lines, line = [], ""
    for w in words:
        if len(w) > width:
            if line: lines.append(line); line = ""
            for i in range(0, len(w), width):
                seg = w[i:i + width]
                if len(seg) == width: lines.append(seg)
                else: line = seg
            continue
        if not line: line = w
        elif len(line) + 1 + len(w) <= width: line += " " + w
        else: lines.append(line); line = w
    if line: lines.append(line)
    return lines

# =========================================================
# COMPONENTE TOAST Y GRÁFICO (DASHBOARD)
# =========================================================

class FloatingToast(QWidget):
    def __init__(self, text, parent=None, is_error=False):
        super().__init__(parent)
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.Tool | Qt.WindowStaysOnTopHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setAttribute(Qt.WA_ShowWithoutActivating)
        
        lay = QVBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        self.lbl = QLabel(text)
        
        bg_color = "#d32f2f" if is_error else "#2e7d32"
        text_color = "#ffffff"
        
        self.lbl.setStyleSheet(f"""
            QLabel {{
                background-color: {bg_color};
                color: {text_color};
                padding: 12px 24px;
                border-radius: 8px;
                font-size: 11pt;
                font-weight: bold;
                border: 1px solid {"#b71c1c" if is_error else "#1b5e20"};
            }}
        """)
        lay.addWidget(self.lbl)
        self.adjustSize()
        
        if parent:
            parent_geo = parent.geometry()
            x = parent_geo.x() + (parent_geo.width() - self.width()) // 2
            y = parent_geo.bottom() - 100
            self.move(x, y)
            
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.close)
        self.timer.setSingleShot(True)
        self.timer.start(2500)

class SimpleBarChart(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.data = {}
        
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed) 
        self.setFixedHeight(250) 
        
    def sizeHint(self):
        return QSize(800, 250)

    def minimumSizeHint(self):
        return QSize(300, 250)
        
    def set_data(self, json_str):
        try:
            self.data = json.loads(json_str)
            self.update()
        except Exception:
            self.data = {}
            self.update()
            
    def paintEvent(self, event):
        if not self.data: return
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        max_val = max([v for k, v in self.data.items() if k != "Total General"] + [1])
        
        margin_left = 60
        margin_bottom = 40
        w = self.width() - margin_left - 10
        h = self.height() - margin_bottom - 20
        
        # Ejes
        painter.setPen(QPen(QColor("#888888"), 2))
        painter.drawLine(margin_left, 20, margin_left, h + 20)
        painter.drawLine(margin_left, h + 20, self.width() - 10, h + 20)
        
        keys = [k for k in self.data.keys() if k != "Total General" and self.data[k] > 0]
        if not keys: return
        
        bar_w = w / len(keys) * 0.6
        spacing = w / len(keys)
        
        font = painter.font()
        font.setPointSize(9)
        font.setBold(True)
        painter.setFont(font)
        
        for i, k in enumerate(keys):
            val = self.data[k]
            bar_h = (val / max_val) * h if max_val > 0 else 0
            x = margin_left + (i * spacing) + (spacing - bar_w) / 2
            y = 20 + h - bar_h
            
            # Barra
            color = QColor(CAT_COLORS.get(k, "#1e88e5"))
            painter.setBrush(QBrush(color))
            painter.setPen(Qt.NoPen)
            painter.drawRect(int(x), int(y), int(bar_w), int(bar_h))
            
            # Etiqueta
            painter.setPen(QColor("#555555"))
            label = k[:5] + "." if len(k) > 6 else k
            painter.drawText(int(x), int(h + 35), label)
            
            # Valor encima
            painter.setPen(QColor("#2e7d32"))
            painter.drawText(int(x), int(y - 5), f"${val:,.0f}")


class ModernBarChart(QWidget):
    """Gráfico liviano para el panel, sin dependencias externas."""

    def __init__(self, accent="#174A96", parent=None):
        super().__init__(parent)
        self.entries = []
        self.accent = QColor(accent)
        self.currency = True
        self.percent = False
        self.show_trend = False
        self.setMinimumHeight(245)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

    def set_entries(self, entries, currency=True, show_trend=False, percent=False):
        self.entries = [(str(label), float(value or 0)) for label, value in entries]
        self.currency = bool(currency)
        self.percent = bool(percent)
        self.show_trend = bool(show_trend)
        self.update()

    def _value_text(self, value):
        if self.percent:
            return f"{value:.1f}%"
        if self.currency:
            if abs(value) >= 1_000_000:
                return f"RD$ {value / 1_000_000:.1f}M"
            if abs(value) >= 1_000:
                return f"RD$ {value / 1_000:.1f}K"
            return f"RD$ {value:,.0f}"
        return f"{int(value):,}"

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.fillRect(self.rect(), QColor("#FFFFFF"))

        if not self.entries or max((value for _, value in self.entries), default=0) <= 0:
            painter.setPen(QColor("#77869A"))
            painter.drawText(self.rect(), Qt.AlignCenter, "No hay datos para los filtros seleccionados")
            return

        left, top, right, bottom = 52, 24, 18, 52
        chart_w = max(1, self.width() - left - right)
        chart_h = max(1, self.height() - top - bottom)
        max_value = max(value for _, value in self.entries) or 1

        painter.setPen(QPen(QColor("#D9E4F2"), 1))
        for line in range(5):
            y = top + chart_h * line / 4
            painter.drawLine(left, int(y), left + chart_w, int(y))

        count = len(self.entries)
        slot = chart_w / max(1, count)
        bar_width = max(12.0, min(48.0, slot * 0.56))
        trend_points = []
        for index, (label, value) in enumerate(self.entries):
            height = chart_h * value / max_value
            x = left + index * slot + (slot - bar_width) / 2
            y = top + chart_h - height
            rect = QRectF(x, y, bar_width, height)

            color = QColor(CAT_COLORS.get(label, self.accent.name()))
            gradient_color = QColor(color)
            gradient_color.setAlpha(215)
            painter.setBrush(QBrush(gradient_color))
            painter.setPen(Qt.NoPen)
            painter.drawRoundedRect(rect, 4, 4)
            trend_points.append(QPointF(x + bar_width / 2, y))

            painter.setPen(QColor("#23344A"))
            font = painter.font(); font.setPointSize(8); font.setBold(True); painter.setFont(font)
            value_rect = QRectF(x - slot * .18, max(0, y - 24), bar_width + slot * .36, 20)
            painter.drawText(value_rect, Qt.AlignCenter, self._value_text(value))

            short_label = label if len(label) <= 12 else label[:11] + "…"
            label_rect = QRectF(left + index * slot, top + chart_h + 8, slot, 34)
            painter.setPen(QColor("#52657A"))
            font.setPointSize(8); font.setBold(False); painter.setFont(font)
            painter.drawText(label_rect, Qt.AlignHCenter | Qt.AlignTop | Qt.TextWordWrap, short_label)

        if self.show_trend and len(trend_points) > 1:
            painter.setPen(QPen(QColor("#F28C28"), 2.5))
            for first, second in zip(trend_points, trend_points[1:]):
                painter.drawLine(first, second)
            painter.setBrush(QColor("#F28C28"))
            painter.setPen(Qt.NoPen)
            for point in trend_points:
                painter.drawEllipse(point, 3.5, 3.5)


class ModernLineChart(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.entries = []
        self.currency = True
        self._hover_points = []
        self.setMouseTracking(True)
        self.setMinimumHeight(250)

    def set_entries(self, entries, currency=True):
        self.entries = [(str(label), float(value or 0)) for label, value in entries]
        self.currency = bool(currency)
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.fillRect(self.rect(), QColor("#FFFFFF"))
        if not self.entries or max((value for _, value in self.entries), default=0) <= 0:
            painter.setPen(QColor("#77869A"))
            painter.drawText(self.rect(), Qt.AlignCenter, "No hay datos para mostrar")
            return
        left, top, right, bottom = 84, 28, 20, 50
        width = max(1, self.width() - left - right)
        height = max(1, self.height() - top - bottom)
        maximum = max(value for _, value in self.entries) or 1
        painter.setPen(QPen(QColor("#DFE8F3"), 1))
        for line in range(5):
            y = top + height * line / 4
            painter.drawLine(left, int(y), left + width, int(y))
            scale_value = maximum * (4 - line) / 4
            shown = (
                f"RD$ {scale_value / 1_000_000:.1f}M" if self.currency and scale_value >= 1_000_000
                else f"RD$ {scale_value / 1_000:.0f}K" if self.currency and scale_value >= 1_000
                else f"RD$ {scale_value:,.0f}" if self.currency
                else f"{int(scale_value):,}"
            )
            painter.setPen(QColor("#687B91"))
            font = painter.font(); font.setPointSize(7); painter.setFont(font)
            painter.drawText(QRectF(2, y - 9, left - 10, 18), Qt.AlignRight | Qt.AlignVCenter, shown)
            painter.setPen(QPen(QColor("#DFE8F3"), 1))
        divisor = max(1, len(self.entries) - 1)
        points = []
        for index, (_label, value) in enumerate(self.entries):
            x = left + width * index / divisor
            y = top + height - height * value / maximum
            points.append(QPointF(x, y))
        self._hover_points = [
            (point, label, value) for point, (label, value) in zip(points, self.entries)
        ]
        painter.setPen(QPen(QColor("#F28C28"), 3))
        for first, second in zip(points, points[1:]):
            painter.drawLine(first, second)
        value_step = max(1, (len(points) + 7) // 8)
        date_step = 1 if len(points) <= 31 else value_step
        for index, (point, (label, value)) in enumerate(zip(points, self.entries)):
            painter.setBrush(QColor("#FFFFFF"))
            painter.setPen(QPen(QColor("#174A96"), 3))
            painter.drawEllipse(point, 4, 4)
            if index % date_step == 0 or index == len(points) - 1:
                painter.setPen(QColor("#52657A"))
                font = painter.font(); font.setPointSize(8); painter.setFont(font)
                date_label = label[-2:] if len(points) <= 31 and len(label) >= 10 else label[:13]
                painter.drawText(QRectF(point.x() - 22, top + height + 10, 44, 25), Qt.AlignHCenter | Qt.AlignTop, date_label)
            if index % value_step == 0 or index == len(points) - 1:
                painter.setPen(QColor("#52657A"))
                font = painter.font(); font.setPointSize(8); painter.setFont(font)
                shown = f"RD$ {value:,.0f}" if self.currency else f"{int(value):,}"
                font.setBold(True); painter.setFont(font)
                painter.drawText(QRectF(point.x() - 55, max(0, point.y() - 28), 110, 20), Qt.AlignCenter, shown)

    def mouseMoveEvent(self, event):
        if not self._hover_points:
            return super().mouseMoveEvent(event)
        point, label, value = min(
            self._hover_points, key=lambda item: abs(item[0].x() - event.position().x())
        )
        if abs(point.x() - event.position().x()) <= 14:
            shown = f"RD$ {value:,.2f}" if self.currency else f"{int(value):,}"
            QToolTip.showText(event.globalPosition().toPoint(), f"{label}\n{shown}", self)
        else:
            QToolTip.hideText()
        super().mouseMoveEvent(event)


class HorizontalBarChart(QWidget):
    """Barras horizontales legibles incluso con muchas ARS."""

    def __init__(self, accent="#174A96", parent=None):
        super().__init__(parent)
        self.entries = []
        self.accent = QColor(accent)
        self.percent = True
        self.setMinimumHeight(245)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

    def set_entries(self, entries, percent=True):
        self.entries = [(str(label), float(value or 0)) for label, value in entries]
        self.percent = bool(percent)
        self.setMinimumHeight(max(245, 52 + len(self.entries) * 31))
        self.updateGeometry()
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.fillRect(self.rect(), QColor("#FFFFFF"))
        if not self.entries or max((value for _, value in self.entries), default=0) <= 0:
            painter.setPen(QColor("#77869A"))
            painter.drawText(self.rect(), Qt.AlignCenter, "No hay datos para los filtros seleccionados")
            return
        label_width, right, top = min(190, max(120, self.width() // 4)), 70, 18
        chart_left = label_width + 18
        chart_width = max(1, self.width() - chart_left - right)
        maximum = max(value for _, value in self.entries) or 1
        row_height = max(25, min(34, (self.height() - top * 2) / len(self.entries)))
        font = painter.font(); font.setPointSize(8); painter.setFont(font)
        for index, (label, value) in enumerate(self.entries):
            y = top + index * row_height
            painter.setPen(QColor("#334A62"))
            painter.drawText(QRectF(6, y, label_width, row_height - 5), Qt.AlignRight | Qt.AlignVCenter, label)
            track = QRectF(chart_left, y + 5, chart_width, max(10, row_height - 15))
            painter.setPen(Qt.NoPen)
            painter.setBrush(QColor("#E8EEF6"))
            painter.drawRoundedRect(track, 5, 5)
            fill = QRectF(track.x(), track.y(), track.width() * value / maximum, track.height())
            painter.setBrush(self.accent)
            painter.drawRoundedRect(fill, 5, 5)
            painter.setPen(QColor("#23344A"))
            font.setBold(True); painter.setFont(font)
            value_text = f"{value:.1f}%" if self.percent else f"{value:,.0f}"
            painter.drawText(QRectF(chart_left + chart_width + 8, y, right - 12, row_height - 5), Qt.AlignLeft | Qt.AlignVCenter, value_text)


class DonutChart(QWidget):
    COLORS = ["#174A96", "#0B7A5A", "#6B35C8", "#EA6A24", "#D14B4B", "#2E91C7", "#8A6D3B", "#607D8B"]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.entries = []
        self.setMinimumHeight(250)

    def set_entries(self, entries):
        self.entries = [(str(label), float(value or 0)) for label, value in entries if float(value or 0) > 0][:8]
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.fillRect(self.rect(), QColor("#FFFFFF"))
        total = sum(value for _, value in self.entries)
        if total <= 0:
            painter.setPen(QColor("#77869A"))
            painter.drawText(self.rect(), Qt.AlignCenter, "No hay datos para mostrar")
            return
        diameter = min(170, self.height() - 56, max(100, self.width() // 3))
        donut_rect = QRectF(28, (self.height() - diameter) / 2, diameter, diameter)
        start_angle = 90 * 16
        for index, (_label, value) in enumerate(self.entries):
            span = -int(value / total * 360 * 16)
            painter.setPen(QPen(QColor(self.COLORS[index]), 28, Qt.SolidLine, Qt.FlatCap))
            inset = donut_rect.adjusted(14, 14, -14, -14)
            painter.drawArc(inset, start_angle, span)
            start_angle += span
        center = donut_rect.center()
        painter.setPen(QColor("#52657A"))
        font = painter.font(); font.setPointSize(8); font.setBold(True); painter.setFont(font)
        painter.drawText(QRectF(center.x() - 55, center.y() - 25, 110, 20), Qt.AlignCenter, "TOTAL")
        painter.setPen(QColor("#123F83"))
        font.setPointSize(11); painter.setFont(font)
        painter.drawText(QRectF(center.x() - 65, center.y() - 5, 130, 28), Qt.AlignCenter, f"RD$ {total:,.0f}")
        legend_x = diameter + 65
        legend_width = max(100, self.width() - legend_x - 20)
        for index, (label, value) in enumerate(self.entries):
            y = 30 + index * 27
            painter.setBrush(QColor(self.COLORS[index]))
            painter.setPen(Qt.NoPen)
            painter.drawRoundedRect(QRectF(legend_x, y, 13, 13), 3, 3)
            painter.setPen(QColor("#334A62"))
            font.setPointSize(8); font.setBold(False); painter.setFont(font)
            painter.drawText(QRectF(legend_x + 22, y - 3, legend_width - 70, 20), Qt.AlignLeft | Qt.AlignVCenter, label[:28])
            font.setBold(True); painter.setFont(font)
            painter.drawText(QRectF(self.width() - 78, y - 3, 62, 20), Qt.AlignRight | Qt.AlignVCenter, f"{value / total * 100:.1f}%")

# =========================================================
# CATEGORY COLORS FOR PDF (Platypus) - UX/UI MODERNIZED
# =========================================================
CAT_PDF_COLORS = {
    "Medicamentos": HexColor("#1A56DB"),  # Azul fuerte
    "Materiales": HexColor("#059669"),    # Verde
    "Laboratorios": HexColor("#7C3AED"),  # Morado
    "Imágenes": HexColor("#EA580C"),      # Naranja
    "Procedimientos": HexColor("#2563EB"),# Azul medio
    "Honorarios": HexColor("#1D4ED8"),    # Azul oscuro
}

# =========================================================
# PARAGRAPH STYLES FOR PDF
# =========================================================
pdf_styles = {
    "header_title": ParagraphStyle(
        "header_title",
        fontSize=28,
        leading=32,
        alignment=TA_LEFT,
        textColor=HexColor("#123F91"),
        spaceAfter=4,
    ),
    "header_subtitle": ParagraphStyle(
        "header_subtitle",
        fontSize=16,
        leading=20,
        alignment=TA_LEFT,
        textColor=HexColor("#123F91"),
        spaceAfter=8,
    ),
    "receipt_label": ParagraphStyle(
        "receipt_label",
        fontSize=10,
        leading=12,
        alignment=TA_LEFT,
        textColor=HexColor("#123F91"),
    ),
    "patient_section_title": ParagraphStyle(
        "patient_section_title",
        fontSize=14,
        leading=18,
        textColor=HexColor("#123F91"),
        spaceAfter=2,
    ),
    "patient_label": ParagraphStyle(
        "patient_label",
        fontSize=11,
        leading=14,
        textColor=HexColor("#6B7280"),
        spaceAfter=2,
    ),
    "patient_value": ParagraphStyle(
        "patient_value",
        fontSize=12,
        leading=15,
        textColor=HexColor("#222222"),
        spaceAfter=4,
    ),
    "category_title": ParagraphStyle(
        "category_title",
        fontSize=14,
        leading=18,
        textColor=HexColor("#123F91"),
        spaceAfter=4,
    ),
    "item_description": ParagraphStyle(
        "item_description",
        fontSize=10,
        leading=12,
        textColor=HexColor("#222222"),
    ),
    "subtotal": ParagraphStyle(
        "subtotal",
        fontSize=11,
        leading=13,
        textColor=white,
    ),
    "total_label": ParagraphStyle(
        "total_label",
        fontSize=14,
        leading=18,
        textColor=HexColor("#123F91"),
    ),
    "total_value": ParagraphStyle(
        "total_value",
        fontSize=48,
        leading=52,
        textColor=HexColor("#123F91"),
    ),
    "footer": ParagraphStyle(
        "footer",
        fontSize=10,
        leading=13,
        textColor=HexColor("#6B7280"),
    ),
}

# =========================================================
# PDF HELPER FUNCTIONS (Platypus)
# =========================================================
def draw_icon_text(c, x, y, icon, size=12, color=HexColor("#0F3D91")):
    """Intenta dibujar un icono usando fuentes que soporten emojis (Windows: Segoe UI Emoji, Linux: Noto Color Emoji)."""
    try:
        for fontname in ["Segoe UI Emoji", "Arial Unicode MS", "Noto Color Emoji", "Apple Color Emoji", "DejaVu Sans"]:
            try:
                c.setFont(fontname, size)
                break
            except:
                continue
        else:
            c.setFont("Helvetica", size)
    except:
        c.setFont("Helvetica", size)
    c.setFillColor(color)
    c.drawString(x, y, icon)


def draw_header(c, page_w, page_h, recibo_number, date_str):
    """Draw PDF header with logo, title, and modern receipt info box."""
    y_pos = page_h - 0.6 * inch

    # Logo
    if LOGO_PATH and os.path.exists(LOGO_PATH):
        try:
            c.drawImage(LOGO_PATH, 0.5*inch, y_pos - 0.6*inch, width=0.7*inch, height=0.7*inch, preserveAspectRatio=True, mask='auto')
        except Exception:
            pass

    # Título principal en dos líneas
    c.setFont("Helvetica-Bold", 22)
    c.setFillColor(HexColor("#0F3D91"))
    c.drawString(1.3 * inch, y_pos - 0.1 * inch, "HOSPITAL PROVINCIAL")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(1.3 * inch, y_pos - 0.35 * inch, "DR. ÁNGEL CONTRERAS MEJÍA")
    c.setFont("Helvetica-Bold", 14)
    c.drawString(1.3 * inch, y_pos - 0.58 * inch, "DETALLE DE FACTURACIÓN DE EMERGENCIA")

    # Tarjeta de Recibo (Derecha)
    card_w = 2.2 * inch
    card_h = 0.7 * inch
    card_x = page_w - 0.5 * inch - card_w
    card_y = y_pos - card_h + 0.1 * inch

    c.setStrokeColor(HexColor("#0F3D91"))
    c.setLineWidth(1.5)
    c.setFillColor(HexColor("#FFFFFF"))
    c.roundRect(card_x, card_y, card_w, card_h, 8, fill=1, stroke=1)

    c.setFillColor(HexColor("#0F3D91"))
    c.setFont("Helvetica-Bold", 14)
    c.drawString(card_x + 0.2*inch, card_y + 0.38*inch, "RECIBO")
    c.setFont("Helvetica-Bold", 18)
    c.drawRightString(card_x + card_w - 0.2*inch, card_y + 0.38*inch, f"#{recibo_number}")

    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#6B7280"))
    c.drawString(card_x + 0.2*inch, card_y + 0.12*inch, "Fecha:")
    c.drawString(card_x + 0.5*inch, card_y + 0.12*inch, f"{date_str}")

    # Línea separadora azul horizontal
    y_pos = y_pos - 0.9 * inch
    c.setStrokeColor(HexColor("#0F3D91"))
    c.setLineWidth(2)
    c.line(0.5 * inch, y_pos, page_w - 0.5 * inch, y_pos)

    return y_pos - 0.2 * inch


def draw_patient_info(c, page_w, y_pos, patient, dx, ars_name, sala):
    """Draw patient info with icons.""" 
    box_x = 0.5 * inch
    box_w = page_w - 1.0 * inch
    box_h = 1.2 * inch
    box_y = y_pos - box_h

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.setLineWidth(1)
    c.setFillColor(HexColor("#F4F6F8"))
    c.roundRect(box_x, box_y, box_w, box_h, 10, fill=1, stroke=1)

    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(HexColor("#0F3D91"))
    c.drawString(box_x + 0.2*inch, box_y + box_h - 0.25*inch, "DATOS DEL PACIENTE")

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.line(box_x, box_y + box_h - 0.4*inch, box_x + box_w, box_y + box_h - 0.4*inch)

    icon_size = 12
    label_y = box_y + 0.55*inch
    draw_icon_text(c, box_x + 0.2*inch, label_y, "👤", icon_size)
    c.setFont("Helvetica-Bold", 11)
    c.setFillColor(HexColor("#6B7280"))
    c.drawString(box_x + 0.4*inch, label_y, "Paciente:")
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(HexColor("#1F2937"))
    c.drawString(box_x + 0.2*inch, box_y + 0.38*inch, patient.upper() if patient else "")

    label_y = box_y + 0.18*inch
    draw_icon_text(c, box_x + 0.2*inch, label_y, "📋", icon_size)
    c.setFont("Helvetica-Bold", 11)
    c.setFillColor(HexColor("#6B7280"))
    c.drawString(box_x + 0.4*inch, label_y, "Diagnóstico:")
    c.setFont("Helvetica", 12)
    c.setFillColor(HexColor("#1F2937"))
    c.drawString(box_x + 0.2*inch, box_y, dx.upper() if dx else "N/A")

    col2_x = box_x + 3.2 * inch
    label_y = box_y + 0.55*inch
    draw_icon_text(c, col2_x, label_y, "🆔", icon_size)
    c.setFont("Helvetica-Bold", 11)
    c.setFillColor(HexColor("#6B7280"))
    c.drawString(col2_x + 0.2*inch, label_y, "ARS:")
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(HexColor("#1F2937"))
    c.drawString(col2_x, box_y + 0.38*inch, ars_name.upper() if ars_name else "N/A")

    sala_box_w = 1.8 * inch
    sala_box_h = 0.65 * inch
    sala_box_x = box_x + box_w - sala_box_w - 0.2*inch
    sala_box_y = box_y + box_h - sala_box_h - 0.2*inch

    c.setFillColor(HexColor("#0F3D91"))
    c.roundRect(sala_box_x, sala_box_y, sala_box_w, sala_box_h, 8, fill=1, stroke=0)
    draw_icon_text(c, sala_box_x + 0.2*inch, sala_box_y + sala_box_h - 0.2*inch, "🛏️", 12, white)
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(sala_box_x + sala_box_w/2, sala_box_y + sala_box_h - 0.18*inch, "COSTO SALA")
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(sala_box_x + sala_box_w/2, sala_box_y + 0.1*inch, f"RD$ {sala:,.2f}")

    return box_y - 0.25 * inch


def draw_section_table(c, page_w, page_h, y_pos, category, items, bottom_margin, current_user):
    """Draw category card with icon, table, and subtotal inside the same box."""
    from reportlab.pdfbase.pdfmetrics import stringWidth

    CAT_PDF_COLORS = {
        "Medicamentos": HexColor("#1A56DB"),
        "Materiales": HexColor("#059669"),
        "Laboratorios": HexColor("#7C3AED"),
        "Imágenes": HexColor("#EA580C"),
        "Honorarios": HexColor("#1D4ED8")
    }
    cat_color = CAT_PDF_COLORS.get(category, HexColor("#0F3D91"))
    icon_text = {
        "Medicamentos": "💊",
        "Materiales": "📦",
        "Laboratorios": "🩸",
        "Imágenes": "📷",
        "Honorarios": "👨‍⚕️"
    }.get(category, "📄")

    table_data = [["Descripción", "Cant.", "Precio Unit.", "Total"]]
    subtotal = 0.0
    for name, pu, qty, total_item, _ in items:
        subtotal += total_item
        table_data.append([
            safe_pdf_text(name),
            str(int(qty)),
            f"{pu:,.2f}",
            f"{total_item:,.2f}"
        ])

    col_widths = [0.55 * (page_w - 1.0*inch), 0.10 * (page_w - 1.0*inch), 0.15 * (page_w - 1.0*inch), 0.20 * (page_w - 1.0*inch)]
    table = Table(table_data, colWidths=col_widths, hAlign='CENTER', repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), HexColor("#EAF3FF")),
        ('TEXTCOLOR', (0,0), (-1,0), HexColor("#0F3D91")),
        ('LINEABOVE', (0,0), (-1,0), 1, HexColor("#0F3D91")),
        ('LINEBELOW', (0,0), (-1,0), 1, HexColor("#D7DEE8")),
        ('ALIGN', (1,0), (-1,-1), 'CENTER'),
        ('ALIGN', (2,0), (-1,-1), 'RIGHT'),
        ('ALIGN', (3,0), (-1,-1), 'RIGHT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('INNERGRID', (0,1), (-1,-1), 0.25, HexColor("#E5E7EB")),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('FONTNAME', (0,0), (-1,0), "Helvetica-Bold"),
        ('FONTNAME', (0,1), (-1,-1), "Helvetica"),
        ('FONTSIZE', (0,0), (-1,-1), 10),
    ]))

    table_width = page_w - 1.0*inch
    w, table_h = table.wrap(table_width, y_pos - bottom_margin)
    badge_h = 0.4 * inch
    total_needed = table_h + badge_h + 0.4 * inch

    if y_pos - total_needed < bottom_margin:
        c.showPage()
        y_pos = page_h - 0.5 * inch

    box_x = 0.5 * inch
    box_w = page_w - 1.0 * inch
    box_h = total_needed
    box_y = y_pos - box_h

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.setLineWidth(1)
    c.setFillColor(white)
    c.roundRect(box_x, box_y, box_w, box_h, 10, fill=1, stroke=1)

    icon_x = box_x + 0.2*inch
    icon_y = box_y + box_h - 0.4*inch
    c.setFillColor(cat_color)
    c.roundRect(icon_x, icon_y, 0.3*inch, 0.3*inch, 6, fill=1, stroke=0)
    draw_icon_text(c, icon_x + 0.05*inch, icon_y + 0.05*inch, icon_text, 14, white)
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(HexColor("#0F3D91"))
    c.drawString(icon_x + 0.4*inch, icon_y + 0.05*inch, category.upper())

    table.drawOn(c, box_x + 0.2*inch, box_y + 0.15*inch + badge_h)

    badge_text_lbl = f"Subtotal {category}:"
    badge_text_val = f"RD$ {subtotal:,.2f}"

    lbl_w = stringWidth(badge_text_lbl, "Helvetica-Bold", 11)
    val_w = stringWidth(badge_text_val, "Helvetica-Bold", 11)
    badge_w = lbl_w + val_w + 34
    badge_h_local = 0.28 * inch
    badge_x = box_x + box_w - badge_w - 0.2*inch
    badge_y = box_y + 0.05*inch

    c.setFillColor(cat_color)
    c.roundRect(badge_x, badge_y, badge_w, badge_h_local, 8, fill=1, stroke=0)

    c.setFillColor(HexColor("#E8F0FE"))
    c.setFont("Helvetica-Bold", 11)
    c.drawString(badge_x + 10, badge_y + 0.08*inch, badge_text_lbl)
    c.setFillColor(white)
    c.drawRightString(badge_x + badge_w - 10, badge_y + 0.08*inch, badge_text_val)

    return box_y - 0.2 * inch


def draw_summary_and_total(c, page_w, page_h, y_pos, subtotales, total_general, bottom_margin, sala):
    """Draw side-by-side Summary box and massive Grand Total block.""" 
    if y_pos < bottom_margin + 1.8 * inch:
        c.showPage()
        y_pos = page_h - 0.5 * inch

    box_y = y_pos - 1.6 * inch
    box_h = 1.6 * inch

    left_w = 3.2 * inch
    left_x = 0.5 * inch

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.setLineWidth(1)
    c.setFillColor(white)
    c.roundRect(left_x, box_y, left_w, box_h, 10, fill=1, stroke=1)

    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(HexColor("#0F3D91"))
    c.drawString(left_x + 0.2*inch, box_y + box_h - 0.25*inch, "RESUMEN DE CARGOS")

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.line(left_x + 0.1*inch, box_y + box_h - 0.4*inch, left_x + left_w - 0.1*inch, box_y + box_h - 0.4*inch)

    y_text = box_y + box_h - 0.6*inch
    c.setFont("Helvetica", 11)
    c.setFillColor(HexColor("#1F2937"))

    c.drawString(left_x + 0.2*inch, y_text, "Costo Sala:")
    c.drawRightString(left_x + left_w - 0.2*inch, y_text, f"RD$ {sala:,.2f}")
    y_text -= 0.22 * inch

    for label, total in subtotales.items():
        if total > 0:
            c.drawString(left_x + 0.2*inch, y_text, f"{label}:")
            c.drawRightString(left_x + left_w - 0.2*inch, y_text, f"RD$ {total:,.2f}")
            y_text -= 0.22 * inch

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.line(left_x + 0.1*inch, y_text + 0.08*inch, left_x + left_w - 0.1*inch, y_text + 0.08*inch)
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(HexColor("#0F3D91"))
    c.drawString(left_x + 0.2*inch, y_text - 0.25*inch, "TOTAL GENERAL:")
    c.drawRightString(left_x + left_w - 0.2*inch, y_text - 0.25*inch, f"RD$ {total_general:,.2f}")

    right_w = page_w - left_w - 0.7 * inch
    right_x = left_x + left_w + 0.2 * inch

    c.setFillColor(HexColor("#0F3D91"))
    c.roundRect(right_x, box_y, right_w, box_h, 10, fill=1, stroke=0)

    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(right_x + right_w/2, box_y + box_h - 0.35*inch, "TOTAL GENERAL")

    c.setFont("Helvetica-Bold", 48)
    c.drawCentredString(right_x + right_w/2, box_y + box_h - 0.95*inch, f"RD$ {total_general:,.2f}")

    number_text = number_to_text(int(total_general)) if total_general < 1000000 else ""
    if number_text:
        c.setFont("Helvetica", 10)
        c.drawCentredString(right_x + right_w/2, box_y + 0.1*inch, number_text)

    return box_y - 0.3 * inch


def draw_footer(c, page_w, current_user):
    """Draw professional footer with signature and three columns.""" 
    from datetime import datetime
    footer_y = 0.9 * inch

    now = datetime.now()

    c.setStrokeColor(HexColor("#D7DEE8"))
    c.setLineWidth(1)

    c.line(0.5 * inch, footer_y + 0.35*inch, page_w - 0.5 * inch, footer_y + 0.35*inch)

    mid_x = page_w / 3
    c.setFont("Helvetica", 11)
    c.setFillColor(HexColor("#6B7280"))
    c.drawString(0.5 * inch, footer_y + 0.18*inch, "Auxiliar:")
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(HexColor("#1F2937"))
    c.drawString(0.5 * inch, footer_y + 0.02*inch, str(current_user.get('full_name', 'N/A')))

    c.line(mid_x - 0.1*inch, footer_y + 0.45*inch, mid_x - 0.1*inch, footer_y - 0.1*inch)

    c.setFont("Helvetica", 11)
    c.setFillColor(HexColor("#6B7280"))
    c.drawCentredString(page_w / 2, footer_y + 0.18*inch, "Fecha y Hora:")
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(HexColor("#1F2937"))
    c.drawCentredString(page_w / 2, footer_y + 0.02*inch, now.strftime("%d/%m/%Y   %H:%M"))

    c.line(2 * mid_x + 0.1*inch, footer_y + 0.45*inch, 2 * mid_x + 0.1*inch, footer_y - 0.1*inch)

    c.setFont("Helvetica", 11)
    c.setFillColor(HexColor("#6B7280"))
    c.drawCentredString(page_w * 2/3 + 0.45*inch, footer_y + 0.18*inch, "Firma:")
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(HexColor("#0F3D91"))
    c.line(page_w - 1.8 * inch, footer_y - 0.02*inch, page_w - 0.5 * inch, footer_y - 0.02*inch)
    c.drawCentredString(page_w * 2/3 + 0.45*inch, footer_y - 0.15*inch, "Auxiliar")

    band_y = footer_y - 0.4 * inch
    c.setFillColor(HexColor("#0F3D91"))
    c.rect(0.5 * inch, band_y, page_w - 1.0 * inch, 0.25 * inch, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(page_w / 2, band_y + 0.03 * inch, "Gracias por confiar en nuestros servicios - Hospital Provincial Dr. Ángel Contreras Mejía")

class PDFStorageSyncWorker(QThread):
    """Sincroniza copias PDF con PostgreSQL fuera del tiempo visible."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.jobs = queue.Queue()
        self._running = True

    @Slot(int, str, str)
    def enqueue(self, recibo_id, filename, pdf_path):
        self.jobs.put((int(recibo_id), str(filename), str(pdf_path)))

    def stop(self):
        self._running = False
        self.jobs.put(None)

    def _load_pending(self):
        try:
            with db_connect() as con:
                cur = con.execute(
                    """SELECT id, pdf_filename FROM recibos
                       WHERE pdf_synced=0 AND is_deleted=0
                         AND COALESCE(pdf_filename, '') <> ''
                       ORDER BY id
                       LIMIT 50"""
                )
                for row in cur.fetchall():
                    filename = str(row["pdf_filename"])
                    pdf_path = stable_storage_path(PDFS_DIR, filename)
                    if os.path.exists(pdf_path):
                        self.jobs.put((int(row["id"]), filename, pdf_path))
        except Exception as exc:
            write_runtime_log(f"No se pudieron cargar PDF pendientes: {exc}")

    def run(self):
        self._load_pending()
        while self._running:
            try:
                job = self.jobs.get(timeout=1.0)
            except queue.Empty:
                continue
            if job is None:
                break

            recibo_id, filename, pdf_path = job
            started = perf_counter()
            error = None
            for attempt in range(1, 4):
                if not self._running:
                    break
                try:
                    with open(pdf_path, "rb") as pdf_file:
                        pdf_bytes = pdf_file.read()
                    with db_connect() as con:
                        con.execute(
                            "INSERT INTO pdf_storage(filename, file_data) VALUES(%s, %s) "
                            "ON CONFLICT(filename) DO UPDATE SET file_data=EXCLUDED.file_data",
                            (filename, psycopg2.Binary(pdf_bytes)),
                        )
                        con.execute(
                            "UPDATE recibos SET pdf_synced=1, pdf_sync_error=NULL WHERE id=%s",
                            (recibo_id,),
                        )
                    error = None
                    write_pdf_performance(
                        f"Recibo {recibo_id} | sincronización={perf_counter() - started:.3f}s "
                        f"| bytes={len(pdf_bytes)} | intento={attempt}"
                    )
                    break
                except Exception as exc:
                    error = exc
                    if attempt < 3 and self._running:
                        time.sleep(float(attempt))

            if error is not None and self._running:
                write_runtime_log(
                    f"No se pudo sincronizar {filename} después de 3 intentos: {error}"
                )
                try:
                    with db_connect() as con:
                        con.execute(
                            "UPDATE recibos SET pdf_sync_error=%s WHERE id=%s",
                            (str(error)[:500], recibo_id),
                        )
                except Exception:
                    pass


class PDFWorkerSignals(QObject):
    finished_signal = Signal(bool, str, str, int)
    sync_requested = Signal(int, str, str)


class PDFDatabaseWorker(threading.Thread):
    def __init__(self):
        super().__init__(name="PDFEngineThread", daemon=True)
        self.signals = PDFWorkerSignals()
        self.jobs = queue.Queue()
        self._running = True
        self.renderer = None

    def warm_up(self):
        started = perf_counter()
        try:
            from pdf_engine import ReceiptPDFRenderer

            self.renderer = ReceiptPDFRenderer(persistent=True)
            self.renderer.start()
            write_pdf_performance(
                f"Motor PDF precargado en {perf_counter() - started:.3f}s"
            )
        except Exception as exc:
            self.renderer = None
            write_runtime_log(f"No se pudo precargar Chromium: {exc}")

    def submit(self, job):
        self.jobs.put(dict(job))

    def stop(self):
        self._running = False
        self.jobs.put(None)

    def run(self):
        self.warm_up()
        try:
            while self._running:
                job = self.jobs.get()
                if job is None:
                    break
                self.process(job)
        finally:
            if self.renderer is not None:
                self.renderer.close()
                self.renderer = None

    def process(self, job):
        self.patient = job["patient"]
        self.date_str = job["date_str"]
        self.dx_raw = job["dx_raw"]
        self.ars_name = job["ars_name"]
        self.coverage = job.get("coverage", "ASEGURADO")
        self.sala = job["sala"]
        self.grouped = job["grouped"]
        self.total_general = job["total_general"]
        self.editing_id = job["editing_id"]
        self.editing_num = job["editing_num"]
        self.current_user = job["current_user"]
        self.is_backdated = job["is_backdated"]
        total_started = perf_counter()
        number_elapsed = 0.0
        prepare_started = perf_counter()
        try:
            generated_at = now_str()
            if self.editing_id is not None:
                try:
                    with db_connect() as con:
                        cur = con.execute("SELECT created_at FROM recibos WHERE id=%s", (self.editing_id,))
                        row = cur.fetchone()
                        if row and row["created_at"]:
                            generated_at = row["created_at"]
                except Exception:
                    pass

            number_started = perf_counter()
            if self.editing_id is not None: recibo_number = self.editing_num
            else: recibo_number = get_next_recibo_number()
            number_elapsed = perf_counter() - number_started
            
            pdf_filename = f"recibo_{recibo_number}.pdf"
            pdf_path = os.path.join(PDFS_DIR, pdf_filename)
            
            # =========================================================
            # NUEVO MOTOR PDF HTML/CSS (Jinja2 + Playwright/Chromium)
            # =========================================================
            # Este bloque intenta generar el recibo moderno usando renderer.py.
            # Si algo falla, usa automáticamente el generador ReportLab anterior
            # para no romper el flujo de facturación.

            icon_map_pdf_html = {
                "Medicamentos": "💊",
                "Materiales": "🧤",
                "Laboratorios": "🧪",
                "Imágenes": "📷",
                "Procedimientos": "🩺",
                "Honorarios": "👨‍⚕️",
            }

            color_map_pdf_html = {
                "Medicamentos": "#1A56DB",
                "Materiales": "#059669",
                "Laboratorios": "#7C3AED",
                "Imágenes": "#EA580C",
                "Procedimientos": "#2563EB",
                "Honorarios": "#1D4ED8",
            }

            categorias_pdf_html = []
            for label, lst in self.grouped:
                subtotal_cat = sum(float(total_item or 0) for _name, _pu, _qty, total_item, _ars in lst)
                items_cat = []
                for name, pu, qty, total_item, _ars in lst:
                    items_cat.append({
                        "descripcion": safe_pdf_text(name),
                        "cantidad": int(qty),
                        "precio": float(pu or 0),
                        "total": float(total_item or 0),
                    })

                categorias_pdf_html.append({
                    "nombre": label,
                    "icon": icon_map_pdf_html.get(label, "📄"),
                    "color": color_map_pdf_html.get(label, "#0F3D91"),
                    "subtotal": subtotal_cat,
                    "items": items_cat,
                })

            total_letras_pdf_html = ""
            try:
                total_letras_pdf_html = number_to_text(int(float(self.total_general))).replace("SON:", "").strip()
            except Exception:
                total_letras_pdf_html = ""

            generado_pdf_html = datetime.now().strftime("%d/%m/%Y %H:%M")
            try:
                if generated_at:
                    generado_pdf_html = datetime.strptime(str(generated_at), "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y %H:%M")
            except Exception:
                pass

            try:
                recibo_data_pdf_html = {
                    "numero": recibo_number,
                    "fecha": self.date_str,
                    "paciente": self.patient,
                    "diagnostico": self.dx_raw or "N/A",
                    "ars": self.ars_name or ("NO ASEGURADO" if self.coverage == "NO_ASEGURADO" else "N/A"),
                    "sala": float(self.sala or 0),
                    "categorias": categorias_pdf_html,
                    "total_general": float(self.total_general or 0),
                    "total_letras": total_letras_pdf_html,
                    "usuario": self.current_user.get("full_name", "N/A"),
                    "generado": generado_pdf_html,
                    "logo_path": LOGO_PATH,
                }
                prepare_elapsed = perf_counter() - prepare_started

                # El motor se importa como paquete para que también funcione
                # dentro del ejecutable generado por PyInstaller.
                if self.renderer is None:
                    from pdf_engine import ReceiptPDFRenderer
                    self.renderer = ReceiptPDFRenderer(persistent=True)

                render_started = perf_counter()
                self.renderer.render_pdf(
                    recibo_data_pdf_html,
                    pdf_path,
                    save_html_preview=False,
                )
                render_elapsed = perf_counter() - render_started

                if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) <= 0:
                    raise RuntimeError("El motor HTML/CSS no creó el archivo PDF.")

            except Exception as render_error:
                render_elapsed = perf_counter() - render_started if 'render_started' in locals() else 0.0
                if self.renderer is not None:
                    try:
                        self.renderer.close()
                    except Exception:
                        pass
                    self.renderer = None
                # Fallback seguro: generador ReportLab anterior.
                print(f"Error usando motor HTML/CSS. Se usará ReportLab: {render_error}")
                write_runtime_log(
                    f"Motor PDF HTML/CSS no disponible para recibo {recibo_number}; "
                    f"se usó ReportLab. Detalle: {render_error}"
                )

                c = rl_canvas.Canvas(pdf_path, pagesize=letter)
                page_w, page_h = letter
                bottom_margin = 0.6 * inch
                y = page_h - 0.5 * inch
                current_user_local = self.current_user

                def ensure_space(h=0.25*inch):
                    nonlocal y
                    if y - h < bottom_margin:
                        c.showPage()
                        y = page_h - 0.5 * inch

                y = draw_header(c, page_w, page_h, recibo_number, self.date_str)
                y = draw_patient_info(c, page_w, y, self.patient, self.dx_raw, self.ars_name, self.sala)

                subtotales = {label: sum(sub for _, _, _, sub, _ in lst) for label, lst in self.grouped}

                for label, lst in self.grouped:
                    y = draw_section_table(c, page_w, page_h, y, label, lst, bottom_margin, current_user_local)
                    ensure_space(0.2*inch)

                y = draw_summary_and_total(c, page_w, page_h, y, subtotales, self.total_general, bottom_margin, self.sala)
                draw_footer(c, page_w, current_user_local)
                c.save()
                render_elapsed = perf_counter() - render_started if 'render_started' in locals() else 0.0

            database_started = perf_counter()
            recibo_id = save_receipt_with_items(
                self.editing_id,
                recibo_number,
                self.patient,
                self.date_str,
                self.dx_raw,
                self.ars_name,
                self.sala,
                self.total_general,
                pdf_filename,
                self.current_user['username'],
                self.is_backdated,
                generated_at,
                self.grouped,
                self.coverage,
            )
            database_elapsed = perf_counter() - database_started
            visible_elapsed = perf_counter() - total_started
            write_pdf_performance(
                f"Recibo {recibo_number} | número={number_elapsed:.3f}s "
                f"| preparación={prepare_elapsed:.3f}s | render={render_elapsed:.3f}s "
                f"| base_datos={database_elapsed:.3f}s | visible={visible_elapsed:.3f}s"
            )
            self.signals.finished_signal.emit(True, "Éxito", pdf_path, recibo_number)
            self.signals.sync_requested.emit(recibo_id, pdf_filename, pdf_path)

        except Exception as e:
            write_runtime_log(f"Error generando PDF: {e}")
            self.signals.finished_signal.emit(False, str(e), "", 0)


def number_to_text(num: int) -> str:
    """Convert number to Spanish text (Dominican Republic format)."""
    if num == 0:
        return "CERO PESOS"
    
    units = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
    teens = ["DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE", "DIECISÉIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE"]
    tens = ["", "", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA", "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
    
    def convert_three(n):
        result = []
        c = n % 100
        d = n // 100
        if d > 0:
            result.append(f"{units[d]} CIENTOS")
        if c >= 10 and c <= 19:
            result.append(teens[c - 10])
        else:
            t = c // 10
            u = c % 10
            if t > 0:
                if u > 0:
                    result.append(f"{tens[t]} Y {units[u]}")
                else:
                    result.append(tens[t])
            elif u > 0:
                result.append(units[u])
        return " ".join(result)
    
    thousands = ["", "MIL", "MIL", "MILLONES"]
    parts = []
    
    milions = num // 1000000
    remainder = num % 1000000
    
    if milions > 0:
        if milions == 1:
            parts.append("UN MILLÓN")
        else:
            parts.append(f"{convert_three(milions)} MILLONES")
    
    thousands_val = remainder // 1000
    if thousands_val > 0:
        if thousands_val == 1:
            parts.append("MIL")
        else:
            parts.append(f"{convert_three(thousands_val)} MIL")
    
    remainder = remainder % 1000
    
    if remainder > 0:
        parts.append(convert_three(remainder))
    
    entero_text = " ".join(parts) if parts else "CERO"
    return f"SON: {entero_text} PESOS"


# -*- coding: utf-8 -*-
def generar_html_factura(
    recibo_numero: int,
    paciente: str,
    fecha: str,
    dx: str,
    ars: str,
    sala: float,
    items_por_categoria: dict,
    total_general: float,
    auxiliar: str = "",
    logo_path: str = None
) -> str:
    logo_base64 = ""
    if logo_path and os.path.exists(logo_path):
        try:
            with open(logo_path, 'rb') as f:
                logo_base64 = base64.b64encode(f.read()).decode('utf-8')
        except Exception:
            pass

    logo_img = f'<img src="data:image/png;base64,{logo_base64}" width="60" height="60" style="vertical-align:middle; margin-right:12px;"/>' if logo_base64 else ''

    subtotales = {}
    for categoria in ["Medicamentos", "Materiales", "Laboratorios", "Imágenes", "Honorarios"]:
        items = items_por_categoria.get(categoria, [])
        subtotales[categoria] = sum(item[3] for item in items)

    ICON_SVG = {
        "user": "👤",
        "hospital": "🏥",
        "file-invoice": "📄",
        "clipboard": "📋",
        "id-card": "🆔",
        "bed": "🛏️",
        "capsules": "💊",
        "bandage": "🩹",
        "flask": "🧪",
        "image": "📷",
        "user-doctor": "👨‍⚕️",
        "calendar-days": "📅",
        "signature": "✍️"
    }

    CAT_COLORS = {
        "Medicamentos": "#1A56DB",
        "Materiales": "#059669",
        "Laboratorios": "#7C3AED",
        "Imágenes": "#EA580C",
        "Honorarios": "#1D4ED8"
    }

    html = f'''
    <html>
    <head>
    <meta charset="UTF-8">
    <style>
    @page {{
        margin: 20px;
        size: letter;
    }}
    body {{
        font-family: 'Segoe UI', Arial, sans-serif;
        font-size: 14px;
        margin: 0;
        padding: 0;
        color: #1F2937;
        line-height: 1.5;
        background: white;
    }}
    .container {{
        padding: 20px;
        background: white;
    }}

    .header {{
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 20px;
        padding-bottom: 15px;
        border-bottom: 4px solid #0F3D91;
        flex-wrap: wrap;
    }}
    .header-left {{
        display: flex;
        align-items: center;
        gap: 12px;
    }}
    .header-titles {{
        display: flex;
        flex-direction: column;
    }}
    .header-title {{
        font-size: 28px;
        font-weight: bold;
        color: #0F3D91;
        margin: 0;
        line-height: 1.1;
        text-transform: uppercase;
    }}
    .header-subtitle {{
        font-size: 20px;
        color: #0F3D91;
        margin: 0;
    }}
    .header-desc {{
        font-size: 16px;
        color: #0F3D91;
        margin-top: 2px;
    }}
    .receipt-card {{
        border: 1px solid #0F3D91;
        border-radius: 12px;
        padding: 12px 20px;
        background: white;
        text-align: center;
        min-width: 160px;
    }}
    .receipt-card .label {{
        font-size: 14px;
        font-weight: bold;
        color: #0F3D91;
    }}
    .receipt-card .value {{
        font-size: 18px;
        font-weight: bold;
        color: #0F3D91;
    }}

    .patient-section {{
        background: #F4F6F8;
        border: 1px solid #D7DEE8;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 2px 8px rgba(0,0,0,.05);
    }}
    .section-header {{
        display: flex;
        align-items: center;
        margin-bottom: 15px;
    }}
    .section-icon {{
        width: 42px;
        height: 42px;
        background: #EAF3FF;
        border-radius: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-right: 10px;
        font-size: 24px;
    }}
    .section-title {{
        font-size: 20px;
        font-weight: bold;
        color: #0F3D91;
        margin: 0;
    }}
    .section-divider {{
        height: 1px;
        background: #D7DEE8;
        margin: 10px 0 15px 0;
    }}
    .patient-grid {{
        display: flex;
        gap: 20px;
        align-items: flex-start;
        flex-wrap: wrap;
    }}
    .patient-col {{
        flex: 1;
        min-width: 150px;
    }}
    .patient-label {{
        font-weight: 600;
        color: #6B7280;
        margin-top: 8px;
        font-size: 14px;
        display: flex;
        align-items: center;
        gap: 6px;
    }}
    .patient-value {{
        color: #1F2937;
        font-size: 15px;
        font-weight: 500;
    }}
    .sala-box {{
        background: #0F3D91;
        color: white;
        border-radius: 8px;
        padding: 12px 16px;
        text-align: center;
        font-weight: 600;
        font-size: 14px;
        min-width: 140px;
    }}
    .sala-box .amount {{
        font-size: 22px;
        font-weight: bold;
        margin-top: 4px;
    }}

    .category-card {{
        border: 1px solid #D7DEE8;
        border-radius: 12px;
        margin: 18px 0;
        padding: 0;
        box-shadow: 0 2px 8px rgba(0,0,0,.05);
        page-break-inside: avoid;
    }}
    .category-header {{
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 12px 16px;
        border-bottom: 1px solid #D7DEE8;
        background: #F9FAFB;
        border-radius: 12px 12px 0 0;
    }}
    .category-icon-box {{
        width: 42px;
        height: 42px;
        border-radius: 8px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 24px;
        color: white;
    }}
    .category-title {{
        font-size: 20px;
        font-weight: bold;
        color: #0F3D91;
        margin: 0;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
    }}
    th {{
        background: #EAF3FF;
        color: #0F3D91;
        padding: 10px 12px;
        text-align: left;
        font-weight: bold;
        font-size: 14px;
        border-bottom: 1px solid #D7DEE8;
    }}
    td {{
        padding: 8px 12px;
        border-bottom: 1px solid #E5E7EB;
        font-size: 14px;
    }}
    .col-desc {{ width: 60%; }}
    .col-cant {{ width: 10%; text-align: center; }}
    .col-price {{ width: 15%; text-align: right; }}
    .col-total {{ width: 15%; text-align: right; }}
    .subtotal-row {{
        display: flex;
        justify-content: flex-end;
        align-items: center;
        padding: 10px 16px;
        background: #F9FAFB;
        border-top: 1px solid #D7DEE8;
        border-radius: 0 0 12px 12px;
    }}
    .subtotal-badge {{
        background: #1A56DB;
        color: white;
        font-weight: bold;
        padding: 6px 20px;
        border-radius: 20px;
        font-size: 14px;
        display: inline-flex;
        align-items: center;
        gap: 8px;
    }}
    .subtotal-badge .label {{
        opacity: 0.85;
    }}
    .subtotal-badge .number {{
        font-weight: 900;
        font-size: 16px;
    }}

    .summary-section {{
        display: flex;
        gap: 20px;
        margin-top: 20px;
        flex-wrap: wrap;
    }}
    .summary-left {{
        flex: 1;
        border: 1px solid #D7DEE8;
        border-radius: 12px;
        padding: 20px;
        background: white;
        min-width: 250px;
    }}
    .summary-row {{
        display: flex;
        justify-content: space-between;
        padding: 4px 0;
        font-size: 14px;
        border-bottom: 1px solid #F3F4F6;
    }}
    .summary-row:last-of-type {{
        border-bottom: none;
    }}
    .total-row {{
        display: flex;
        justify-content: space-between;
        margin-top: 15px;
        padding-top: 15px;
        border-top: 2px solid #0F3D91;
        font-weight: bold;
        color: #0F3D91;
        font-size: 20px;
    }}
    .total-card {{
        flex: 0 0 45%;
        background: #0F3D91;
        color: white;
        border-radius: 12px;
        padding: 30px 20px;
        text-align: center;
        min-width: 200px;
    }}
    .total-title {{
        font-size: 18px;
        font-weight: bold;
        margin-bottom: 10px;
        letter-spacing: 1px;
    }}
    .total-value {{
        font-size: 58px;
        font-weight: bold;
        letter-spacing: -1px;
        line-height: 1.1;
    }}
    .total-text {{
        font-size: 12px;
        margin-top: 10px;
        opacity: 0.9;
    }}

    .footer {{
        margin-top: 30px;
        padding-top: 20px;
        border-top: 1px solid #D7DEE8;
    }}
    .footer-columns {{
        display: flex;
        gap: 20px;
        justify-content: space-between;
        flex-wrap: wrap;
    }}
    .footer-col {{
        flex: 1;
        text-align: center;
        min-width: 120px;
    }}
    .footer-col:not(:last-child) {{
        border-right: 1px dashed #D7DEE8;
        padding-right: 20px;
    }}
    .footer-col-label {{
        color: #6B7280;
        font-size: 14px;
        margin-bottom: 4px;
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 6px;
    }}
    .footer-col-value {{
        color: #1F2937;
        font-weight: bold;
        font-size: 15px;
    }}
    .footer-band {{
        background: #0F3D91;
        color: white;
        text-align: center;
        padding: 15px;
        margin: 30px -20px 0 -20px;
        font-size: 14px;
        border-radius: 0 0 12px 12px;
    }}
    </style>
    </head>
    <body>
    <div class="container">

    <div class="header">
        <div class="header-left">
            {logo_img}
            <div class="header-titles">
                <div class="header-title">HOSPITAL PROVINCIAL</div>
                <div class="header-subtitle">DR. ÁNGEL CONTRERAS MEJÍA</div>
                <div class="header-desc">DETALLE DE FACTURACIÓN DE EMERGENCIA</div>
            </div>
        </div>
        <div class="receipt-card">
            <div class="label">{ICON_SVG['file-invoice']} RECIBO</div>
            <div class="value">#{recibo_numero}</div>
            <div style="font-size:12px; color:#6B7280; margin-top:2px;">{ICON_SVG['calendar-days']} {fecha}</div>
        </div>
    </div>

    <div class="patient-section">
        <div class="section-header">
            <div class="section-icon">{ICON_SVG['hospital']}</div>
            <div class="section-title">DATOS DEL PACIENTE</div>
        </div>
        <div class="section-divider"></div>
        <div class="patient-grid">
            <div class="patient-col">
                <div class="patient-label">{ICON_SVG['user']} Paciente:</div>
                <div class="patient-value">{safe_pdf_text(paciente) or ''}</div>
                <div class="patient-label">{ICON_SVG['clipboard']} Diagnóstico:</div>
                <div class="patient-value">{safe_pdf_text(dx) or 'N/A'}</div>
                <div class="patient-label">{ICON_SVG['id-card']} ARS:</div>
                <div class="patient-value">{safe_pdf_text(ars) or 'N/A'}</div>
            </div>
            <div class="patient-col" style="flex:0 0 auto;">
                <div class="sala-box">
                    <div style="display:flex;align-items:center;justify-content:center;gap:6px;margin-bottom:4px;">
                        {ICON_SVG['bed']} COSTO SALA
                    </div>
                    <div class="amount">RD$ {sala:,.2f}</div>
                </div>
            </div>
        </div>
    </div>
    '''

    icon_map = {
        "Medicamentos": "capsules",
        "Materiales": "bandage",
        "Laboratorios": "flask",
        "Imágenes": "image",
        "Honorarios": "user-doctor"
    }

    for categoria in ["Medicamentos", "Materiales", "Laboratorios", "Imágenes", "Honorarios"]:
        items = items_por_categoria.get(categoria, [])
        if not items:
            continue

        subtotal = subtotales.get(categoria, 0)
        cat_color = CAT_COLORS.get(categoria, "#0F3D91")
        icono = ICON_SVG.get(icon_map.get(categoria, ""), "")

        html += f'''
        <div class="category-card">
            <div class="category-header">
                <div class="category-icon-box" style="background:{cat_color};">{icono}</div>
                <div class="category-title">{categoria.upper()}</div>
            </div>
            <table>
                <tr>
                    <th class="col-desc">Descripción</th>
                    <th class="col-cant">Cant.</th>
                    <th class="col-price">Precio Unit.</th>
                    <th class="col-total">Total</th>
                </tr>
        '''
        for name, pu, qty, total_item, _ in items:
            html += f'''
                <tr>
                    <td>{safe_pdf_text(name)}</td>
                    <td style="text-align:center;">{int(qty)}</td>
                    <td style="text-align:right;">RD$ {pu:,.2f}</td>
                    <td style="text-align:right;">RD$ {total_item:,.2f}</td>
                </tr>
            '''
        html += f'''
            </table>
            <div class="subtotal-row">
                <div class="subtotal-badge" style="background:{cat_color};">
                    <span class="label">Subtotal {categoria}:</span>
                    <span class="number">RD$ {subtotal:,.2f}</span>
                </div>
            </div>
        </div>
        '''

    number_text = number_to_text(int(total_general)) if total_general < 1000000 else ""

    from datetime import datetime
    now = datetime.now()

    html += f'''
    <div class="summary-section">
        <div class="summary-left">
            <div class="section-header">
                <div class="section-icon" style="background:#EAF3FF;">{ICON_SVG['file-invoice']}</div>
                <div class="section-title">RESUMEN DE CARGOS</div>
            </div>
            <div class="section-divider"></div>
            <div class="summary-row">
                <span>{ICON_SVG['bed']} Costo Sala:</span>
                <span>RD$ {sala:,.2f}</span>
            </div>
    '''
    for categoria, subtotal in subtotales.items():
        if subtotal > 0:
            icono = ICON_SVG.get(icon_map.get(categoria, ""), "")
            html += f'''
            <div class="summary-row">
                <span>{icono} {categoria}:</span>
                <span>RD$ {subtotal:,.2f}</span>
            </div>
            '''
    html += f'''
            <div class="total-row">
                <span>TOTAL GENERAL</span>
                <span>RD$ {total_general:,.2f}</span>
            </div>
        </div>

        <div class="total-card">
            <div class="total-title">TOTAL GENERAL</div>
            <div class="total-value">RD$ {total_general:,.2f}</div>
            <div class="total-text">{number_text}</div>
        </div>
    </div>
    '''

    html += f'''
    <div class="footer">
        <div class="footer-columns">
            <div class="footer-col">
                <div class="footer-col-label">{ICON_SVG['user']} Auxiliar:</div>
                <div class="footer-col-value">{safe_pdf_text(auxiliar) or 'N/A'}</div>
            </div>
            <div class="footer-col">
                <div class="footer-col-label">{ICON_SVG['calendar-days']} Fecha y Hora:</div>
                <div class="footer-col-value">{now.strftime('%d/%m/%Y   %H:%M')}</div>
            </div>
            <div class="footer-col">
                <div class="footer-col-label">{ICON_SVG['signature']} Firma Auxiliar</div>
                <div style="width:180px;border-top:1px solid #0F3D91;margin:5px auto 0 auto;"></div>
            </div>
        </div>
        <div class="footer-band">
            Gracias por confiar en nuestros servicios - Hospital Provincial Dr. Ángel Contreras Mejía
        </div>
    </div>
    </div>
    </body>
    </html>
    '''
    return html
def imprimir_factura_desde_html(html: str, parent=None, is_preview: bool = False):
    """Muestra vista previa o imprime/exporta a PDF usando QTextDocument.
    
    Args:
        html: HTML content del documento
        parent: Widget padre
        is_preview: Si es True, retorna QTextBrowser con el documento (no muestra diálogo)
    """
    doc = QTextDocument()
    doc.setHtml(html)
    
    if is_preview:
        from PySide6.QtWidgets import QTextBrowser
        browser = QTextBrowser()
        browser.setDocument(doc)
        browser.setReadOnly(True)
        return browser
    
    printer = QPrinter()
    printer.setOutputFormat(QPrinter.PdfFormat)
    
    if parent:
        dialog = QPrintDialog(printer, parent)
        dialog.setWindowTitle("Guardar Factura como PDF")
        if dialog.exec() == QDialog.Accepted:
            doc.print_(printer)
    else:
        printer.setOutputFileName(os.path.join(PDFS_DIR, f"factura_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"))
        doc.print_(printer)


def draw_pdf_banner(c, page_w, page_h, title_text):
    top_y = page_h - 0.32 * inch
    if LOGO_PATH and os.path.exists(LOGO_PATH):
        try:
            c.drawImage(
                LOGO_PATH,
                0.05 * inch,
                top_y - 0.84 * inch,
                width=page_w - 0.10 * inch,
                height=1.28 * inch,
                preserveAspectRatio=True,
                anchor='n',
                mask='auto'
            )
        except Exception:
            pass
            
    font_size = 14
    c.setFont("Helvetica-Bold", font_size)
    margin_right = 0.5 * inch
    max_width = page_w - (0.20 * inch) - margin_right
    
    from reportlab.pdfbase.pdfmetrics import stringWidth  # type: ignore
    while stringWidth(title_text, "Helvetica-Bold", font_size) > max_width and font_size > 8:
        font_size -= 0.5
        c.setFont("Helvetica-Bold", font_size)
        
    c.drawString(0.20 * inch, top_y - 1.10 * inch, title_text)
    
    return top_y - 1.34 * inch

def _legacy_create_report_pdf(filename: str, title: str, date_subtitle: str, totals: dict, generated_by: str) -> str:
    os.makedirs(REPORTS_DIR, exist_ok=True)
    path = os.path.join(REPORTS_DIR, filename)
    c = rl_canvas.Canvas(path, pagesize=letter)
    page_w, page_h = letter

    y = draw_pdf_banner(c, page_w, page_h, title)

    def ensure_space(h):
        nonlocal y
        if y - h < 0.85 * inch:
            c.setFont("Helvetica-Bold", 9)
            c.drawRightString(page_w - 0.50 * inch, 0.50 * inch, f"Generado por: {generated_by}")
            c.showPage()
            y = draw_pdf_banner(c, page_w, page_h, title)
            c.setFont("Helvetica", 10)

    c.setFont("Helvetica-Bold", 12)
    c.drawString(1 * inch, y, date_subtitle)
    y -= 0.22 * inch

    c.setFont("Helvetica", 9)
    c.drawString(1 * inch, y, "Criterio de reporte: fecha real de generación del recibo.")
    y -= 0.28 * inch

    c.setFont("Helvetica", 11)
    for cat, val in totals.items():
        if cat not in ["Total General", "Sala Emergencia"] and not str(cat).startswith("_"):
            ensure_space(0.25 * inch)
            c.drawString(1 * inch, y, f"{cat}:")
            c.drawRightString(page_w - 1 * inch, y, f"${val:,.2f}")
            y -= 0.25 * inch

    y -= 0.1 * inch
    ensure_space(0.6 * inch)
    c.setFont("Helvetica-Bold", 11)
    c.drawString(1 * inch, y, "Sala de Emergencia:")
    c.drawRightString(page_w - 1 * inch, y, f"${totals.get('Sala Emergencia', 0.0):,.2f}")
    y -= 0.3 * inch

    c.setFont("Helvetica-Bold", 14)
    c.drawString(1 * inch, y, "TOTAL GENERAL RECAUDADO:")
    c.drawRightString(page_w - 1 * inch, y, f"${totals.get('Total General', 0.0):,.2f}")

    total_recibos = totals.get("_total_recibos", 0)
    if total_recibos > 0:
        y -= 0.5 * inch
        ensure_space(0.5 * inch)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(1 * inch, y, f"TOTAL DE RECIBOS EMITIDOS: {total_recibos}")
        y -= 0.3 * inch

        ars_counts = totals.get("_ars_counts", {})
        user_counts = totals.get("_user_counts", {})

        if ars_counts:
            c.setFont("Helvetica-Bold", 10)
            c.drawString(1 * inch, y, "Desglose por ARS:")
            y -= 0.2 * inch
            c.setFont("Helvetica", 10)
            for a_name, a_count in ars_counts.items():
                ensure_space(0.2 * inch)
                c.drawString(1.2 * inch, y, f"• {a_name}: {a_count} recibo(s)")
                y -= 0.15 * inch
            y -= 0.1 * inch

        if user_counts:
            ensure_space(0.3 * inch)
            c.setFont("Helvetica-Bold", 10)
            c.drawString(1 * inch, y, "Desglose por Auxiliar / Usuario:")
            y -= 0.2 * inch
            c.setFont("Helvetica", 10)
            for u_name, u_count in user_counts.items():
                ensure_space(0.2 * inch)
                c.drawString(1.2 * inch, y, f"• {u_name}: {u_count} recibo(s)")
                y -= 0.15 * inch

    y -= 0.5 * inch
    ensure_space(0.5 * inch)
    c.setFont("Helvetica-Bold", 10)
    c.drawRightString(page_w - 0.5 * inch, y, f"Generado por: {generated_by}")
    y -= 0.3 * inch

    c.save()

    with open(path, 'rb') as f:
        pdf_bytes = f.read()
    with db_connect() as con:
        con.execute(
            "INSERT INTO pdf_storage(filename, file_data) VALUES(%s, %s) "
            "ON CONFLICT(filename) DO UPDATE SET file_data=EXCLUDED.file_data",
            (filename, psycopg2.Binary(pdf_bytes))
        )
    return path


def _create_report_pdf(filename: str, title: str, date_subtitle: str, totals: dict, generated_by: str) -> str:
    """Genera el reporte normal con el mismo motor HTML/CSS de los recibos."""
    os.makedirs(REPORTS_DIR, exist_ok=True)
    path = os.path.join(REPORTS_DIR, filename)
    category_rows = []
    for category in ALL_CATEGORIES:
        value = float(totals.get(category, 0) or 0)
        if value > 0:
            category_rows.append((category, value))
    room = float(totals.get("Sala Emergencia", 0) or 0)
    if room > 0:
        category_rows.append(("Sala Emergencia", room))

    context = {
        "mode": "standard",
        "title": title,
        "subtitle": date_subtitle,
        "generated_by": generated_by,
        "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "logo_path": LOGO_PATH or "",
        "totals": totals,
        "category_rows": category_rows,
        "ars_rows": list((totals.get("_ars_counts") or {}).items()),
        "user_rows": list((totals.get("_user_counts") or {}).items()),
    }
    try:
        ReportHTMLRenderer().render_pdf(context, path, landscape=False)
    except Exception as exc:
        write_runtime_log(f"Reporte HTML/CSS falló ({filename}): {exc}")
        raise RuntimeError(
            "No se pudo generar el reporte institucional con el motor HTML/CSS. "
            f"Detalle: {exc}"
        ) from exc

    with open(path, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()
    with db_connect() as con:
        con.execute(
            "INSERT INTO pdf_storage(filename, file_data) VALUES(%s, %s) "
            "ON CONFLICT(filename) DO UPDATE SET file_data=EXCLUDED.file_data",
            (filename, psycopg2.Binary(pdf_bytes)),
        )
    return path


# ---> FUNCIONES PDF PARA REPORTES CON FILTRO ARS/USUARIO <---
def _report_filters_title(ars_filter, user_filter):
    ars_selection = _normalize_report_filter(ars_filter, "Todas las ARS")
    user_selection = _normalize_report_filter(user_filter, "Todos los Usuarios")
    parts = []
    if ars_selection["values"]:
        action = "excluidas" if ars_selection["mode"] == "exclude" else "incluidas"
        parts.append(f"ARS {action}: {', '.join(ars_selection['values'])}")
    if user_selection["values"]:
        action = "excluidos" if user_selection["mode"] == "exclude" else "incluidos"
        parts.append(f"Facturadores {action}: {', '.join(user_selection['values'])}")
    return " · ".join(parts), ars_selection, user_selection


def generate_daily_report_pdf(report_date: str, generated_by: str, is_backdated: int = 0, ars_filter="Todas las ARS", user_filter="Todos los Usuarios") -> str:
    include_all_histories = is_backdated is None or str(is_backdated) == "-1"
    totals = get_receipt_stats_by_date(report_date, is_backdated, ars_filter, user_filter)
    tipo = "PRINCIPAL + ALTERNO" if include_all_histories else ("ALTERNO" if int(is_backdated) == 1 else "PRINCIPAL")
    tipo_file = "ambos_historiales" if include_all_histories else tipo.lower()

    _filter_title, ars_selection, user_selection = _report_filters_title(ars_filter, user_filter)
    title_extra = ""

    filename = f"reporte_diario_{tipo_file}_{report_date}.pdf"
    path = _create_report_pdf(
        filename,
        f"REPORTE DIARIO ({tipo}){title_extra}",
        f"Fecha de generación de recibos: {report_date}",
        totals,
        generated_by
    )

    if (not include_all_histories and int(is_backdated) == 0
            and _report_filter_is_all(ars_selection) and _report_filter_is_all(user_selection)):
        save_daily_report_record(report_date, path, totals, generated_by)
    else:
        full_type = f"Diario ({tipo.title()}){title_extra}"
        save_report_history(full_type, report_date, report_date, path, totals, generated_by)

    return path

def generate_period_report_pdf(
    report_type: str,
    start_date: str,
    end_date: str,
    generated_by: str,
    is_backdated: int = 0,
    ars_filter="Todas las ARS",
    user_filter="Todos los Usuarios",
    period_metadata: dict | None = None,
) -> str:
    include_all_histories = is_backdated is None or str(is_backdated) == "-1"
    totals = get_receipt_stats_between(start_date, end_date, is_backdated, ars_filter, user_filter)
    tipo = "PRINCIPAL + ALTERNO" if include_all_histories else ("ALTERNO" if int(is_backdated) == 1 else "PRINCIPAL")
    tipo_file = "ambos_historiales" if include_all_histories else tipo.lower()

    _filter_title, _ars_selection, _user_selection = _report_filters_title(
        ars_filter, user_filter
    )
    title_extra = ""

    if period_metadata:
        totals["_period"] = copy.deepcopy(period_metadata)
    safe_report_type = re.sub(r"[^a-zA-Z0-9_-]+", "_", report_type).strip("_") or "periodo"
    filename = f"reporte_{safe_report_type.lower()}_{tipo_file}_{start_date}_al_{end_date}.pdf"
    path = _create_report_pdf(
        filename,
        f"REPORTE {report_type.upper()} ({tipo}){title_extra}",
        f"Período por fecha de generación: {start_date} al {end_date}",
        totals,
        generated_by
    )

    full_type = f"{report_type} ({tipo.title()}){title_extra}"
    save_report_history(full_type, start_date, end_date, path, totals, generated_by)
    return path


def generate_comparison_report_pdf(data: dict, generated_by: str) -> str:
    """Genera el comparativo desde la copia exacta ya cargada en el panel."""
    snapshot = copy.deepcopy(data or {})
    if not snapshot.get("filters", {}).get("compare_previous"):
        raise ValueError("La comparación no está activada.")
    if not snapshot.get("previous", {}).get("receipts"):
        raise ValueError("No existen datos válidos en el período anterior.")
    period = snapshot.get("period", {})
    current_label = period.get("period_label", "Período actual")
    previous_label = period.get("comparison_label", "Período anterior")
    start_date = snapshot.get("start_date", period.get("start_date", ""))
    end_date = snapshot.get("end_date", period.get("end_date", ""))
    filename = (
        f"Comparacion_{start_date}_al_{end_date}_vs_"
        f"{snapshot.get('previous_start', '')}_al_{snapshot.get('previous_end', '')}.pdf"
    )
    os.makedirs(REPORTS_DIR, exist_ok=True)
    path = os.path.join(REPORTS_DIR, filename)
    context = {
        "mode": "comparison",
        "title": "REPORTE COMPARATIVO",
        "subtitle": f"{current_label} frente a {previous_label}",
        "generated_by": generated_by,
        "generated_at": datetime.now().strftime("%d/%m/%Y %H:%M"),
        "logo_path": LOGO_PATH or "",
        "data": snapshot,
    }
    ReportHTMLRenderer().render_pdf(context, path, landscape=True)
    with open(path, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()
    with db_connect() as con:
        con.execute(
            "INSERT INTO pdf_storage(filename, file_data) VALUES(%s, %s) "
            "ON CONFLICT(filename) DO UPDATE SET file_data=EXCLUDED.file_data",
            (filename, psycopg2.Binary(pdf_bytes)),
        )
    history_payload = {
        "period": period,
        "filters": snapshot.get("filters", {}),
        "summary": snapshot.get("summary", {}),
        "previous": snapshot.get("previous", {}),
        "category_comparison": snapshot.get("category_comparison", []),
    }
    save_report_history(
        f"Reporte comparativo: {current_label} vs {previous_label}",
        start_date,
        end_date,
        path,
        history_payload,
        generated_by,
    )
    return path

def safe_generate_pending_reports(username: str):
    yesterday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    generated = False
    errors = False
    try:
        if not report_exists(yesterday):
            generate_daily_report_pdf(yesterday, username)
            generated = True
    except Exception as e:
        errors = True
        print(f"Error generando reporte pendiente: {e}")
    return generated, errors

def import_word_to_universal_catalog(path: str):
    parsed = parse_word_for_universal_categories(path)
    summary = {cat: {"updated": 0, "inserted": 0} for cat in UNIVERSAL_CATEGORIES}
    existing_maps = {cat: {} for cat in UNIVERSAL_CATEGORIES}
    for cat in UNIVERSAL_CATEGORIES:
        for existing_name in get_universal(cat).keys():
            existing_maps[cat].setdefault(catalog_identity_key(existing_name), []).append(existing_name)
    with db_connect() as con:
        for cat in UNIVERSAL_CATEGORIES:
            for name, price in parsed[cat]:
                display_name = normalize_catalog_name(name)
                key = catalog_identity_key(display_name)
                previous_names = list(existing_maps[cat].get(key, []))
                for previous_name in previous_names:
                    if previous_name != display_name:
                        con.execute(
                            "UPDATE universal_items SET is_active=0 WHERE categoria=%s AND nombre=%s",
                            (cat, previous_name),
                        )
                con.execute(
                    "INSERT INTO universal_items(categoria, nombre, precio, is_active) VALUES(%s,%s,%s, 1) "
                    "ON CONFLICT(categoria, nombre) DO UPDATE SET precio=EXCLUDED.precio, is_active=1",
                    (cat, display_name.strip(), float(price))
                )
                if previous_names: summary[cat]["updated"] += 1
                else:
                    summary[cat]["inserted"] += 1
                existing_maps[cat][key] = [display_name]
    return summary

def import_word_to_ars_catalog(path: str, ars_name: str):
    if not ars_name: raise ValueError("Debes seleccionar una ARS.")
    parsed = parse_word_for_ars_categories(path)
    summary = {cat: {"updated": 0, "inserted": 0} for cat in ARS_CATEGORIES}
    existing_maps = {cat: {name.strip().casefold(): name for name in get_ars_items(cat, ars_name).keys()} for cat in ARS_CATEGORIES}
    total_loaded = 0
    _ensure_ars(ars_name)
    with db_connect() as con:
        for cat in ARS_CATEGORIES:
            for name, price in parsed[cat]:
                key = name.casefold()
                display_name = existing_maps[cat].get(key, name)
                con.execute(
                    """
                    INSERT INTO ars_items(ars_id, categoria, nombre, precio, is_active)
                    SELECT a.id, %s, %s, %s, 1 FROM ars a WHERE a.nombre=%s
                    ON CONFLICT(ars_id, categoria, nombre) DO UPDATE SET precio=EXCLUDED.precio, is_active=1
                    """,
                    (cat, display_name.strip(), float(price), ars_name)
                )
                if key in existing_maps[cat]: summary[cat]["updated"] += 1
                else:
                    summary[cat]["inserted"] += 1
                    existing_maps[cat][key] = display_name
                total_loaded += 1
    if total_loaded == 0: raise ValueError("No se detectaron encabezados válidos en el archivo Word.")
    return summary

def parse_word_for_universal_categories(path: str):
    if Document is None: raise RuntimeError("Falta dependencia python-docx")
    doc = Document(path); parsed = {"Medicamentos": [], "Materiales": []}; current_cat = None
    previous_names = {"Medicamentos": "", "Materiales": ""}
    def push(cat: str, raw_name: str, raw_price_text: str):
        if cat not in UNIVERSAL_CATEGORIES: return
        if _looks_like_month_or_month_year(raw_name): return
        combined = str(raw_name or "")
        if "$" not in combined and raw_price_text:
            combined = f"{combined} ${raw_price_text}"
        entries = parse_universal_catalog_line(combined, previous_names.get(cat, ""))
        for name, price in entries:
            parsed[cat].append((name, apply_price_rule(cat, price)))
            previous_names[cat] = name
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells): continue
            if len(cells) == 1:
                one = cells[0]; cat = _norm_cat(one)
                if cat: current_cat = cat; continue
                if current_cat: push(current_cat, one, one)
                continue
            joined_lower = "|".join(cells).casefold()
            if ("nombre" in joined_lower and "precio" in joined_lower) or ("item" in joined_lower and "precio" in joined_lower): continue
            cat = _norm_cat(cells[0]) or current_cat or _norm_cat(" ".join(cells)) or "Medicamentos"
            name = cells[1] if _norm_cat(cells[0]) and len(cells) >= 2 else cells[0]
            push(cat, name, cells[-1])
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt or _looks_like_month_or_month_year(txt): continue
        maybe_cat = _norm_cat(txt)
        if maybe_cat: current_cat = maybe_cat; continue
        m = re.match(r'^\s*(?P<cat>Medicamentos?|Materiales?)\s*[:\\-|]\s*(?P<name>.+?)\s*[:\\-|]\s*(?P<price>.+?)\s*$', txt, re.I)
        if m: push(_norm_cat(m.group('cat')), m.group('name').strip(), m.group('price')); continue
        if '$' in txt and current_cat: push(current_cat, txt, txt)
    dedup = {"Medicamentos": {}, "Materiales": {}}
    for cat in parsed:
        # Si el Word repite un ítem, conserva únicamente la última aparición.
        for name, price in parsed[cat]: dedup[cat][catalog_identity_key(name)] = (name, price)
    return {cat: list(dedup[cat].values()) for cat in dedup}

def parse_word_for_ars_categories(path: str):
    if Document is None: raise RuntimeError("Falta dependencia python-docx")
    doc = Document(path); parsed = {cat: [] for cat in ARS_CATEGORIES}; current_cat = None
    def push(cat: str, raw_name: str, raw_price_text: str):
        if cat not in ARS_CATEGORIES: return
        if _looks_like_month_or_month_year(raw_name): return
        name = _clean_name(raw_name)
        if not name: return
        price = _parse_price(raw_price_text)
        if price is None: return
        parsed[cat].append((name, apply_price_rule(cat, price)))
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells): continue
            if len(cells) == 1:
                one = cells[0]; cat = _norm_ars_cat(one)
                if cat: current_cat = cat; continue
                if current_cat: push(current_cat, one, one)
                continue
            joined_lower = "|".join(cells).casefold()
            if ("nombre" in joined_lower and "precio" in joined_lower) or ("item" in joined_lower and "precio" in joined_lower): continue
            cat = _norm_ars_cat(cells[0]) or current_cat or _norm_ars_cat(" ".join(cells))
            if not cat: continue
            name = cells[1] if _norm_ars_cat(cells[0]) and len(cells) >= 2 else cells[0]
            push(cat, name, cells[-1])
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if not txt or _looks_like_month_or_month_year(txt): continue
        maybe_cat = _norm_ars_cat(txt)
        if maybe_cat: current_cat = maybe_cat; continue
        m = re.match(r'^\s*(?P<cat>Laboratorios?|Imagenes?|Imágenes?|Procedimientos?|Honorarios?)\s*[:\\-|]\s*(?P<name>.+?)\s*[:\\-|]\s*(?P<price>.+?)\s*$', txt, re.I)
        if m: push(_norm_ars_cat(m.group('cat')), m.group('name').strip(), m.group('price')); continue
        if ('$' in txt or re.search(r'\d', txt)) and current_cat: push(current_cat, txt, txt)
    dedup = {cat: {} for cat in ARS_CATEGORIES}
    for cat in parsed:
        for name, price in parsed[cat]: dedup[cat][name.casefold()] = (name, price)
    return {cat: list(dedup[cat].values()) for cat in dedup}

class QInputDialogWithText(QDialog):
    @staticmethod
    def get(parent, title, label):
        dlg = QInputDialogWithText(parent, title, label)
        ok = dlg.exec()
        return dlg.line.text(), ok == QDialog.Accepted
    def __init__(self, parent, title, label):
        super().__init__(parent)
        self.setWindowTitle(title)
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel(label))
        self.line = QLineEdit()
        layout.addWidget(self.line)
        dbtns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        dbtns.accepted.connect(self.accept)
        dbtns.rejected.connect(self.reject)
        layout.addWidget(dbtns)
        for btn in dbtns.buttons(): set_button_role(btn, 'success' if btn.text().lower().startswith('ok') else 'neutral')

class QtyDialog(QDialog):
    def __init__(self, current_qty: int, parent=None):
        super().__init__(parent)
        self.parent_ref = parent
        self.setWindowTitle("Modificar cantidad")
        self.setFixedSize(360, 170)
        root = QVBoxLayout(self)
        root.setContentsMargins(24, 18, 24, 18)
        root.setSpacing(14)

        lbl = QLabel("Nueva cantidad:")
        lbl.setStyleSheet("font-size: 10pt; font-weight: 500;")
        root.addWidget(lbl, alignment=Qt.AlignLeft)

        spin_container = QHBoxLayout()
        spin_container.setSpacing(6)

        self.spin = QSpinBox()
        self.spin.setRange(1, 9999)
        self.spin.setValue(current_qty)
        self.spin.setButtonSymbols(QAbstractSpinBox.NoButtons)
        is_dark = parent.is_dark_mode if parent else False
        self.spin.setStyleSheet(f"""
            QSpinBox {{
                border: 1px solid {'#555' if is_dark else '#aaa'};
                border-radius: 6px;
                padding: 5px 12px;
                font-size: 11pt;
                min-height: 28px;
                background: {'#2d2d2d' if is_dark else '#f8f9fa'};
                color: {'white' if is_dark else '#212529'};
            }}
        """)

        self.spin.setFixedWidth(110)

        btn_down = QToolButton()
        btn_down.setText("−")
        btn_down.setToolTip("Disminuir cantidad")
        btn_down.setFixedSize(40, 40)
        btn_down.setStyleSheet(f"""
            QToolButton {{
                border: 1px solid {'#555' if is_dark else '#c4d0dc'};
                border-radius: 6px;
                background: {'#37474f' if is_dark else '#e7edf3'};
                color: {'white' if is_dark else '#263238'};
                font-weight: bold;
                font-size: 16pt;
            }}
            QToolButton:hover {{ background: {'#455a64' if is_dark else '#d5e0ea'}; }}
        """)
        btn_down.clicked.connect(lambda: self.spin.setValue(max(1, self.spin.value() - 1)))

        btn_up = QToolButton()
        btn_up.setText("+")
        btn_up.setToolTip("Aumentar cantidad")
        btn_up.setFixedSize(40, 40)
        btn_up.setStyleSheet(btn_down.styleSheet())
        btn_up.clicked.connect(lambda: self.spin.setValue(self.spin.value() + 1))

        spin_container.addStretch()
        spin_container.addWidget(btn_down)
        spin_container.addWidget(self.spin)
        spin_container.addWidget(btn_up)
        spin_container.addStretch()
        root.addLayout(spin_container)

        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.setStyleSheet("font-size: 10pt;")
        btn_ok = btns.button(QDialogButtonBox.Ok)
        btn_cancel = btns.button(QDialogButtonBox.Cancel)
        btn_ok.setText("Aceptar")
        btn_cancel.setText("Cancelar")
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        root.addWidget(btns)
        set_button_role(btn_ok, 'success')
        set_button_role(btn_cancel, 'neutral')

class RegisterUserDialog(QDialog):
    def __init__(self, parent=None, allow_role_choice=False):
        super().__init__(parent)
        self.setWindowTitle("Crear Nueva Cuenta")
        self.setFixedSize(450, 520)

        main_lay = QVBoxLayout(self)
        main_lay.setContentsMargins(30, 30, 30, 30)
        main_lay.setSpacing(10)

        title = QLabel("Registro de Usuario")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 22pt; font-weight: 900; color: #1565c0; margin-bottom: 10px; border: none;")
        main_lay.addWidget(title)

        form = QFormLayout()
        form.setSpacing(12)
        
        self.full_name = QLineEdit()
        self.full_name.setPlaceholderText("Nombre y Apellido")
        
        self.username = QLineEdit()
        self.username.setPlaceholderText("👤 Nombre de usuario")
        
        self.password = QLineEdit(); self.password.setEchoMode(QLineEdit.Password)
        self.password.setPlaceholderText("🔒 Contraseña")
        
        self.password2 = QLineEdit(); self.password2.setEchoMode(QLineEdit.Password)
        self.password2.setPlaceholderText("🔒 Confirmar contraseña")
        
        self.question = QLineEdit()
        self.question.setPlaceholderText("Ej.: ¿Ciudad de nacimiento?")
        
        self.answer = QLineEdit(); self.answer.setEchoMode(QLineEdit.Password)
        self.answer.setPlaceholderText("Respuesta secreta")
        
        self.role = QComboBox()
        if allow_role_choice:
            self.role.addItems([ROLE_AUX, ROLE_ADMIN, ROLE_AUDIT])
        else:
            self.role.addItems([ROLE_AUX])
        
        form.addRow("Nombre Completo:", self.full_name)
        form.addRow("Usuario:", self.username)
        form.addRow("Contraseña:", self.password)
        form.addRow("Confirmar:", self.password2)
        form.addRow("Preg. Seguridad:", self.question)
        form.addRow("Respuesta:", self.answer)
        form.addRow("Rol en sistema:", self.role)
        
        main_lay.addLayout(form)
        
        info = QLabel("<i>La pregunta de seguridad te permitirá recuperar tu clave si la olvidas.</i>")
        info.setWordWrap(True)
        info.setStyleSheet("color: #78909c; font-size: 9pt; font-weight: normal;")
        main_lay.addWidget(info)
        
        main_lay.addStretch()
        
        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.accepted.connect(self._accept)
        btns.rejected.connect(self.reject)
        main_lay.addWidget(btns)
        
        for btn in btns.buttons():
            if btn.text().lower().startswith('ok'):
                btn.setText("Crear Cuenta")
                btn.setStyleSheet("background-color: #2e7d32; color: white; font-weight: bold;")
            else:
                btn.setText("Cancelar")
                btn.setStyleSheet("background-color: #b0bec5; color: #212529;")

    def _accept(self):
        if self.password.text() != self.password2.text():
            FloatingToast("Las contraseñas no coinciden", self, is_error=True).show()
            return
        try:
            create_user(self.full_name.text().strip(), self.username.text().strip(), self.password.text(), self.question.text().strip(), self.answer.text(), self.role.currentText())
            self.accept()
        except psycopg2.IntegrityError:
            FloatingToast("Ese usuario ya existe", self, is_error=True).show()
        except Exception as e:
            FloatingToast(str(e), self, is_error=True).show()

class RecoverPasswordDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Recuperar Contraseña")
        self.setFixedSize(450, 500)
        
        main_lay = QVBoxLayout(self)
        main_lay.setContentsMargins(30, 30, 30, 30)
        main_lay.setSpacing(12)
        
        title = QLabel("Recuperación de Acceso")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 22pt; font-weight: 900; color: #c62828; margin-bottom: 10px; border: none;")
        main_lay.addWidget(title)

        form = QFormLayout()
        form.setSpacing(12)
        
        self.username = QLineEdit()
        self.username.setPlaceholderText("👤 Ingresa tu usuario")
        
        self.btn_load = QPushButton("Buscar Pregunta Secreta")
        self.btn_load.setStyleSheet("background-color: #1565c0; color: white; font-weight: bold;")
        self.btn_load.clicked.connect(self.load_question)
        
        self.question_lbl = QLabel("<i>(Escribe tu usuario y busca la pregunta)</i>")
        self.question_lbl.setWordWrap(True)
        self.question_lbl.setStyleSheet("color: #1565c0; font-size: 10pt;")
        
        self.answer = QLineEdit(); self.answer.setEchoMode(QLineEdit.Password)
        self.answer.setPlaceholderText("Tu respuesta secreta")
        
        self.new_password = QLineEdit(); self.new_password.setEchoMode(QLineEdit.Password)
        self.new_password.setPlaceholderText("🔒 Nueva contraseña")
        
        self.new_password2 = QLineEdit(); self.new_password2.setEchoMode(QLineEdit.Password)
        self.new_password2.setPlaceholderText("🔒 Confirmar nueva contraseña")
        
        form.addRow("Usuario:", self.username)
        form.addRow("", self.btn_load)
        form.addRow("Pregunta:", self.question_lbl)
        form.addRow("Respuesta:", self.answer)
        form.addRow("Nueva clave:", self.new_password)
        form.addRow("Confirmar:", self.new_password2)
        
        main_lay.addLayout(form)
        main_lay.addStretch()
        
        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.accepted.connect(self._accept)
        btns.rejected.connect(self.reject)
        main_lay.addWidget(btns)
        
        for btn in btns.buttons():
            if btn.text().lower().startswith('ok'):
                btn.setText("Restablecer")
                btn.setStyleSheet("background-color: #2e7d32; color: white; font-weight: bold;")
            else:
                btn.setText("Cancelar")
                btn.setStyleSheet("background-color: #b0bec5; color: #212529;")

    def load_question(self):
        q = get_security_question(self.username.text().strip())
        self.question_lbl.setText(f"<b>{q}</b>" if q else "<span style='color:#c62828;'>Usuario no encontrado.</span>")

    def _accept(self):
        username = self.username.text().strip()
        if self.new_password.text() != self.new_password2.text():
            FloatingToast("Las contraseñas no coinciden", self, is_error=True).show()
            return
        if not username:
            FloatingToast("Indica el usuario", self, is_error=True).show()
            return
        ok = reset_password_by_security(username, self.answer.text(), self.new_password.text())
        if not ok:
            FloatingToast("Respuesta incorrecta", self, is_error=True).show()
            return
        log_action(username, "Recuperación de contraseña", "Contraseña restablecida con pregunta de seguridad")
        FloatingToast("✅ Clave actualizada", self).show()
        self.accept()

class LoginDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.user = None
        self.setWindowTitle("Bienvenido - Hospital Provincial")
        self.setFixedSize(450, 560)
        
        self.setStyleSheet("""
            QPushButton#LoginBtn { background-color: #1565c0; color: white; font-size: 12pt; font-weight: bold; border-radius: 6px; padding: 12px; }
            QPushButton#LoginBtn:hover { background-color: #0d47a1; }
            QPushButton#LinkBtn { color: #1e88e5; background: transparent; border: none; font-weight: bold; font-size: 10pt;}
            QPushButton#LinkBtn:hover { text-decoration: underline; color: #0d47a1; }
        """)
        
        main_lay = QVBoxLayout(self)
        main_lay.setContentsMargins(40, 40, 40, 40)
        main_lay.setSpacing(15)

        if LOGO_PATH and os.path.exists(LOGO_PATH):
            logo = QLabel()
            pix = QPixmap(LOGO_PATH)
            if not pix.isNull():
                logo.setPixmap(pix.scaledToWidth(130, Qt.SmoothTransformation))
                logo.setAlignment(Qt.AlignCenter)
                main_lay.addWidget(logo)

        title = QLabel("Iniciar Sesión")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 20pt; font-weight: 900; color: #1565c0; margin-bottom: 10px; border: none;")
        main_lay.addWidget(title)

        lbl_user = QLabel("Usuario")
        main_lay.addWidget(lbl_user)
        self.username = QLineEdit()
        self.username.setPlaceholderText("👤 Ingresa tu usuario")
        self.username_completer = QCompleter(list_usernames())
        self.username_completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.username_completer.setFilterMode(Qt.MatchContains)
        self.username.setCompleter(self.username_completer)
        main_lay.addWidget(self.username)

        lbl_pass = QLabel("Contraseña")
        main_lay.addWidget(lbl_pass)
        self.password = QLineEdit()
        self.password.setEchoMode(QLineEdit.Password)
        self.password.setPlaceholderText("🔒 Ingresa tu contraseña")
        main_lay.addWidget(self.password)

        main_lay.addSpacing(15)

        self.btn_login = QPushButton("Entrar al Sistema")
        self.btn_login.setObjectName("LoginBtn")
        self.btn_login.setCursor(QCursor(Qt.PointingHandCursor))
        main_lay.addWidget(self.btn_login)

        main_lay.addSpacing(10)

        links_lay = QHBoxLayout()
        self.btn_register = QPushButton("Crear cuenta")
        self.btn_register.setObjectName("LinkBtn")
        self.btn_register.setCursor(QCursor(Qt.PointingHandCursor))
        
        self.btn_recover = QPushButton("¿Olvidaste tu contraseña?")
        self.btn_recover.setObjectName("LinkBtn")
        self.btn_recover.setCursor(QCursor(Qt.PointingHandCursor))
        
        links_lay.addWidget(self.btn_recover)
        main_lay.addLayout(links_lay)

        main_lay.addStretch()

        self.lbl_error = QLabel("")
        self.lbl_error.setAlignment(Qt.AlignCenter)
        self.lbl_error.setStyleSheet("background-color: #d32f2f; color: white; padding: 10px; border-radius: 6px; font-weight: bold; font-size: 11pt;")
        self.lbl_error.hide() 
        main_lay.addWidget(self.lbl_error)

        self.btn_login.clicked.connect(self.try_login)
        self.btn_register.clicked.connect(self.register_user)
        self.btn_recover.clicked.connect(self.recover_password)
        
        self.password.returnPressed.connect(self.try_login)

    def try_login(self):
        user = authenticate_user(self.username.text().strip(), self.password.text())
        if not user:
            self.lbl_error.setText("❌ Usuario o contraseña incorrectos")
            self.lbl_error.show()
            self.password.clear() 
            self.password.setFocus()
            return
            
        self.lbl_error.hide()
        self.user = user
        log_action(user["username"], "Inicio de sesión", f"Rol: {user['role']}")
        self.accept()

    def register_user(self):
        dlg = RegisterUserDialog(self, allow_role_choice=False)
        if dlg.exec() == QDialog.Accepted:
            log_action(dlg.username.text().strip(), "Registro de usuario", "Usuario creado desde pantalla de inicio")
            FloatingToast("✅ Cuenta creada. Ya puedes iniciar.", self).show()
            self.username_completer.model().setStringList(list_usernames())

    def recover_password(self):
        RecoverPasswordDialog(self).exec()

class AddCatalogItemDialog(QDialog):
    def __init__(
        self,
        category: str,
        ars_name: str = "",
        parent=None,
        allow_category_navigation: bool = True,
    ):
        super().__init__(parent)
        self.category = category if category in ALL_CATEGORIES else "Medicamentos"
        self.ars_name = str(ars_name or "").strip()
        self.allow_category_navigation = bool(allow_category_navigation)
        self.setMinimumSize(660, 420 if self.allow_category_navigation else 340)

        lay = QVBoxLayout(self)
        lay.setContentsMargins(22, 20, 22, 20)
        lay.setSpacing(14)

        title = QLabel("Agregar ítem al catálogo" if self.allow_category_navigation else "Editar ítem")
        title.setStyleSheet("font-size: 18pt; font-weight: 900; color: #123F83;")
        title.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        lay.addWidget(title)

        subtitle = QLabel(
            "Selecciona primero el destino correcto. El panel azul indica exactamente dónde se guardará."
            if self.allow_category_navigation
            else "Modifica el nombre o precio del elemento seleccionado."
        )
        subtitle.setWordWrap(True)
        subtitle.setStyleSheet("color: #607289; font-size: 10.5pt;")
        subtitle.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        lay.addWidget(subtitle)

        if self.allow_category_navigation:
            navigation = QHBoxLayout()
            self.btn_previous = QPushButton("Anterior")
            self.btn_previous.setIcon(self.style().standardIcon(QStyle.SP_ArrowLeft))
            self.btn_previous.setToolTip("Ir a la categoría anterior")
            self.category_combo = QComboBox()
            self.category_combo.addItems(ALL_CATEGORIES)
            self.category_combo.setCurrentText(self.category)
            self.btn_next = QPushButton("Siguiente")
            self.btn_next.setIcon(self.style().standardIcon(QStyle.SP_ArrowRight))
            self.btn_next.setToolTip("Ir a la categoría siguiente")
            navigation.addWidget(self.btn_previous)
            navigation.addWidget(self.category_combo, 1)
            navigation.addWidget(self.btn_next)
            lay.addLayout(navigation)
            set_button_role(self.btn_previous, "neutral")
            set_button_role(self.btn_next, "neutral")
            self.btn_previous.clicked.connect(lambda: self.navigate(-1))
            self.btn_next.clicked.connect(lambda: self.navigate(1))
            self.category_combo.currentTextChanged.connect(self.set_category)

        self.destination_panel = QLabel()
        self.destination_panel.setWordWrap(True)
        self.destination_panel.setMinimumHeight(72)
        self.destination_panel.setMaximumHeight(78)
        self.destination_panel.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        self.destination_panel.setAlignment(Qt.AlignVCenter | Qt.AlignLeft)
        lay.addWidget(self.destination_panel)

        form = QFormLayout()
        form.setHorizontalSpacing(16)
        form.setVerticalSpacing(13)
        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("Descripción exacta del servicio o producto")
        self.price_spin = QDoubleSpinBox()
        self.price_spin.setRange(0, 1_000_000)
        self.price_spin.setDecimals(2)
        self.price_spin.setPrefix("RD$ ")
        self.price_spin.setButtonSymbols(QAbstractSpinBox.NoButtons)
        self.price_spin.setMinimumHeight(38)
        up = QToolButton(); up.setText("+"); up.setToolTip("Aumentar precio")
        dn = QToolButton(); dn.setText("−"); dn.setToolTip("Disminuir precio")
        for button in (up, dn):
            button.setFixedSize(40, 38)
            button.setStyleSheet(
                "QToolButton { border: 1px solid #B7C9E2; border-radius: 7px; "
                "background: #F5F8FC; color: #174A96; font-size: 16pt; font-weight: 900; } "
                "QToolButton:hover { background: #E8F1FD; }"
            )
        hl = QHBoxLayout(); hl.setContentsMargins(0, 0, 0, 0)
        hl.addWidget(self.price_spin, 1); hl.addWidget(dn); hl.addWidget(up)
        price_box = QWidget(); price_box.setLayout(hl)
        form.addRow("Nombre:", self.name_edit)
        form.addRow("Precio:", price_box)
        lay.addLayout(form)
        up.clicked.connect(lambda: self.price_spin.setValue(self.price_spin.value() + 1.0))
        dn.clicked.connect(lambda: self.price_spin.setValue(max(0.0, self.price_spin.value() - 1.0)))

        btns = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        self.btn_save = btns.button(QDialogButtonBox.Save)
        self.btn_save.setText("Guardar")
        self.btn_cancel = btns.button(QDialogButtonBox.Cancel)
        self.btn_cancel.setText("Cancelar")
        btns.accepted.connect(self.confirm_and_accept)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)
        for btn in btns.buttons():
            set_button_role(btn, "success" if btn is self.btn_save else "neutral")

        self.name_edit.returnPressed.connect(self.confirm_and_accept)
        self._update_destination()

    def available_categories(self):
        if self.ars_name:
            return list(ALL_CATEGORIES)
        return list(UNIVERSAL_CATEGORIES)

    def navigate(self, direction: int):
        categories = self.available_categories()
        if not categories:
            return
        current = categories.index(self.category) if self.category in categories else 0
        self.category_combo.setCurrentText(categories[(current + direction) % len(categories)])

    def set_category(self, category: str):
        if category in ARS_CATEGORIES and not self.ars_name:
            FloatingToast(
                "Selecciona una ARS en la pantalla principal para usar esta categoría.",
                self,
                is_error=True,
            ).show()
            if hasattr(self, "category_combo"):
                self.category_combo.blockSignals(True)
                self.category_combo.setCurrentText(self.category)
                self.category_combo.blockSignals(False)
            return
        if category in ALL_CATEGORIES:
            self.category = category
            self._update_destination()

    def _update_destination(self):
        color = CAT_COLORS.get(self.category, "#174A96")
        if self.category in UNIVERSAL_CATEGORIES:
            catalog_type = "CATÁLOGO UNIVERSAL"
            destination = f"Disponible para todas las ARS"
        else:
            catalog_type = "CATÁLOGO POR ARS"
            destination = f"ARS de destino: {self.ars_name or 'NO SELECCIONADA'}"
        self.destination_panel.setText(
            f"DESTINO  ·  {self.category}\n{catalog_type}  •  {destination}"
        )
        self.destination_panel.setStyleSheet(
            f"QLabel {{ background: {color}; color: white; border-radius: 10px; "
            "padding: 11px 15px; font-size: 10.5pt; font-weight: 800; }}"
        )
        suffix = f" · {self.ars_name}" if self.category in ARS_CATEGORIES else ""
        self.btn_save.setText("Guardar ítem")
        self.setWindowTitle(f"Agregar en {self.category}{suffix}")

    def confirm_and_accept(self):
        name = self.name_edit.text().strip()
        if not name:
            FloatingToast("Escribe el nombre del ítem", self, is_error=True).show()
            self.name_edit.setFocus()
            return
        if self.category in ARS_CATEGORIES and not self.ars_name:
            FloatingToast("Selecciona una ARS antes de guardar", self, is_error=True).show()
            return

        current = (
            get_universal(self.category)
            if self.category in UNIVERSAL_CATEGORIES
            else get_ars_items(self.category, self.ars_name)
        )
        existing = next((item for item in current if item.casefold() == name.casefold()), None)
        action = "actualizará el precio existente" if existing else "creará un nuevo ítem"
        destination = (
            "Catálogo universal"
            if self.category in UNIVERSAL_CATEGORIES
            else f"ARS: {self.ars_name}"
        )
        answer = QMessageBox.question(
            self,
            "Confirmar destino",
            f"Se {action}.\n\nCategoría: {self.category}\n{destination}\n"
            f"Nombre: {name}\nPrecio: RD$ {self.price_spin.value():,.2f}\n\n¿Guardar aquí?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes,
        )
        if answer == QMessageBox.Yes:
            super().accept()

    def values(self):
        return self.name_edit.text().strip(), float(self.price_spin.value())

class CatalogEditorDialog(QDialog):
    def __init__(self, category: str, ars_name: str = "", parent=None):
        super().__init__(parent)
        self.category = category if category in ALL_CATEGORIES else "Medicamentos"
        self.ars_name = (ars_name or "").strip()
        self.has_changes = False
        self._dirty = False
        self._loading = False
        self._changing_category = False

        self.setMinimumSize(900, 640)
        lay = QVBoxLayout(self)
        lay.setContentsMargins(18, 16, 18, 16)
        lay.setSpacing(10)

        nav_row = QHBoxLayout()
        self.category_combo = QComboBox()
        self.category_combo.addItems(ALL_CATEGORIES)
        self.category_combo.setCurrentText(self.category)

        self.btn_prev_category = QPushButton("Anterior")
        self.btn_prev_category.setIcon(self.style().standardIcon(QStyle.SP_ArrowLeft))
        self.btn_next_category = QPushButton("Siguiente")
        self.btn_next_category.setIcon(self.style().standardIcon(QStyle.SP_ArrowRight))
        self.btn_prev_category.setToolTip("Ir a la categoría anterior")
        self.btn_next_category.setToolTip("Ir a la categoría siguiente")

        self.lbl_ars_context = QLabel("")
        self.lbl_ars_context.setStyleSheet("font-weight: bold; color: #1565c0;")

        nav_row.addWidget(QLabel("Catálogo:"))
        nav_row.addWidget(self.category_combo, 1)
        nav_row.addWidget(self.btn_prev_category)
        nav_row.addWidget(self.btn_next_category)
        nav_row.addSpacing(12)
        nav_row.addWidget(self.lbl_ars_context)
        lay.addLayout(nav_row)

        self.search_edit = QLineEdit()
        self.search_edit.textChanged.connect(lambda t: filter_table_widget(self.table, t))
        lay.addWidget(self.search_edit)

        self.table = QTableWidget(0, 2)
        self.table.verticalHeader().setVisible(False)
        self.table.setAlternatingRowColors(True)
        self.table.setShowGrid(False)
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.Fixed)
        self.table.setColumnWidth(1, 140)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.table.itemChanged.connect(self._mark_dirty)
        lay.addWidget(self.table)

        actions = QHBoxLayout()
        self.btn_clear_prices = QPushButton("Poner precio en RD$ 0.00")
        self.btn_delete_rows = QPushButton("Eliminar selección")
        self.btn_clear_prices.setToolTip("Asigna RD$ 0.00 a los precios seleccionados; se aplica al guardar")
        self.btn_delete_rows.setToolTip("Elimina los elementos seleccionados del catálogo")
        actions.addWidget(self.btn_clear_prices)
        actions.addWidget(self.btn_delete_rows)
        actions.addStretch(1)
        lay.addLayout(actions)

        bottom = QHBoxLayout()
        self.btn_save_changes = QPushButton("Guardar cambios")
        self.btn_close = QPushButton("Cerrar")
        self.lbl_pending_changes = QLabel("Sin cambios pendientes")
        self.lbl_pending_changes.setStyleSheet("color: #607289; font-weight: 600;")
        bottom.addWidget(self.lbl_pending_changes)
        bottom.addStretch(1)
        bottom.addWidget(self.btn_save_changes)
        bottom.addWidget(self.btn_close)
        lay.addLayout(bottom)

        self.btn_prev_category.clicked.connect(lambda: self.navigate_category(-1))
        self.btn_next_category.clicked.connect(lambda: self.navigate_category(1))
        self.category_combo.currentTextChanged.connect(self.on_category_changed)
        self.btn_clear_prices.clicked.connect(self.clear_selected_prices)
        self.btn_delete_rows.clicked.connect(self.delete_selected_rows)
        self.btn_save_changes.clicked.connect(lambda: self.save_changes(close_after=False))
        self.btn_close.clicked.connect(self.close)
        self.table.itemSelectionChanged.connect(self._update_selection_actions)

        set_button_role(self.btn_prev_category, 'neutral')
        set_button_role(self.btn_next_category, 'neutral')
        set_button_role(self.btn_clear_prices, 'warning')
        set_button_role(self.btn_delete_rows, 'danger')
        set_button_role(self.btn_save_changes, 'success')
        set_button_role(self.btn_close, 'neutral')

        self._update_dialog_context()
        self.load_rows()
        self._update_selection_actions()
        self._update_dirty_state()

    def _singular_label(self, category: str) -> str:
        labels = {
            "Medicamentos": "Medicamento",
            "Materiales": "Material",
            "Laboratorios": "Laboratorio",
            "Imágenes": "Imagen",
            "Procedimientos": "Procedimiento",
            "Honorarios": "Honorario",
        }
        return labels.get(category, category)

    def _mark_dirty(self, *_args):
        if not self._loading:
            self._dirty = True
            self._update_dirty_state()

    def _update_dirty_state(self):
        dirty = bool(getattr(self, "_dirty", False))
        self.btn_save_changes.setEnabled(dirty)
        self.lbl_pending_changes.setText("● Hay cambios sin guardar" if dirty else "Sin cambios pendientes")
        self.lbl_pending_changes.setStyleSheet(
            "color: #c25e00; font-weight: 800;" if dirty else "color: #607289; font-weight: 600;"
        )

    def _update_selection_actions(self):
        count = len({index.row() for index in self.table.selectedIndexes()})
        self.btn_clear_prices.setEnabled(count > 0)
        self.btn_delete_rows.setEnabled(count > 0)
        self.btn_clear_prices.setText(f"Poner precio en RD$ 0.00 ({count})")
        self.btn_delete_rows.setText(f"Eliminar selección ({count})")

    def _update_dialog_context(self):
        suffix = f" - {self.ars_name}" if self.category in ARS_CATEGORIES and self.ars_name else ""
        self.setWindowTitle(f"Gestión de {self.category}{suffix}")
        self.search_edit.setPlaceholderText(f"🔍 Buscar en {self.category}...")
        self.table.setHorizontalHeaderLabels([self._singular_label(self.category), "Precio"])
        if self.category in ARS_CATEGORIES:
            self.lbl_ars_context.setText(f"ARS seleccionada: {self.ars_name or 'Ninguna'}")
        else:
            self.lbl_ars_context.setText("Catálogo universal")

    def navigate_category(self, direction: int):
        idx = self.category_combo.currentIndex()
        if idx < 0:
            return
        new_idx = (idx + direction) % self.category_combo.count()
        self.category_combo.setCurrentIndex(new_idx)

    def _confirm_discard_or_save(self) -> bool:
        if not getattr(self, '_dirty', False):
            return True

        res = QMessageBox.question(
            self,
            "Cambios sin guardar",
            "Hay cambios sin guardar en este catálogo.\n\n¿Deseas guardarlos antes de cambiar?",
            QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
            QMessageBox.Yes
        )
        if res == QMessageBox.Cancel:
            return False
        if res == QMessageBox.Yes:
            return self.save_changes(close_after=False, silent=True)

        self._dirty = False
        return True

    def on_category_changed(self, new_category: str):
        if self._changing_category:
            return
        if new_category == self.category:
            return

        if new_category in ARS_CATEGORIES and not self.ars_name:
            FloatingToast("Selecciona una ARS en la pantalla principal antes de abrir categorías por ARS", self, is_error=True).show()
            self._changing_category = True
            self.category_combo.setCurrentText(self.category)
            self._changing_category = False
            return

        if not self._confirm_discard_or_save():
            self._changing_category = True
            self.category_combo.setCurrentText(self.category)
            self._changing_category = False
            return

        self.category = new_category
        self.search_edit.clear()
        self._update_dialog_context()
        self.load_rows()

    def _get_data(self):
        return get_universal(self.category) if self.category in UNIVERSAL_CATEGORIES else get_ars_items(self.category, self.ars_name)

    def load_rows(self):
        self._loading = True
        self.table.setRowCount(0)
        for name, price in sorted(self._get_data().items(), key=lambda x: x[0]):
            r = self.table.rowCount()
            self.table.insertRow(r)
            name_item = QTableWidgetItem(name)
            name_item.setFlags(name_item.flags() & ~Qt.ItemIsEditable)
            self.table.setItem(r, 0, name_item)
            shown_price = get_effective_price(self.category, price)
            price_item = QTableWidgetItem(f"{shown_price:,.2f}")
            price_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            self.table.setItem(r, 1, price_item)
        self._loading = False
        self._dirty = False
        self._update_dirty_state()
        self._update_selection_actions()

    def clear_selected_prices(self):
        rows = sorted({idx.row() for idx in self.table.selectedIndexes()})
        if not rows:
            FloatingToast("Selecciona al menos una fila", self, is_error=True).show()
            return
        for r in rows:
            item = self.table.item(r, 1)
            if item is None:
                item = QTableWidgetItem("0.00")
                self.table.setItem(r, 1, item)
            else:
                item.setText("0.00")
        self._dirty = True
        self._update_dirty_state()

    def delete_selected_rows(self):
        rows = sorted({idx.row() for idx in self.table.selectedIndexes()}, reverse=True)
        if not rows:
            FloatingToast("Selecciona al menos una fila", self, is_error=True).show()
            return
        names = [self.table.item(r, 0).text().strip() for r in rows]
        preview = "\n".join(f"• {name}" for name in names[:5])
        if len(names) > 5:
            preview += f"\n• … y {len(names) - 5} más"
        answer = QMessageBox.question(
            self,
            "Confirmar eliminación",
            f"Se eliminarán {len(names)} elemento(s) de {self.category}:\n\n{preview}\n\n¿Deseas continuar?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if answer != QMessageBox.Yes:
            return
        with db_connect() as con:
            for r in rows:
                name = self.table.item(r, 0).text().strip()
                if self.category in UNIVERSAL_CATEGORIES:
                    con.execute("UPDATE universal_items SET is_active=0 WHERE categoria=%s AND nombre=%s", (self.category, name))
                else:
                    con.execute("UPDATE ars_items SET is_active=0 WHERE categoria=%s AND nombre=%s AND ars_id=(SELECT id FROM ars WHERE nombre=%s)", (self.category, name, self.ars_name))
                self.table.removeRow(r)
        self.has_changes = True
        self._dirty = False
        self._update_dirty_state()
        self._update_selection_actions()
        FloatingToast("✅ Ítem(s) eliminado(s)", self).show()

    def save_changes(self, close_after: bool = False, silent: bool = False) -> bool:
        try:
            with db_connect() as con:
                for r in range(self.table.rowCount()):
                    name_item = self.table.item(r, 0)
                    price_item = self.table.item(r, 1)
                    if name_item is None:
                        continue
                    name = name_item.text().strip()
                    price_text = price_item.text() if price_item else ""
                    price = _parse_price(price_text)
                    if price is None:
                        FloatingToast(f"Precio inválido en '{name}'", self, is_error=True).show()
                        return False

                    if self.category == 'Medicamentos':
                        price = round(float(price) / 1.20, 2)

                    if self.category in UNIVERSAL_CATEGORIES:
                        con.execute(
                            "INSERT INTO universal_items(categoria, nombre, precio, is_active) VALUES(%s,%s,%s, 1) "
                            "ON CONFLICT(categoria, nombre) DO UPDATE SET precio=EXCLUDED.precio, is_active=1",
                            (self.category, name, float(price))
                        )
                    else:
                        con.execute(
                            """
                            INSERT INTO ars_items(ars_id, categoria, nombre, precio, is_active)
                            SELECT a.id, %s, %s, %s, 1 FROM ars a WHERE a.nombre=%s
                            ON CONFLICT(ars_id, categoria, nombre) DO UPDATE SET precio=EXCLUDED.precio, is_active=1
                            """,
                            (self.category, name, float(price), self.ars_name)
                        )
            self.has_changes = True
            self._dirty = False
            self._update_dirty_state()
            if not silent:
                FloatingToast("✅ Cambios guardados correctamente", self).show()
            if close_after:
                self.accept()
            return True
        except Exception as e:
            QMessageBox.critical(self, "Catálogo", f"No se pudieron guardar los cambios:\n{e}")
            return False

    def closeEvent(self, event):
        if getattr(self, '_dirty', False):
            res = QMessageBox.question(
                self,
                "Cambios sin guardar",
                "Hay cambios sin guardar.\n\n¿Deseas guardarlos antes de cerrar?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel,
                QMessageBox.Yes
            )
            if res == QMessageBox.Cancel:
                event.ignore()
                return
            if res == QMessageBox.Yes and not self.save_changes(close_after=False, silent=True):
                event.ignore()
                return
        super().closeEvent(event)

class ARSManagerDialog(QDialog):
    def __init__(self, current_user: dict, parent=None):
        super().__init__(parent)
        self.current_user = current_user
        self.setWindowTitle("Gestión de ARS")
        self.setMinimumSize(720, 460)
        self.resize(860, 540)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 18, 20, 18)
        layout.setSpacing(12)
        title = QLabel("Gestión de ARS")
        title.setStyleSheet("font-size: 18pt; font-weight: 900; color: #123F83;")
        subtitle = QLabel("Configura tarifas, catálogos e importaciones de cada aseguradora.")
        subtitle.setStyleSheet("color: #607289; font-size: 10.5pt;")
        layout.addWidget(title)
        layout.addWidget(subtitle)
        row = QHBoxLayout()
        self.ars_combo = QComboBox(); self.refresh_ars()
        row.addWidget(QLabel("ARS:")); row.addWidget(self.ars_combo, 1)
        layout.addLayout(row)
        form = QFormLayout()
        self.sala_spin = QDoubleSpinBox(); self.sala_spin.setRange(0, 1_000_000); self.sala_spin.setDecimals(2); self.sala_spin.setButtonSymbols(QAbstractSpinBox.NoButtons)
        up = QToolButton(); up.setIcon(self.style().standardIcon(QStyle.SP_ArrowUp)); up.setIconSize(QSize(14,14))
        dn = QToolButton(); dn.setIcon(self.style().standardIcon(QStyle.SP_ArrowDown)); dn.setIconSize(QSize(14,14))
        hl = QHBoxLayout(); hl.addWidget(self.sala_spin); hl.addWidget(up); hl.addWidget(dn); hl.addStretch(1)
        box = QWidget(); box.setLayout(hl)
        form.addRow("Sala Emergencia:", box)
        layout.addLayout(form)
        row1 = QHBoxLayout()
        self.btn_add = QPushButton("+ Agregar ARS")
        self.btn_del = QPushButton("Eliminar ARS")
        self.btn_save = QPushButton("Guardar cambios")
        row1.addWidget(self.btn_add); row1.addWidget(self.btn_del); row1.addStretch(1); row1.addWidget(self.btn_save)
        layout.addLayout(row1)
        catalogs = QGroupBox("Catálogos y servicios")
        row2 = QGridLayout(catalogs)
        self.btn_import_ars_word = QPushButton("Importar catálogo desde Word")
        self.btn_edit_labs = QPushButton("Laboratorios")
        self.btn_edit_imgs = QPushButton("Imágenes")
        self.btn_edit_procs = QPushButton("Procedimientos")
        self.btn_edit_hono = QPushButton("Honorarios")
        row2.addWidget(self.btn_edit_labs, 0, 0); row2.addWidget(self.btn_edit_imgs, 0, 1)
        row2.addWidget(self.btn_edit_procs, 1, 0); row2.addWidget(self.btn_edit_hono, 1, 1)
        row2.addWidget(self.btn_import_ars_word, 2, 0, 1, 2)
        layout.addWidget(catalogs)
        if user_is_admin(self.current_user):
            self.btn_toggle_migrate = QPushButton("Mostrar migración entre ARS")
            set_button_role(self.btn_toggle_migrate, 'warning')
            layout.addWidget(self.btn_toggle_migrate)
            mig = QGroupBox("Migrar datos entre ARS")
            mig.setStyleSheet("QGroupBox { font-weight: bold; }")
            mlay = QHBoxLayout(mig)
            self.mig_src = QComboBox(); self.refresh_ars()
            self.mig_dst = QComboBox(); self.refresh_ars()
            self.mig_src.setPlaceholderText("ARS origen")
            self.mig_dst.setPlaceholderText("ARS destino")
            mlay.addWidget(QLabel("Origen:"))
            mlay.addWidget(self.mig_src, 1)
            mlay.addSpacing(12)
            mlay.addWidget(QLabel("Destino:"))
            mlay.addWidget(self.mig_dst, 1)
            self.btn_exec_migrate = QPushButton("Ejecutar migración")
            set_button_role(self.btn_exec_migrate, 'warning')
            mlay.addWidget(self.btn_exec_migrate)
            mig.hide()
            layout.addWidget(mig)
            self.btn_toggle_migrate.clicked.connect(lambda: mig.setVisible(not mig.isVisible()))
            self.btn_exec_migrate.clicked.connect(self.migrate_ars)
        self.ars_combo.currentTextChanged.connect(self.on_ars_change)
        self.btn_add.clicked.connect(self.add_ars)
        self.btn_del.clicked.connect(self.delete_ars_clicked)
        self.btn_save.clicked.connect(self.save_price)
        self.btn_import_ars_word.clicked.connect(self.import_ars_word)
        self.btn_edit_labs.clicked.connect(lambda: self.open_category_editor('Laboratorios'))
        self.btn_edit_imgs.clicked.connect(lambda: self.open_category_editor('Imágenes'))
        self.btn_edit_procs.clicked.connect(lambda: self.open_category_editor('Procedimientos'))
        self.btn_edit_hono.clicked.connect(lambda: self.open_category_editor('Honorarios'))
        up.clicked.connect(lambda: self.sala_spin.setValue(self.sala_spin.value()+SALA_STEP))
        dn.clicked.connect(lambda: self.sala_spin.setValue(max(0.0, self.sala_spin.value()-SALA_STEP)))
        if self.ars_combo.count() > 0:
            self.on_ars_change(self.ars_combo.currentText())
        dbtns = QDialogButtonBox(QDialogButtonBox.Close)
        dbtns.rejected.connect(self.reject)
        layout.addWidget(dbtns)
        set_button_role(self.btn_add, 'success')
        set_button_role(self.btn_del, 'danger')
        set_button_role(self.btn_save, 'success')
        set_button_role(self.btn_import_ars_word, 'warning')
        set_button_role(self.btn_edit_labs, 'info')
        set_button_role(self.btn_edit_imgs, 'info')
        set_button_role(self.btn_edit_procs, 'info')
        set_button_role(self.btn_edit_hono, 'info')
        for btn in dbtns.buttons(): set_button_role(btn, 'neutral')

    def refresh_ars(self):
        items = ars_list()
        self.ars_combo.clear(); self.ars_combo.addItems(items)
        if hasattr(self, 'mig_src'):
            self.mig_src.clear(); self.mig_src.addItems(items)
        if hasattr(self, 'mig_dst'):
            self.mig_dst.clear(); self.mig_dst.addItems(items)

    def on_ars_change(self, name):
        self.sala_spin.setValue(get_emergency_price(name) if name else 0.0)

    def migrate_ars(self):
        src = self.mig_src.currentText().strip()
        dst = self.mig_dst.currentText().strip()
        if not src or not dst or src == dst:
            FloatingToast("Selecciona ARS origen y destino distintas", self, is_error=True).show()
            return
        confirm = QMessageBox.question(self, "Confirmar migración", f"¿Migrar todos los ítems de '{src}' a '{dst}'?\n\nSolo se agregarán los ítems que NO existan en '{dst}'. Los ítems ya existentes no se modificarán.")
        if confirm != QMessageBox.Yes:
            return
        try:
            with db_connect() as con:
                for cat in ARS_CATEGORIES:
                    con.execute("""
                        INSERT INTO ars_items(ars_id, categoria, nombre, precio, is_active)
                        SELECT a2.id, ai.categoria, ai.nombre, ai.precio, ai.is_active
                        FROM ars_items ai
                        JOIN ars a1 ON a1.id = ai.ars_id
                        JOIN ars a2 ON a2.nombre = %s
                        WHERE a1.nombre = %s AND ai.categoria = %s
                        ON CONFLICT(ars_id, categoria, nombre) DO NOTHING
                    """, (dst, src, cat))
            log_action(self.current_user["username"], "Migrar ARS", f"{src} -> {dst}")
            FloatingToast(f"✅ Datos migrados: {src} → {dst}", self).show()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo migrar: {e}")

    def add_ars(self):
        name, ok = QInputDialogWithText.get(self, "Nueva ARS", "Nombre de la ARS:")
        if not ok or not name.strip(): return
        if name.strip() in ars_list():
            FloatingToast("Ya existe una ARS con ese nombre", self, is_error=True).show()
            return
        upsert_ars(name.strip(), 0.0)
        log_action(self.current_user["username"], "Agregar ARS", name.strip())
        self.refresh_ars(); self.ars_combo.setCurrentText(name.strip())
        FloatingToast("✅ ARS agregada", self).show()

    def delete_ars_clicked(self):
        name = self.ars_combo.currentText()
        if not name: return
        if QMessageBox.question(self, "Eliminar", f"¿Eliminar '{name}'?") != QMessageBox.Yes: return
        delete_ars(name)
        log_action(self.current_user["username"], "Eliminar ARS", name)
        self.refresh_ars()
        FloatingToast("✅ ARS eliminada", self).show()

    def save_price(self):
        name = self.ars_combo.currentText()
        if not name: return
        set_emergency_price(name, self.sala_spin.value())
        log_action(self.current_user["username"], "Actualizar sala ARS", f"{name} -> {self.sala_spin.value():.2f}")
        FloatingToast("✅ Precio de sala guardado", self).show()

    def import_ars_word(self):
        if Document is None:
            QMessageBox.warning(self, "Importar Word", "Falta dependencia: python-docx.\nInstala con:\n\npip install python-docx")
            return
        ars_name = self.ars_combo.currentText().strip()
        if not ars_name:
            FloatingToast("Selecciona una ARS primero", self, is_error=True).show()
            return
        path, _ = QFileDialog.getOpenFileName(self, "Seleccionar archivo Word", "", "Word (*.docx)")
        if not path: return
        try:
            summary = import_word_to_ars_catalog(path, ars_name)
            log_action(self.current_user['username'], "Importar Word ARS", f"ARS: {ars_name} | Archivo: {os.path.basename(path)}")
            QMessageBox.information(
                self, "Importación completada",
                (f"ARS: {ars_name}\n\n"
                 f"Laboratorios — Actualizados: {summary['Laboratorios']['updated']} | Nuevos: {summary['Laboratorios']['inserted']}\n"
                 f"Imágenes     — Actualizados: {summary['Imágenes']['updated']} | Nuevos: {summary['Imágenes']['inserted']}\n"
                 f"Procedimientos — Actualizados: {summary['Procedimientos']['updated']} | Nuevos: {summary['Procedimientos']['inserted']}\n"
                 f"Honorarios     — Actualizados: {summary['Honorarios']['updated']} | Nuevos: {summary['Honorarios']['inserted']}")
            )
        except Exception as e:
            QMessageBox.critical(self, "Importar Word", f"Error al procesar el archivo:\n{e}")

    def open_category_editor(self, category: str):
        ars_name = self.ars_combo.currentText().strip()
        if not ars_name:
            FloatingToast("Selecciona una ARS primero", self, is_error=True).show()
            return
        CatalogEditorDialog(category, ars_name, self).exec()

class UsersAdminDialog(QDialog):
    def __init__(self, current_user: dict, parent=None):
        super().__init__(parent)
        self.current_user = current_user
        self.setWindowTitle("Administración de usuarios")
        self.setMinimumSize(1050, 600)
        lay = QVBoxLayout(self)
        
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Buscar usuario...")
        self.search_edit.textChanged.connect(lambda t: filter_table_widget(self.table, t))
        lay.addWidget(self.search_edit)

        self.table = QTableWidget(0, 8)
        self.table.setHorizontalHeaderLabels(["Usuario", "Nombre", "Rol", "Activo", "Sesión actual", "Creado", "Último login", "Recibos"])
        self.table.verticalHeader().hide()
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Stretch)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        # ---> ACTIVAR MENÚ CONTEXTUAL (CLIC DERECHO) <---
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_context_menu)
        
        lay.addWidget(self.table)
        
        row = QHBoxLayout()
        self.btn_refresh = QPushButton("Actualizar")
        self.btn_add = QPushButton("Crear usuario")
        self.btn_toggle = QPushButton("Activar / Desactivar")
        self.btn_role = QPushButton("Cambiar rol")
        self.btn_reset = QPushButton("Restablecer contraseña")
        self.btn_remote_logout = QPushButton("Cerrar sesión remota")
        row.addWidget(self.btn_refresh); row.addWidget(self.btn_add); row.addWidget(self.btn_toggle); row.addWidget(self.btn_role); row.addWidget(self.btn_reset); row.addWidget(self.btn_remote_logout); row.addStretch(1)
        lay.addLayout(row)
        
        self.btn_refresh.clicked.connect(self.load_users)
        self.btn_add.clicked.connect(self.create_user)
        self.btn_toggle.clicked.connect(self.toggle_active)
        self.btn_role.clicked.connect(self.change_role)
        self.btn_reset.clicked.connect(self.reset_password)
        self.btn_remote_logout.clicked.connect(self.close_remote_session)
        close_btn = QDialogButtonBox(QDialogButtonBox.Close)
        close_btn.rejected.connect(self.reject)
        lay.addWidget(close_btn)
        self.load_users()
        set_button_role(self.btn_refresh, 'info')
        set_button_role(self.btn_add, 'success')
        set_button_role(self.btn_toggle, 'warning')
        set_button_role(self.btn_role, 'info')
        set_button_role(self.btn_reset, 'danger')
        set_button_role(self.btn_remote_logout, 'warning')
        if not user_is_admin(self.current_user):
            self.btn_add.setEnabled(False)
            self.btn_toggle.setEnabled(False)
            self.btn_role.setEnabled(False)
            self.btn_reset.setEnabled(False)
        for btn in close_btn.buttons(): set_button_role(btn, 'neutral')

    def load_users(self):
        counts = get_user_receipt_counts()
        active_sessions = get_active_sessions_map()
        self.table.setRowCount(0)
        for user in list_users():
            r = self.table.rowCount(); self.table.insertRow(r)
            username = user["username"]
            session_status = "Activa" if username in active_sessions else "No activa"
            values = [
                username,
                user["full_name"],
                user["role"],
                "Sí" if int(user["is_active"]) else "No",
                session_status,
                user["created_at"] or "",
                user.get("last_login") or "",
                counts.get(username, 0)
            ]
            for c, val in enumerate(values):
                item = QTableWidgetItem(str(val))
                if c == 4 and session_status == "Activa":
                    item.setForeground(QColor("#2e7d32"))
                    font = item.font()
                    font.setBold(True)
                    item.setFont(font)
                self.table.setItem(r, c, item)

    def _selected_username(self):
        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un usuario", self, is_error=True).show()
            return None
        return self.table.item(row, 0).text().strip()

    def create_user(self):
        if not user_is_admin(self.current_user):
            FloatingToast("Solo el administrador puede crear usuarios", self, is_error=True).show()
            return
        dlg = RegisterUserDialog(self, allow_role_choice=True)
        if dlg.exec() == QDialog.Accepted:
            log_action(self.current_user["username"], "Crear usuario", dlg.username.text().strip())
            FloatingToast("✅ Usuario creado", self).show()
            self.load_users()

    def toggle_active(self):
        if not user_is_admin(self.current_user):
            FloatingToast("Solo el administrador puede activar o desactivar usuarios", self, is_error=True).show()
            return
        username = self._selected_username()
        if not username: return
        user = get_user(username)
        if not user: return
        new_status = not bool(int(user["is_active"]))
        set_user_active(username, new_status)
        log_action(self.current_user["username"], "Cambiar estado usuario", f"{username} -> {'activo' if new_status else 'inactivo'}")
        FloatingToast("✅ Estado actualizado", self).show()
        self.load_users()

    def change_role(self):
        if not user_is_admin(self.current_user):
            FloatingToast("Solo el administrador puede cambiar roles", self, is_error=True).show()
            return
        username = self._selected_username()
        if not username: return
        user = get_user(username)
        if not user: return
        
        roles = [ROLE_AUX, ROLE_ADMIN, ROLE_AUDIT]
        current_idx = roles.index(user["role"]) if user["role"] in roles else 0
        new_role, ok = QInputDialog.getItem(self, "Cambiar rol", "Selecciona el nuevo rol:", roles, current_idx, False)
        
        if ok and new_role and new_role != user["role"]:
            update_user_role(username, new_role)
            log_action(self.current_user["username"], "Cambiar rol", f"{username} -> {new_role}")
            FloatingToast("✅ Rol actualizado", self).show()
            self.load_users()

    def reset_password(self):
        if not user_is_admin(self.current_user):
            FloatingToast("Solo el administrador puede restablecer contraseñas", self, is_error=True).show()
            return
        username = self._selected_username()
        if not username: return
        new_pass, ok = QInputDialogWithText.get(self, "Restablecer contraseña", f"Nueva contraseña para {username}:")
        if not ok or not new_pass.strip(): return
        admin_reset_password(username, new_pass.strip())
        log_action(self.current_user["username"], "Reset de contraseña", f"Contraseña restablecida para {username}")
        FloatingToast("✅ Contraseña restablecida", self).show()

    def show_context_menu(self, pos):
        row = self.table.rowAt(pos.y())
        if row < 0: return
        self.table.selectRow(row)
        
        menu = QMenu(self)
        
        if user_is_admin(self.current_user):
            a_toggle = QAction("🔄 Activar / Desactivar", self)
            a_role = QAction("🏷️ Cambiar rol", self)
            a_reset = QAction("🔑 Restablecer contraseña", self)
            
            a_toggle.triggered.connect(self.toggle_active)
            a_role.triggered.connect(self.change_role)
            a_reset.triggered.connect(self.reset_password)
            
            menu.addAction(a_toggle)
            menu.addAction(a_role)
            menu.addAction(a_reset)
            menu.addSeparator()

        a_remote = QAction("🚪 Cerrar sesión remota", self)
        a_remote.triggered.connect(self.close_remote_session)
        menu.addAction(a_remote)
        
        if user_is_admin(self.current_user):
            menu.addSeparator()
            a_delete = QAction("❌ Eliminar usuario", self)
            a_delete.triggered.connect(self.delete_user)
            
            font = a_delete.font()
            font.setBold(True)
            a_delete.setFont(font)
            
            menu.addAction(a_delete)
            
        menu.exec(self.table.viewport().mapToGlobal(pos))

    def close_remote_session(self):
        if not user_can_manage_sessions(self.current_user):
            FloatingToast("No tienes permiso para cerrar sesiones remotas", self, is_error=True).show()
            return

        username = self._selected_username()
        if not username:
            return

        if username == self.current_user.get("username"):
            FloatingToast("No puedes cerrar remotamente tu propia sesión activa", self, is_error=True).show()
            return

        active_sessions = get_active_sessions_map()
        if username not in active_sessions:
            FloatingToast("Ese usuario no tiene una sesión activa en este momento", self, is_error=True).show()
            self.load_users()
            return

        confirm = QMessageBox.question(
            self,
            "Cerrar sesión remota",
            f"¿Deseas cerrar la sesión activa del usuario '{username}'?\n\nSi ese usuario tiene el sistema abierto, volverá al login automáticamente.",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm != QMessageBox.Yes:
            return

        try:
            request_remote_logout(username, self.current_user.get("username", "Sistema"), "Sesión cerrada")
            log_action(self.current_user["username"], "Cerrar sesión remota", f"Usuario: {username}")
            FloatingToast("✅ Orden de cierre enviada", self).show()
            self.load_users()
        except Exception as e:
            QMessageBox.critical(self, "Sesiones", f"No se pudo cerrar la sesión remota:\n{e}")

    def delete_user(self):
        if not user_is_admin(self.current_user):
            FloatingToast("Solo el administrador puede eliminar usuarios", self, is_error=True).show()
            return
        username = self._selected_username()
        if not username: return
        
        if username == self.current_user.get("username"):
            FloatingToast("No puedes eliminar tu propio usuario activo", self, is_error=True).show()
            return
            
        res = QMessageBox.question(
            self, 
            "Eliminar Usuario", 
            f"¿Estás seguro de que deseas ELIMINAR permanentemente al usuario '{username}'?\n\nEsta acción no se puede deshacer.", 
            QMessageBox.Yes | QMessageBox.No
        )
        
        if res == QMessageBox.Yes:
            try:
                delete_user_db(username)
                log_action(self.current_user["username"], "Eliminar usuario", f"Usuario eliminado del sistema: {username}")
                FloatingToast("✅ Usuario eliminado", self).show()
                self.load_users()
            except Exception as e:
                QMessageBox.critical(self, "Error", f"No se pudo eliminar el usuario:\n{e}")

class HistoryDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Historial de acciones")
        self.setMinimumSize(950, 600)
        lay = QVBoxLayout(self)
        
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Buscar acción...")
        self.search_edit.textChanged.connect(lambda t: filter_table_widget(self.table, t))
        lay.addWidget(self.search_edit)

        self.table = QTableWidget(0, 4)
        self.table.setHorizontalHeaderLabels(["Fecha", "Usuario", "Acción", "Detalle"])
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Interactive)
        header.setStretchLastSection(True)
        self.table.setColumnWidth(0, 150)
        self.table.setColumnWidth(1, 150)
        self.table.setColumnWidth(2, 250)
        lay.addWidget(self.table)
        
        btns = QDialogButtonBox(QDialogButtonBox.Close)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)
        for btn in btns.buttons(): set_button_role(btn, 'neutral')
        self.load_rows()

    def load_rows(self):
        self.table.setRowCount(0)
        for row in get_recent_history(500):
            r = self.table.rowCount(); self.table.insertRow(r)
            vals = [row["created_at"], row["username"], row["action"], row["details"]]
            for c, val in enumerate(vals): self.table.setItem(r, c, QTableWidgetItem(str(val)))

class ReceiptTrashDialog(QDialog):
    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.setWindowTitle("Papelera de recibos")
        self.setMinimumSize(1100, 650)
        lay = QVBoxLayout(self)

        info = QLabel("Los recibos se eliminan automáticamente después de 30 días en la papelera.")
        info.setWordWrap(True)
        lay.addWidget(info)

        self.table = QTableWidget(0, 9)
        self.table.setHorizontalHeaderLabels([
            "ID", "Recibo N°", "Paciente", "Fecha seleccionada", "ARS", "Total", "Usuario", "Eliminado el", "Archivo PDF"
        ])
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Interactive)
        header.setStretchLastSection(True)
        self.table.setColumnWidth(0, 60)
        self.table.setColumnWidth(1, 90)
        self.table.setColumnWidth(2, 220)
        self.table.setColumnWidth(3, 120)
        self.table.setColumnWidth(4, 120)
        self.table.setColumnWidth(5, 100)
        self.table.setColumnWidth(6, 120)
        self.table.setColumnWidth(7, 150)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setAlternatingRowColors(True)
        lay.addWidget(self.table)

        row = QHBoxLayout()
        self.btn_restore = QPushButton("♻️ Restaurar")
        self.btn_delete_forever = QPushButton("🗑️ Eliminar permanentemente")
        self.btn_purge = QPushButton("🧹 Vaciar papelera (>30 días)")
        self.btn_close = QPushButton("Cerrar")
        row.addWidget(self.btn_restore)
        row.addWidget(self.btn_delete_forever)
        row.addWidget(self.btn_purge)
        row.addStretch(1)
        row.addWidget(self.btn_close)
        lay.addLayout(row)

        self.btn_restore.clicked.connect(self.restore_selected)
        self.btn_delete_forever.clicked.connect(self.delete_selected_forever)
        self.btn_purge.clicked.connect(self.purge_old)
        self.btn_close.clicked.connect(self.close)

        set_button_role(self.btn_restore, 'success')
        set_button_role(self.btn_delete_forever, 'danger')
        set_button_role(self.btn_purge, 'warning')
        set_button_role(self.btn_close, 'neutral')

        self.load_rows()

    def load_rows(self):
        self.table.setUpdatesEnabled(False)
        self.table.setRowCount(0)
        for row_data in list_deleted_receipts():
            r = self.table.rowCount()
            self.table.insertRow(r)
            pdf_filename = row_data.get("pdf_filename") or ""
            pdf_path = os.path.join(PDFS_DIR, pdf_filename) if pdf_filename else ""
            vals = [
                row_data.get("id", ""), row_data.get("numero", ""), row_data.get("nombre", ""),
                row_data.get("fecha", ""), row_data.get("ars", ""),
                f"${float(row_data.get('total', 0.0)):,.2f}", row_data.get("username", ""),
                row_data.get("deleted_at", ""), pdf_path
            ]
            for c, val in enumerate(vals):
                self.table.setItem(r, c, QTableWidgetItem(str(val)))
        self.table.setUpdatesEnabled(True)

    def restore_selected(self):
        rows = sorted(set(i.row() for i in self.table.selectedIndexes()), reverse=True)
        if not rows:
            FloatingToast("Selecciona al menos un recibo", self, is_error=True).show()
            return
        for r in rows:
            recibo_id = int(self.table.item(r, 0).text())
            restore_recibo(recibo_id)
        self.load_rows()
        FloatingToast(f"♻️ {len(rows)} recibo(s) restaurado(s)", self).show()

    def delete_selected_forever(self):
        rows = sorted(set(i.row() for i in self.table.selectedIndexes()), reverse=True)
        if not rows:
            FloatingToast("Selecciona al menos un recibo", self, is_error=True).show()
            return
        confirm = QMessageBox.question(self, "Confirmar", f"¿Eliminar permanentemente {len(rows)} recibo(s)?", QMessageBox.Yes | QMessageBox.No)
        if confirm != QMessageBox.Yes:
            return
        for r in rows:
            recibo_id = int(self.table.item(r, 0).text())
            permanently_delete_recibo(recibo_id)
        self.load_rows()
        FloatingToast(f"🗑️ {len(rows)} recibo(s) eliminado(s) permanentemente", self).show()

    def purge_old(self):
        confirm = QMessageBox.question(self, "Vaciar papelera", "¿Eliminar permanentemente todos los registros con más de 30 días?", QMessageBox.Yes | QMessageBox.No)
        if confirm != QMessageBox.Yes:
            return
        purged = purge_old_deleted_receipts(30)
        self.load_rows()
        FloatingToast(f"🧹 {len(purged)} registro(s) antiguos eliminados permanentemente", self).show()

class ReceiptHistoryDialog(QDialog):
    def __init__(self, main_window, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.setWindowTitle("Historial de recibos")
        self.setMinimumSize(1100, 650)
        lay = QVBoxLayout(self)

        search_row = QHBoxLayout()
        self.filter_combo = QComboBox()
        self.filter_combo.addItems(["Mostrar Todos", "Solo Flujo Principal (Normal)", "Solo Historial Alterno (Atrasados)"])
        self.filter_combo.currentIndexChanged.connect(self.load_rows)
        
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Buscar recibo (por paciente, número, ARS)...")
        self.search_edit.textChanged.connect(lambda t: filter_table_widget(self.table, t))
        
        search_row.addWidget(self.filter_combo)
        search_row.addWidget(self.search_edit, 1)
        lay.addLayout(search_row)

        self.table = QTableWidget(0, 10)
        self.table.setHorizontalHeaderLabels([
            "ID", "Recibo N°", "Paciente", "Fecha seleccionada", "Tipo", "Fecha generado", "ARS", "Total", "Usuario", "Archivo PDF"
        ])
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Interactive)
        header.setStretchLastSection(True)
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 80)
        self.table.setColumnWidth(2, 220)
        self.table.setColumnWidth(3, 100)
        self.table.setColumnWidth(4, 120)
        self.table.setColumnWidth(5, 150)
        self.table.setColumnWidth(6, 150)
        self.table.setColumnWidth(7, 100)
        self.table.setColumnWidth(8, 120)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setAlternatingRowColors(True)
        self.table.itemDoubleClicked.connect(self.open_selected_receipt)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_context_menu)
        lay.addWidget(self.table)

        self._receipt_offset = 0
        self._receipt_limit = 500
        self._receipt_batch = 200
        self._receipt_total = 0
        self._loading_more = False

        self.lbl_count = QLabel("")
        lay.addWidget(self.lbl_count)

        row = QHBoxLayout()
        self.btn_open_receipt = QPushButton("📄 Abrir PDF")
        self.btn_edit_receipt = QPushButton("✏️ Editar Recibo")
        self.btn_delete_receipt = QPushButton("❌ Eliminar Recibo")
        self.btn_delete_selected = QPushButton("🗑️ Eliminar seleccionados")
        self.btn_trash = QPushButton("🗑️ Papelera")
        self.btn_close = QPushButton("Cerrar")

        row.addWidget(self.btn_open_receipt)
        row.addWidget(self.btn_edit_receipt)
        row.addWidget(self.btn_delete_receipt)
        row.addWidget(self.btn_delete_selected)
        row.addWidget(self.btn_trash)
        row.addStretch(1)
        row.addWidget(self.btn_close)
        lay.addLayout(row)

        self.btn_open_receipt.clicked.connect(self.open_selected_receipt)
        self.btn_edit_receipt.clicked.connect(self.edit_selected_receipt)
        self.btn_delete_receipt.clicked.connect(self.delete_selected_receipt)
        self.btn_delete_selected.clicked.connect(self.delete_selected_receipts)
        self.btn_trash.clicked.connect(self.main_window.open_trash_dialog)
        self.btn_close.clicked.connect(self.close)

        set_button_role(self.btn_open_receipt, 'report')
        set_button_role(self.btn_edit_receipt, 'info')
        set_button_role(self.btn_delete_receipt, 'danger')
        set_button_role(self.btn_delete_selected, 'danger')
        set_button_role(self.btn_trash, 'warning')
        set_button_role(self.btn_close, 'neutral')

        self.load_rows()

    def show_context_menu(self, pos):
        row = self.table.rowAt(pos.y())
        if row < 0: return
        self.table.selectRow(row)
        
        menu = QMenu(self)
        a_open = QAction("📄 Abrir PDF", self)
        a_edit = QAction("✏️ Editar Recibo", self)
        
        a_open.triggered.connect(self.open_selected_receipt)
        a_edit.triggered.connect(self.edit_selected_receipt)
        
        menu.addAction(a_open)
        menu.addAction(a_edit)
        
        if user_can_delete_receipts(self.main_window.current_user):
            a_del = QAction("❌ Eliminar Recibo", self)
            a_del.triggered.connect(self.delete_selected_receipt)
            a_del_sel = QAction("🗑️ Eliminar seleccionados", self)
            a_del_sel.triggered.connect(self.delete_selected_receipts)
            menu.addSeparator()
            menu.addAction(a_del)
            menu.addAction(a_del_sel)
            
        menu.exec(self.table.viewport().mapToGlobal(pos))

    def delete_selected_receipts(self):
        if not user_can_delete_receipts(self.main_window.current_user):
            FloatingToast("Solo el administrador o facturación de auditoría puede eliminar recibos", self, is_error=True).show()
            return
        rows = sorted(set(idx.row() for idx in self.table.selectionModel().selectedRows()))
        if not rows:
            FloatingToast("Selecciona al menos un recibo para eliminar", self, is_error=True).show()
            return
        count = len(rows)
        if QMessageBox.question(self, "Eliminar Recibos", f"¿Estás seguro de que deseas ELIMINAR DEFINITIVAMENTE {count} recibo(s)?\n\nEsta acción no se puede deshacer.", QMessageBox.Yes | QMessageBox.No) != QMessageBox.Yes:
            return
        try:
            for row in reversed(rows):
                recibo_id = int(self.table.item(row, 0).text())
                delete_recibo(recibo_id)
                log_action(self.main_window.current_user["username"], "Eliminar recibo múltiple", f"Recibo ID {recibo_id} eliminado")
            FloatingToast(f"✅ {count} recibo(s) eliminado(s)", self).show()
            self.load_rows()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudieron eliminar: {e}")

    def load_rows(self, reset: bool = True):
        if reset:
            self.table.setRowCount(0)
            self._receipt_offset = 0
            self._receipt_total = 0
        filter_idx = self.filter_combo.currentIndex()
        batch = self._receipt_limit if reset else self._receipt_batch
        rows = list_receipts_history(limit=batch, offset=self._receipt_offset)
        if not rows and reset:
            self.lbl_count.setText("Sin recibos")
            return
        for row_data in rows:
            is_backdated = row_data.get("is_backdated", 0)
            if filter_idx == 1 and is_backdated: continue
            if filter_idx == 2 and not is_backdated: continue
            r = self.table.rowCount()
            self.table.insertRow(r)
            pdf_filename = row_data.get("pdf_filename") or ""
            pdf_path = os.path.join(PDFS_DIR, pdf_filename) if pdf_filename else ""
            tipo = "⏱️ Alterno" if is_backdated else "✅ Principal"
            vals = [
                row_data.get("id", ""), row_data.get("numero", ""), row_data.get("nombre", ""),
                row_data.get("fecha", ""), tipo, row_data.get("created_at", ""), row_data.get("ars", ""),
                f"${float(row_data.get('total', 0.0)):,.2f}", row_data.get("username", ""), pdf_path
            ]
            for c, val in enumerate(vals):
                self.table.setItem(r, c, QTableWidgetItem(str(val)))
        self._receipt_offset += len(rows)
        self._receipt_total = self.table.rowCount()
        self.lbl_count.setText(f"Mostrando {self._receipt_total} recibos")
        if self._receipt_total > 0:
            self.table.verticalScrollBar().setValue(0)
        self.table.verticalScrollBar().valueChanged.connect(self._on_receipt_scroll)

    def _on_receipt_scroll(self, value):
        if self._loading_more:
            return
        bar = self.table.verticalScrollBar()
        if value >= bar.maximum() - 2:
            self._loading_more = True
            self._load_more_receipts()

    def _load_more_receipts(self):
        rows = list_receipts_history(limit=self._receipt_batch, offset=self._receipt_offset)
        if not rows:
            self.lbl_count.setText(f"Mostrando {self._receipt_total} recibos (fin de la lista)")
            self._loading_more = False
            return
        filter_idx = self.filter_combo.currentIndex()
        added = 0
        for row_data in rows:
            is_backdated = row_data.get("is_backdated", 0)
            if filter_idx == 1 and is_backdated: continue
            if filter_idx == 2 and not is_backdated: continue
            r = self.table.rowCount()
            self.table.insertRow(r)
            pdf_filename = row_data.get("pdf_filename") or ""
            pdf_path = os.path.join(PDFS_DIR, pdf_filename) if pdf_filename else ""
            tipo = "⏱️ Alterno" if is_backdated else "✅ Principal"
            vals = [
                row_data.get("id", ""), row_data.get("numero", ""), row_data.get("nombre", ""),
                row_data.get("fecha", ""), tipo, row_data.get("created_at", ""), row_data.get("ars", ""),
                f"${float(row_data.get('total', 0.0)):,.2f}", row_data.get("username", ""), pdf_path
            ]
            for c, val in enumerate(vals):
                self.table.setItem(r, c, QTableWidgetItem(str(val)))
            added += 1
        self._receipt_offset += len(rows)
        self._receipt_total = self.table.rowCount()
        self.lbl_count.setText(f"Mostrando {self._receipt_total} recibos")
        self._loading_more = False

    def open_selected_receipt(self):
        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un recibo", self, is_error=True).show()
            return

        stored_path = self.table.item(row, 9).text().strip()
        filename = os.path.basename(stored_path)

        if not filename:
            QMessageBox.warning(self, "Recibos", "Este recibo no tiene un archivo PDF asociado.")
            return

        pdf_path = stored_path if os.path.exists(stored_path) else stable_storage_path(PDFS_DIR, filename)

        if not os.path.exists(pdf_path):
            with db_connect() as con:
                cur = con.execute("SELECT file_data FROM pdf_storage WHERE filename=%s", (filename,))
                row_data = cur.fetchone()
                if row_data:
                    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
                    with open(pdf_path, 'wb') as f:
                        f.write(bytes(row_data['file_data']))

        if not os.path.exists(pdf_path):
            QMessageBox.warning(self, "Recibos", "No se encontró el PDF correspondiente a este recibo en la base de datos.")
            return

        if not open_file_path(pdf_path):
            QMessageBox.warning(self, "Recibos", "No fue posible abrir el recibo.")

    def edit_selected_receipt(self):
        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un recibo para editar", self, is_error=True).show()
            return
        recibo_id = int(self.table.item(row, 0).text())
        self.main_window.load_recibo_for_editing(recibo_id)
        self.accept()

    def delete_selected_receipt(self):
        if not user_can_delete_receipts(self.main_window.current_user):
            FloatingToast("Solo el administrador o facturación de auditoría puede eliminar recibos", self, is_error=True).show()
            return

        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un recibo para eliminar", self, is_error=True).show()
            return
            
        recibo_id = int(self.table.item(row, 0).text())
        numero = self.table.item(row, 1).text()
        
        if QMessageBox.question(self, "Eliminar Recibo", f"¿Estás seguro de que deseas ELIMINAR DEFINITIVAMENTE el recibo N° {numero}?\n\nEsta acción no se puede deshacer.", QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
            try:
                delete_recibo(recibo_id)
                log_action(self.main_window.current_user["username"], "Eliminar recibo", f"Recibo N° {numero} eliminado permanentemente")
                FloatingToast("✅ Recibo eliminado", self).show()
                self.load_rows()
            except Exception as e:
                QMessageBox.critical(self, "Error", f"No se pudo eliminar: {e}")

class LegacyReportsDialog(QDialog):
    def __init__(self, current_user: dict, parent=None):
        super().__init__(parent)
        self.current_user = current_user
        self.setWindowTitle("Reportes y Estadísticas")
        self.setMinimumSize(1100, 750)
        lay = QVBoxLayout(self)

        config_box = QGroupBox("Generar reporte")
        cfg = QGridLayout(config_box) 
        
        self.report_type = QComboBox()
        self.report_type.addItems(["Diario", "Semanal", "Mensual", "Anual"])
        self.report_type.addItem("Período personalizado")
        
        self.date_from = QDateEdit(); self.date_from.setCalendarPopup(True); self.date_from.setDisplayFormat("yyyy-MM-dd"); self.date_from.setDate(QDate.currentDate())
        self.date_to = QDateEdit(); self.date_to.setCalendarPopup(True); self.date_to.setDisplayFormat("yyyy-MM-dd"); self.date_to.setDate(QDate.currentDate())
        self.date_to.dateChanged.connect(self._on_date_to_changed)
        
        self.ars_filter = QComboBox()
        self.ars_filter.addItem("Todas las ARS")
        self.ars_filter.addItems(ars_list())
        
        self.user_filter = QComboBox()
        self.user_filter.addItem("Todos los Usuarios")
        self.user_filter.addItems(list_usernames())
        
        self.btn_generate = QPushButton("Generar reporte")
        self.btn_generate.setMinimumHeight(40)
        
        self.date_from.dateChanged.connect(self._on_date_from_changed)
        
        cfg.addWidget(QLabel("Tipo de Reporte:"), 0, 0)
        cfg.addWidget(self.report_type, 0, 1)
        
        cfg.addWidget(QLabel("Filtro ARS:"), 1, 0)
        cfg.addWidget(self.ars_filter, 1, 1)
        cfg.addWidget(QLabel("Filtro Usuario:"), 1, 2)
        cfg.addWidget(self.user_filter, 1, 3)
        
        cfg.addWidget(QLabel("Desde la fecha:"), 2, 0)
        cfg.addWidget(self.date_from, 2, 1)
        cfg.addWidget(QLabel("Hasta la fecha:"), 2, 2)
        cfg.addWidget(self.date_to, 2, 3)
        cfg.addWidget(self.date_to, 2, 3)
        
        cfg.addWidget(self.btn_generate, 3, 0, 1, 4)
        lay.addWidget(config_box)

        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("🔍 Buscar reporte...")
        self.search_edit.textChanged.connect(lambda t: filter_table_widget(self.table, t))
        lay.addWidget(self.search_edit)

        self.table = QTableWidget(0, 9)
        self.table.setHorizontalHeaderLabels(["Tipo", "Desde", "Hasta", "Generado en la fecha", "Por", "Archivo", "JSON", "ID", "Tabla"])
        self.table.setColumnHidden(6, True)
        self.table.setColumnHidden(7, True)
        self.table.setColumnHidden(8, True)
        
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Interactive)
        header.setStretchLastSection(True)
        self.table.setColumnWidth(0, 180)
        self.table.setColumnWidth(1, 100)
        self.table.setColumnWidth(2, 100)
        self.table.setColumnWidth(3, 150)
        self.table.setColumnWidth(4, 150)
        
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        
        self.table.itemDoubleClicked.connect(self.open_selected)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_context_menu)

        lay.addWidget(self.table, 1)

        btnrow = QHBoxLayout()
        self.btn_open = QPushButton("Abrir archivo PDF")
        self.btn_delete = QPushButton("Eliminar reporte")
        self.btn_close = QPushButton("Cerrar")
        btnrow.addWidget(self.btn_open)
        btnrow.addWidget(self.btn_delete)
        btnrow.addStretch(1)
        btnrow.addWidget(self.btn_close)
        lay.addLayout(btnrow)

        self.report_type.currentTextChanged.connect(self.sync_dates_for_type)
        self.btn_generate.clicked.connect(self.generate_selected_report)
        self.btn_open.clicked.connect(self.open_selected)
        self.btn_delete.clicked.connect(self.delete_selected_report)
        self.btn_close.clicked.connect(self.close)

        set_button_role(self.btn_generate, 'report')
        set_button_role(self.btn_open, 'report')
        set_button_role(self.btn_delete, 'danger')
        set_button_role(self.btn_close, 'neutral')

        self.load_rows()

    def _on_date_from_changed(self, new_date):
        if self.report_type.currentText() == "Diario":
            self.date_to.setDate(new_date)
        # Auto-switch to "Período personalizado" when date is manually changed
        if self.report_type.currentText() != "Período personalizado":
            self.report_type.setCurrentText("Período personalizado")

    def _on_date_to_changed(self, new_date):
        # Auto-switch to "Período personalizado" when date is manually changed
        if self.report_type.currentText() != "Período personalizado":
            self.report_type.setCurrentText("Período personalizado")

    def sync_dates_for_type(self, report_type: str):
        today = QDate.currentDate()
        self.date_from.blockSignals(True)
        self.date_to.blockSignals(True)
        self.date_from.setEnabled(True)
        
        if report_type == "Diario":
            self.date_from.setDate(today)
            self.date_to.setDate(today)
            self.date_to.setEnabled(False) 
        elif report_type == "Semanal":
            start = today.addDays(-(today.dayOfWeek() - 1))
            end = start.addDays(6)
            self.date_from.setDate(start)
            self.date_to.setDate(end)
            self.date_to.setEnabled(True)
        elif report_type == "Mensual":
            start = QDate(today.year(), today.month(), 1)
            end = start.addMonths(1).addDays(-1)
            self.date_from.setDate(start)
            self.date_to.setDate(end)
            self.date_to.setEnabled(True)
        elif report_type == "Anual":
            self.date_from.setDate(QDate(today.year(), 1, 1))
            self.date_to.setDate(QDate(today.year(), 12, 31))
            self.date_to.setEnabled(True)
        else:
            self.date_to.setEnabled(True)
            
        self.date_from.blockSignals(False)
        self.date_to.blockSignals(False)

    def generate_selected_report(self):
        report_type = self.report_type.currentText()
        ars_filter = self.ars_filter.currentText()
        user_filter = self.user_filter.currentText()

        start_date = self.date_from.date().toString("yyyy-MM-dd")
        end_date = self.date_to.date().toString("yyyy-MM-dd")
        if end_date < start_date:
            FloatingToast("Fechas inválidas", self, is_error=True).show()
            return
        try:
            if report_type == "Diario": path = generate_daily_report_pdf(start_date, self.current_user['username'], None, ars_filter, user_filter)
            else: path = generate_period_report_pdf(report_type, start_date, end_date, self.current_user['username'], None, ars_filter, user_filter)
            if not path:
                QMessageBox.information(self, "Reportes", "No hay datos para ese período.")
                return
            self.load_rows()
            FloatingToast("✅ Reporte generado", self).show()
            if not open_file_path(path):
                QMessageBox.warning(self, "Reportes", f"El reporte se generó, pero no se pudo abrir automáticamente:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Reportes", f"No se pudo generar el reporte:\n{e}")

    def load_rows(self):
        self.table.setRowCount(0)
        for row in list_report_history():
            r = self.table.rowCount()
            self.table.insertRow(r)

            stored_path = row.get("filepath", "")
            display_path = stable_storage_path(REPORTS_DIR, stored_path) if stored_path else ""

            vals = [
                row.get("report_type", "Diario"), 
                row.get("start_date", ""), 
                row.get("end_date", ""), 
                row.get("generated_at", ""), 
                row.get("generated_by", ""), 
                display_path,
                row.get("totals_json", "{}"),
                row.get("record_id", ""),
                row.get("source_table", "")
            ]
            for c, val in enumerate(vals):
                self.table.setItem(r, c, QTableWidgetItem(str(val)))

        if self.table.rowCount() > 0:
            self.table.selectRow(0)

    def show_context_menu(self, pos):
        row = self.table.rowAt(pos.y())
        if row < 0:
            return
        self.table.selectRow(row)

        menu = QMenu(self)
        a_open = QAction("📄 Abrir PDF", self)
        a_delete = QAction("❌ Eliminar reporte", self)

        a_open.triggered.connect(self.open_selected)
        a_delete.triggered.connect(self.delete_selected_report)

        menu.addAction(a_open)
        menu.addSeparator()
        menu.addAction(a_delete)
        menu.exec(self.table.viewport().mapToGlobal(pos))

    def open_selected(self):
        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un reporte", self, is_error=True).show()
            return

        stored_path = self.table.item(row, 5).text().strip()
        filename = os.path.basename(stored_path)

        if not filename:
            QMessageBox.warning(self, "Reportes", "Este reporte no tiene un archivo PDF asociado.")
            return

        path = stored_path if os.path.exists(stored_path) else stable_storage_path(REPORTS_DIR, filename)

        if not os.path.exists(path):
            with db_connect() as con:
                cur = con.execute("SELECT file_data FROM pdf_storage WHERE filename=%s", (filename,))
                row_data = cur.fetchone()
                if row_data:
                    os.makedirs(os.path.dirname(path), exist_ok=True)
                    with open(path, 'wb') as f:
                        f.write(bytes(row_data['file_data']))

        if not os.path.exists(path):
            QMessageBox.warning(self, "Reportes", "No se encontró el PDF de este reporte en la base de datos.")
            return

        if not open_file_path(path):
            QMessageBox.warning(self, "Reportes", "No fue posible abrir el archivo.")

    def delete_selected_report(self):
        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un reporte para eliminar", self, is_error=True).show()
            return

        record_item = self.table.item(row, 7)
        source_item = self.table.item(row, 8)
        filepath_item = self.table.item(row, 5)

        record_id = int(record_item.text()) if record_item and record_item.text().strip() else 0
        source_table = source_item.text().strip() if source_item else ""
        filepath = filepath_item.text().strip() if filepath_item else ""
        report_type = self.table.item(row, 0).text() if self.table.item(row, 0) else "reporte"

        if not record_id or not source_table:
            QMessageBox.warning(self, "Reportes", "No se pudo identificar el reporte seleccionado.")
            return

        confirm = QMessageBox.question(
            self,
            "Eliminar reporte",
            f"¿Seguro que deseas eliminar este reporte?\n\n{report_type}\n{os.path.basename(filepath)}\n\nEsta acción quitará el registro del historial y el PDF almacenado.",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm != QMessageBox.Yes:
            return

        try:
            delete_report_record(source_table, record_id, filepath)
            log_action(self.current_user["username"], "Eliminar reporte", f"{report_type} | {os.path.basename(filepath)}")
            FloatingToast("✅ Reporte eliminado", self).show()
            self.load_rows()
        except Exception as e:
            QMessageBox.critical(self, "Reportes", f"No se pudo eliminar el reporte:\n{e}")


class DashboardLoadWorker(QThread):
    """Carga agregaciones del panel sin bloquear el hilo de la interfaz."""

    loaded = Signal(object)
    failed = Signal(str)

    def __init__(self, parameters: tuple, parent=None):
        super().__init__(parent)
        self.parameters = parameters

    def run(self):
        try:
            self.loaded.emit(get_dashboard_statistics(*self.parameters))
        except Exception as exc:
            self.failed.emit(str(exc))


class DashboardExportWorker(QThread):
    """Genera Excel o PDF fuera del hilo visual y conserva el mismo snapshot."""

    exported = Signal(str)
    failed = Signal(str)

    def __init__(self, mode, data, path, username, logo_path, parent=None):
        super().__init__(parent)
        self.mode = mode
        self.data = data
        self.path = path
        self.username = username
        self.logo_path = logo_path

    def run(self):
        try:
            if self.mode == "xlsx":
                result = export_panel_xlsx(
                    self.data, self.path, self.username, self.logo_path
                )
            else:
                view = self.data.get("view", {})
                context = {
                    "mode": "panel",
                    "title": "PANEL DE REPORTES Y GRÁFICOS",
                    "subtitle": (
                        f"Período analizado: {self.data['start_date']} al "
                        f"{self.data['end_date']}"
                    ),
                    "generated_by": self.username,
                    "logo_path": self.logo_path,
                    "data": self.data,
                    "ars_metric": view.get("ars_metric", "total"),
                    "evolution_metric": view.get("evolution_metric", "total"),
                }
                result = ReportHTMLRenderer().render_pdf(
                    context, self.path, landscape=True
                )
            self.exported.emit(result)
        except Exception as exc:
            self.failed.emit(str(exc))


class ReportsDialog(LegacyReportsDialog):
    """Centro moderno de reportes, estadísticas e historial."""

    def __init__(self, current_user: dict, parent=None):
        QDialog.__init__(self, parent)
        self.current_user = current_user
        self.panel_access = is_administrator(current_user)
        self.dashboard_data = None
        self._dashboard_worker = None
        self._dashboard_export_worker = None
        self.setWindowTitle("Centro de Reportes y Estadísticas")
        self.setMinimumSize(860, 680)
        self.resize(1320, 880)

        root = QVBoxLayout(self)
        root.setContentsMargins(16, 14, 16, 14)
        root.setSpacing(12)

        header = QWidget()
        header.setStyleSheet(
            "QWidget { background: #123F83; border-radius: 12px; } "
            "QLabel { color: white; background: transparent; }"
        )
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(22, 15, 22, 15)
        header_text = QVBoxLayout()
        title = QLabel("CENTRO DE REPORTES Y ESTADÍSTICAS")
        title.setStyleSheet("font-size: 18pt; font-weight: 900;")
        subtitle = QLabel("Análisis por fecha real de generación del recibo")
        subtitle.setStyleSheet("font-size: 10.5pt; color: #D9E8FF;")
        header_text.addWidget(title)
        header_text.addWidget(subtitle)
        header_layout.addLayout(header_text, 1)
        close_top = QPushButton("Cerrar")
        close_top.clicked.connect(self.close)
        set_button_role(close_top, "neutral")
        header_layout.addWidget(close_top)
        root.addWidget(header)

        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)
        self.tabs.setStyleSheet(
            "QTabBar::tab { min-width: 190px; padding: 11px 18px; font-weight: 800; } "
            "QTabBar::tab:selected { color: #123F83; border-bottom: 3px solid #123F83; }"
        )
        root.addWidget(self.tabs, 1)
        if self.panel_access:
            self._build_dashboard_tab()
        self._build_generation_tab()
        self._build_history_tab()

        self.load_rows()
        if self.panel_access:
            self.update_dashboard_period(refresh=False)
            self.refresh_dashboard()

    def _ensure_panel_access(self, show_message=True) -> bool:
        allowed = is_administrator(self.current_user)
        if not allowed and show_message:
            QMessageBox.warning(
                self,
                "Acceso restringido",
                "El Panel y gráficos está disponible exclusivamente para administradores.",
            )
        return allowed

    def _build_dashboard_tab(self):
        if not self._ensure_panel_access(show_message=False):
            return
        tab = QWidget()
        outer = QVBoxLayout(tab)
        outer.setContentsMargins(4, 10, 4, 4)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        content = QWidget()
        content.setMinimumWidth(0)
        content.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Preferred)
        layout = QVBoxLayout(content)
        layout.setSpacing(13)

        self._updating_dashboard_dates = False

        self.dashboard_period_box = QGroupBox("1. Seleccione el período")
        self.dashboard_period_layout = QGridLayout(self.dashboard_period_box)
        self.dashboard_period_selector = PeriodSelectorWidget()
        self.dashboard_period_layout.addWidget(self.dashboard_period_selector, 0, 0)
        self.dashboard_period_fields = []

        self.dashboard_main_box = QGroupBox("2. ¿Qué información desea analizar?")
        self.dashboard_main_layout = QGridLayout(self.dashboard_main_box)
        self.dashboard_ars = MultiSelectFilter(
            ars_list(), "Todas las ARS", "ARS", feminine=True
        )
        self.dashboard_user = MultiSelectFilter(
            list_usernames(), "Todos los facturadores", "facturador"
        )
        self.dashboard_coverage = QComboBox()
        self.dashboard_coverage.addItems(["Todas", "Asegurados", "No asegurados"])
        self.dashboard_main_fields = [
            (QLabel("ARS:"), self.dashboard_ars),
            (QLabel("Facturadores:"), self.dashboard_user),
            (QLabel("Cobertura:"), self.dashboard_coverage),
        ]
        self.btn_dashboard_advanced = QToolButton()
        self.btn_dashboard_advanced.setCheckable(True)
        self.btn_dashboard_advanced.setText("▸ Más filtros")
        self.btn_dashboard_advanced.setToolButtonStyle(Qt.ToolButtonTextOnly)
        self.btn_dashboard_advanced.setMaximumWidth(250)
        self.btn_dashboard_advanced.setStyleSheet(
            "QToolButton { color: #123F83; border: none; background: transparent; "
            "font-weight: 800; padding: 6px 2px; text-align: left; } "
            "QToolButton:hover { color: #0B7A5A; }"
        )

        self.dashboard_advanced_box = QGroupBox("Filtros adicionales")
        self.dashboard_advanced_layout = QGridLayout(self.dashboard_advanced_box)
        self.dashboard_category = QComboBox()
        self.dashboard_category.addItem("Todas las categorías")
        self.dashboard_category.addItems(ALL_CATEGORIES + ["Sala de Emergencia"])
        self.dashboard_granularity = QComboBox()
        self.dashboard_granularity.addItems(["Automático", "Diario", "Semanal", "Mensual"])
        self.dashboard_advanced_fields = [
            (QLabel("Categoría de servicio:"), self.dashboard_category),
            (QLabel("Mostrar evolución por:"), self.dashboard_granularity),
        ]
        self.dashboard_advanced_box.setVisible(False)

        self.btn_dashboard_compare = QPushButton("Comparar período actual")
        self.btn_dashboard_compare.setCheckable(True)
        self.btn_dashboard_refresh = QPushButton("Actualizar panel")
        self.btn_dashboard_clear = QPushButton("Restablecer filtros")
        self.btn_export_excel = QPushButton("Exportar Excel")
        self.btn_export_panel_pdf = QPushButton("Exportar panel PDF")
        self.btn_dashboard_export = QToolButton()
        self.btn_dashboard_export.setText("Exportar")
        self.btn_dashboard_export.setPopupMode(QToolButton.InstantPopup)
        export_menu = QMenu(self.btn_dashboard_export)
        export_excel_action = export_menu.addAction("Exportar a Excel")
        export_pdf_action = export_menu.addAction("Exportar panel a PDF")
        export_excel_action.triggered.connect(self.export_dashboard_excel)
        export_pdf_action.triggered.connect(self.export_dashboard_pdf)
        self.btn_dashboard_export.setMenu(export_menu)
        self.btn_dashboard_export.setStyleSheet(
            "QToolButton { background: #2E7D32; color: white; border: none; border-radius: 6px; "
            "padding: 8px 14px; font-weight: 800; } "
            "QToolButton:hover { background: #1F6424; } "
            "QToolButton:disabled { background: #B0BEC5; color: #757575; }"
        )
        self.btn_print_comparison = QPushButton("Imprimir comparación")
        self.btn_print_comparison.setEnabled(False)
        self.dashboard_compare_summary = QLabel()
        self.dashboard_compare_summary.setWordWrap(True)
        self.dashboard_compare_summary.setStyleSheet(
            "color: #36516E; background: transparent; border: none; padding: 4px 2px;"
        )
        self.dashboard_compare_summary.setVisible(False)
        self.dashboard_action_query_label = QLabel("CONSULTA")
        self.dashboard_action_docs_label = QLabel("DOCUMENTOS")
        for label in (self.dashboard_action_query_label, self.dashboard_action_docs_label):
            label.setStyleSheet(
                "color: #66788A; background: transparent; border: none; "
                "font-size: 8pt; font-weight: 900; letter-spacing: .5px; padding: 2px;"
            )
        set_button_role(self.btn_dashboard_refresh, "report")
        set_button_role(self.btn_dashboard_compare, "neutral")
        set_button_role(self.btn_dashboard_clear, "neutral")
        set_button_role(self.btn_export_excel, "success")
        set_button_role(self.btn_export_panel_pdf, "report")
        set_button_role(self.btn_print_comparison, "report")
        self.btn_dashboard_clear.setMaximumWidth(180)
        self.btn_dashboard_clear.setStyleSheet(
            "QPushButton { background: #FFF3E0; color: #8A4B08; border: 1px solid #E7B36D; "
            "border-radius: 6px; padding: 8px 12px; font-weight: 800; } "
            "QPushButton:hover { background: #FFE2B8; border-color: #D99538; } "
            "QPushButton:pressed { background: #F8CE92; }"
        )
        self.btn_dashboard_clear.setCursor(Qt.PointingHandCursor)
        for button in (self.btn_dashboard_refresh, self.btn_dashboard_compare,
                       self.btn_export_excel, self.btn_export_panel_pdf,
                       self.btn_print_comparison):
            button.setMinimumHeight(38)
        self.btn_dashboard_export.setMinimumHeight(38)

        self.dashboard_filter_fields = self.dashboard_main_fields + self.dashboard_advanced_fields
        self.dashboard_filter_buttons = [
            self.btn_dashboard_refresh, self.btn_dashboard_compare, self.btn_dashboard_clear,
            self.btn_print_comparison, self.btn_dashboard_export,
        ]
        self.dashboard_actions_widget = QWidget()
        self.dashboard_actions_widget.setStyleSheet(
            "QWidget { background: #FFFFFF; border: 1px solid #D7E2F0; border-radius: 9px; } "
            "QPushButton { border-radius: 6px; }"
        )
        dashboard_actions = QGridLayout(self.dashboard_actions_widget)
        self.dashboard_actions_layout = dashboard_actions
        dashboard_actions.setContentsMargins(10, 8, 10, 9)
        dashboard_actions.setSpacing(8)
        for _label, field in self.dashboard_filter_fields:
            field.setMinimumWidth(120)
            field.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            if isinstance(field, QComboBox):
                field.setSizeAdjustPolicy(QComboBox.AdjustToMinimumContentsLengthWithIcon)
                field.setMinimumContentsLength(12)
        self._arrange_dashboard_filters(False)
        layout.addWidget(self.dashboard_period_box)
        layout.addWidget(self.dashboard_main_box)
        layout.addWidget(self.dashboard_advanced_box)

        self.dashboard_filter_summary = QLabel()
        self.dashboard_filter_summary.setWordWrap(True)
        self.dashboard_filter_summary.setStyleSheet(
            "background: #EDF4FD; color: #24415F; border: 1px solid #C4D7EF; "
            "border-radius: 9px; padding: 9px 13px; font-weight: 700;"
        )
        layout.addWidget(self.dashboard_filter_summary)
        layout.addWidget(self.dashboard_actions_widget)

        self.dashboard_status = QLabel("Preparando estadísticas...")
        self.dashboard_status.setStyleSheet(
            "color: #52657A; padding: 5px 2px; font-size: 9.5pt;"
        )
        layout.addWidget(self.dashboard_status)

        self.dashboard_results_state = QLabel(
            "Cargando la información del período seleccionado..."
        )
        self.dashboard_results_state.setWordWrap(True)
        self.dashboard_results_state.setStyleSheet(
            "background: #EDF4FD; color: #24415F; border: 1px solid #C4D7EF; "
            "border-radius: 8px; padding: 11px 13px; font-weight: 700;"
        )
        layout.addWidget(self.dashboard_results_state)

        self.dashboard_results_tabs = QTabWidget()
        self.dashboard_results_tabs.setTabBar(NoWheelTabBar(self.dashboard_results_tabs))
        self.dashboard_results_tabs.setObjectName("DashboardResultsTabs")
        self.dashboard_results_tabs.setDocumentMode(True)
        self.dashboard_results_tabs.setMinimumHeight(560)
        self.dashboard_results_tabs.setMinimumWidth(0)
        self.dashboard_results_tabs.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Expanding)
        self.dashboard_summary_page = QWidget()
        self.dashboard_distribution_page = QWidget()
        self.dashboard_evolution_page = QWidget()
        self.dashboard_comparison_page = QWidget()
        self.dashboard_detail_page = QWidget()
        for page in (
            self.dashboard_summary_page, self.dashboard_distribution_page,
            self.dashboard_evolution_page, self.dashboard_comparison_page,
            self.dashboard_detail_page,
        ):
            page.setMinimumWidth(0)
            page.setSizePolicy(QSizePolicy.Ignored, QSizePolicy.Expanding)
        self.dashboard_summary_layout = QVBoxLayout(self.dashboard_summary_page)
        self.dashboard_distribution_layout = QVBoxLayout(self.dashboard_distribution_page)
        self.dashboard_evolution_layout = QVBoxLayout(self.dashboard_evolution_page)
        self.dashboard_comparison_page_layout = QVBoxLayout(self.dashboard_comparison_page)
        self.dashboard_detail_layout = QVBoxLayout(self.dashboard_detail_page)
        self.dashboard_summary_tab_index = self.dashboard_results_tabs.addTab(
            self.dashboard_summary_page, "Resumen"
        )
        self.dashboard_distribution_tab_index = self.dashboard_results_tabs.addTab(
            self.dashboard_distribution_page, "Distribuciones"
        )
        self.dashboard_evolution_tab_index = self.dashboard_results_tabs.addTab(
            self.dashboard_evolution_page, "Evolución"
        )
        self.dashboard_comparison_tab_index = self.dashboard_results_tabs.addTab(
            self.dashboard_comparison_page, "Comparación"
        )
        self.dashboard_detail_tab_index = self.dashboard_results_tabs.addTab(
            self.dashboard_detail_page, "Detalle"
        )
        self.dashboard_results_tabs.setTabVisible(self.dashboard_comparison_tab_index, False)
        layout.addWidget(self.dashboard_results_tabs)

        kpis = QGridLayout()
        self.dashboard_kpi_layout = kpis
        self.kpi_receipts = self._create_kpi_card(kpis, 0, "TOTAL DE RECIBOS", "0", "#123F83")
        self.kpi_total = self._create_kpi_card(kpis, 1, "TOTAL EMITIDO", "RD$ 0.00", "#0B7A5A")
        self.kpi_average = self._create_kpi_card(kpis, 2, "PROMEDIO POR RECIBO", "RD$ 0.00", "#D16619")
        self.kpi_room = self._create_kpi_card(kpis, 3, "SALA DE EMERGENCIA", "RD$ 0.00", "#6B35C8")
        self.dashboard_kpi_cards = [
            self.kpi_receipts._card_widget,
            self.kpi_total._card_widget,
            self.kpi_average._card_widget,
            self.kpi_room._card_widget,
        ]
        self.dashboard_summary_layout.addLayout(kpis)

        self.dashboard_comparison_panel = QGroupBox("Comparación con período anterior")
        comparison_panel_layout = QVBoxLayout(self.dashboard_comparison_panel)
        self.dashboard_comparison_periods = QLabel()
        self.dashboard_comparison_periods.setWordWrap(True)
        self.dashboard_comparison_periods.setStyleSheet(
            "background: #EDF4FD; color: #24415F; border: 1px solid #C4D7EF; "
            "border-radius: 7px; padding: 9px 12px; font-weight: 700;"
        )
        comparison_panel_layout.addWidget(self.dashboard_comparison_periods)

        self.dashboard_comparison_table = QTableWidget(0, 5)
        self.dashboard_comparison_table.setHorizontalHeaderLabels(
            ["Indicador", "Actual", "Anterior", "Diferencia", "Variación"]
        )
        self.dashboard_comparison_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        for column in range(1, 5):
            self.dashboard_comparison_table.horizontalHeader().setSectionResizeMode(
                column, QHeaderView.ResizeToContents
            )
        self.dashboard_comparison_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.dashboard_comparison_table.setSelectionMode(QAbstractItemView.NoSelection)
        self.dashboard_comparison_table.setAlternatingRowColors(True)
        self.dashboard_comparison_table.setMinimumHeight(190)
        comparison_panel_layout.addWidget(self.dashboard_comparison_table)

        categories_comparison_box = QGroupBox("Comparación por categorías")
        self.dashboard_categories_comparison_box = categories_comparison_box
        categories_comparison_layout = QVBoxLayout(categories_comparison_box)
        self.dashboard_categories_comparison = QTableWidget(0, 2)
        self.dashboard_categories_comparison.setHorizontalHeaderLabels(["Categoría", "Variación"])
        self.dashboard_categories_comparison.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.dashboard_categories_comparison.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.dashboard_categories_comparison.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.dashboard_categories_comparison.setSelectionMode(QAbstractItemView.NoSelection)
        self.dashboard_categories_comparison.setAlternatingRowColors(True)
        self.dashboard_categories_comparison.setMinimumHeight(210)
        categories_comparison_layout.addWidget(self.dashboard_categories_comparison)
        comparison_panel_layout.addWidget(categories_comparison_box)
        self.dashboard_comparison_panel.setVisible(False)
        self.dashboard_comparison_page_layout.addWidget(self.dashboard_comparison_panel)
        self.dashboard_comparison_page_layout.addStretch(1)

        charts = QGridLayout()
        charts.setHorizontalSpacing(13)
        charts.setVerticalSpacing(13)
        self.dashboard_charts_layout = charts

        donut_box = QGroupBox("Distribución porcentual por categoría")
        self.donut_box = donut_box
        donut_layout = QVBoxLayout(donut_box)
        donut_hint = QLabel("Participación de cada categoría sobre el total clasificado, incluida Sala de Emergencia.")
        donut_hint.setWordWrap(True)
        donut_hint.setStyleSheet("color: #66788A; font-size: 8.5pt;")
        donut_layout.addWidget(donut_hint)
        self.donut_chart = DonutChart()
        donut_layout.addWidget(self.donut_chart)

        category_box = QGroupBox("Distribución por categorías - Total emitido")
        self.category_box = category_box
        category_layout = QVBoxLayout(category_box)
        category_hint = QLabel("Monto acumulado por categoría, expresado en pesos dominicanos (RD$).")
        category_hint.setStyleSheet("color: #66788A; font-size: 8.5pt;")
        category_hint.setWordWrap(True)
        category_layout.addWidget(category_hint)
        self.category_bar_chart = ModernBarChart("#0B7A5A")
        category_layout.addWidget(self.category_bar_chart)

        comparison_box = QGroupBox("Distribución porcentual por ARS")
        self.comparison_box = comparison_box
        comparison_layout = QVBoxLayout(comparison_box)
        comparison_controls = QHBoxLayout()
        self.dashboard_comparison_metric = QComboBox()
        self.dashboard_comparison_metric.addItems(
            ["Total emitido", "Cantidad de recibos"]
        )
        comparison_controls.addWidget(QLabel("Porcentaje según:"))
        comparison_controls.addWidget(self.dashboard_comparison_metric, 1)
        comparison_layout.addLayout(comparison_controls)
        self.comparison_chart = HorizontalBarChart("#174A96")
        comparison_layout.addWidget(self.comparison_chart)

        line_box = QGroupBox("Evolución diaria")
        self.line_box = line_box
        line_layout = QVBoxLayout(line_box)
        line_hint = QLabel("Cada punto representa el intervalo seleccionado; coloca el cursor para consultar el valor exacto.")
        line_hint.setWordWrap(True)
        line_hint.setStyleSheet("color: #66788A; font-size: 8.5pt;")
        line_layout.addWidget(line_hint)
        evolution_controls = QHBoxLayout()
        self.dashboard_evolution_metric = QComboBox()
        self.dashboard_evolution_metric.addItems(
            ["Total emitido", "Total de recibos", "Promedio por recibo"]
        )
        evolution_controls.addWidget(QLabel("Evolución de:"))
        evolution_controls.addWidget(self.dashboard_evolution_metric, 1)
        line_layout.addLayout(evolution_controls)
        self.line_chart = ModernLineChart()
        line_layout.addWidget(self.line_chart)

        coverage_box = QGroupBox("Distribución por tipo de cobertura")
        self.coverage_box = coverage_box
        coverage_layout = QVBoxLayout(coverage_box)
        coverage_hint = QLabel("Porcentaje de recibos asegurados y no asegurados dentro del período.")
        coverage_hint.setWordWrap(True)
        coverage_hint.setStyleSheet("color: #66788A; font-size: 8.5pt;")
        coverage_layout.addWidget(coverage_hint)
        self.coverage_chart = HorizontalBarChart("#0B7A5A")
        coverage_layout.addWidget(self.coverage_chart)

        self._dashboard_chart_widgets = [donut_box, category_box, comparison_box, coverage_box]
        charts.addWidget(donut_box, 0, 0)
        charts.addWidget(category_box, 0, 1)
        charts.addWidget(comparison_box, 1, 0, 1, 2)
        charts.addWidget(coverage_box, 2, 0, 1, 2)
        charts.setColumnStretch(0, 1)
        charts.setColumnStretch(1, 1)
        self.dashboard_distribution_layout.addLayout(charts)
        self.dashboard_evolution_layout.addWidget(line_box)

        table_box = QGroupBox("Tabla resumen de los datos mostrados")
        self.dashboard_table_box = table_box
        table_layout = QVBoxLayout(table_box)
        self.dashboard_table = QTableWidget(0, 6)
        self.dashboard_table.setHorizontalHeaderLabels(
            ["ARS", "Total emitido", "Recibos", "% del dinero", "% de recibos", "Promedio / recibo"]
        )
        self.dashboard_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.dashboard_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.dashboard_table.setAlternatingRowColors(True)
        self.dashboard_table.setMinimumHeight(220)
        table_layout.addWidget(self.dashboard_table)
        self.dashboard_detail_layout.addWidget(table_box)

        guide_box = QGroupBox("Lectura guiada de los resultados")
        guide_layout = QVBoxLayout(guide_box)
        self.dashboard_guide = QLabel("Selecciona los filtros para analizar el período.")
        self.dashboard_guide.setWordWrap(True)
        self.dashboard_guide.setStyleSheet(
            "font-size: 10.5pt; color: #354C65; background: #F7FAFE; "
            "border-radius: 8px; padding: 13px;"
        )
        guide_layout.addWidget(self.dashboard_guide)
        self.dashboard_summary_layout.addWidget(guide_box)
        self.dashboard_summary_layout.addStretch(1)

        scroll.setWidget(content)
        outer.addWidget(scroll)
        self.tabs.addTab(tab, "Panel y gráficos")

        self.dashboard_period_selector.periodChanged.connect(self.update_dashboard_period)
        for combo in (
            self.dashboard_coverage, self.dashboard_category, self.dashboard_granularity,
        ):
            combo.currentTextChanged.connect(self._mark_dashboard_stale)
        self.dashboard_ars.selectionChanged.connect(self._mark_dashboard_stale)
        self.dashboard_user.selectionChanged.connect(self._mark_dashboard_stale)
        self.btn_dashboard_compare.toggled.connect(self._dashboard_compare_toggled)
        self.btn_dashboard_advanced.toggled.connect(self._toggle_dashboard_advanced)
        self.dashboard_comparison_metric.currentTextChanged.connect(self.render_dashboard)
        self.dashboard_evolution_metric.currentTextChanged.connect(self.render_dashboard)
        self.btn_dashboard_refresh.clicked.connect(self.refresh_dashboard)
        self.btn_dashboard_clear.clicked.connect(self.clear_dashboard_filters)
        self.btn_export_excel.clicked.connect(self.export_dashboard_excel)
        self.btn_export_panel_pdf.clicked.connect(self.export_dashboard_pdf)
        self.btn_print_comparison.clicked.connect(self.print_dashboard_comparison)
        self.update_dashboard_period(refresh=False)

    def _create_kpi_card(self, grid, column, title, initial, color):
        card = QWidget()
        card.setStyleSheet(
            f"QWidget {{ background: white; border: 1px solid #D7E2F0; border-radius: 11px; }} "
            f"QLabel {{ border: none; background: transparent; }}"
        )
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(16, 13, 16, 13)
        title_label = QLabel(title)
        title_label.setStyleSheet(f"color: {color}; font-size: 9pt; font-weight: 900;")
        value_label = QLabel(initial)
        value_label.setWordWrap(True)
        value_label.setStyleSheet("color: #1D2B3D; font-size: 19pt; font-weight: 900;")
        card_layout.addWidget(title_label)
        card_layout.addWidget(value_label)
        grid.addWidget(card, 0, column)
        value_label._card_widget = card
        return value_label

    def _arrange_dashboard_filters(self, compact):
        for button in self.dashboard_filter_buttons:
            self.dashboard_actions_layout.removeWidget(button)
        self.dashboard_actions_layout.removeWidget(self.dashboard_compare_summary)
        self.dashboard_actions_layout.removeWidget(self.dashboard_action_query_label)
        self.dashboard_actions_layout.removeWidget(self.dashboard_action_docs_label)
        for column in range(7):
            self.dashboard_actions_layout.setColumnStretch(column, 0)
        if compact:
            self.dashboard_actions_layout.addWidget(self.dashboard_action_query_label, 0, 0, 1, 2)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_refresh, 1, 0)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_compare, 1, 1)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_clear, 2, 0, 1, 2)
            self.dashboard_actions_layout.addWidget(self.dashboard_action_docs_label, 3, 0, 1, 2)
            self.dashboard_actions_layout.addWidget(self.btn_print_comparison, 4, 0, 1, 2)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_export, 5, 0, 1, 2)
            self.dashboard_actions_layout.addWidget(self.dashboard_compare_summary, 6, 0, 1, 2)
            self.dashboard_actions_layout.setColumnStretch(0, 1)
            self.dashboard_actions_layout.setColumnStretch(1, 1)
        else:
            self.dashboard_actions_layout.addWidget(self.dashboard_action_query_label, 0, 0, 1, 3)
            self.dashboard_actions_layout.addWidget(self.dashboard_action_docs_label, 0, 4, 1, 3)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_refresh, 1, 0)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_compare, 1, 1)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_clear, 1, 2)
            self.dashboard_actions_layout.addWidget(self.btn_print_comparison, 1, 4)
            self.dashboard_actions_layout.addWidget(self.btn_dashboard_export, 1, 5, 1, 2)
            self.dashboard_actions_layout.addWidget(self.dashboard_compare_summary, 2, 0, 1, 7)
            self.dashboard_actions_layout.setColumnStretch(3, 1)

        def clear_fields(target_layout, fields):
            for label, field in fields:
                target_layout.removeWidget(label)
                target_layout.removeWidget(field)

        clear_fields(self.dashboard_main_layout, self.dashboard_main_fields)
        clear_fields(self.dashboard_advanced_layout, self.dashboard_advanced_fields)
        self.dashboard_main_layout.removeWidget(self.btn_dashboard_advanced)

        if compact:
            for row, (label, field) in enumerate(self.dashboard_main_fields):
                self.dashboard_main_layout.addWidget(label, row, 0)
                self.dashboard_main_layout.addWidget(field, row, 1)
            self.dashboard_main_layout.addWidget(
                self.btn_dashboard_advanced, len(self.dashboard_main_fields), 0, 1, 2, Qt.AlignLeft
            )
            self.dashboard_main_layout.setColumnStretch(1, 1)

            for row, (label, field) in enumerate(self.dashboard_advanced_fields):
                self.dashboard_advanced_layout.addWidget(label, row, 0)
                self.dashboard_advanced_layout.addWidget(field, row, 1)
            self.dashboard_advanced_layout.setColumnStretch(1, 1)
        else:
            for column, (label, field) in enumerate(self.dashboard_main_fields):
                self.dashboard_main_layout.addWidget(label, 0, column)
                self.dashboard_main_layout.addWidget(field, 1, column)
                self.dashboard_main_layout.setColumnStretch(column, 1)
            self.dashboard_main_layout.addWidget(
                self.btn_dashboard_advanced, 2, 0, 1, 3, Qt.AlignLeft
            )

            for column, (label, field) in enumerate(self.dashboard_advanced_fields):
                self.dashboard_advanced_layout.addWidget(label, 0, column)
                self.dashboard_advanced_layout.addWidget(field, 1, column)
                self.dashboard_advanced_layout.setColumnStretch(column, 1)

    def _build_generation_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 14, 10, 10)
        intro = QLabel(
            "Genera documentos oficiales usando la fecha real en que cada recibo fue creado. "
            "Los filtros por ARS y usuario abarcan ambos historiales."
        )
        intro.setWordWrap(True)
        intro.setStyleSheet("background: #EDF4FD; padding: 13px; border-radius: 9px; color: #24415F;")
        layout.addWidget(intro)

        config_box = QGroupBox("Configuración del reporte PDF")
        cfg = QGridLayout(config_box)
        self.report_period_selector = PeriodSelectorWidget()
        self.ars_filter = MultiSelectFilter(
            ars_list(), "Todas las ARS", "ARS", feminine=True
        )
        self.user_filter = MultiSelectFilter(
            list_usernames(), "Todos los facturadores", "facturador"
        )
        self.btn_generate = QPushButton("Generar y abrir reporte PDF")
        self.btn_generate.setMinimumHeight(44)
        set_button_role(self.btn_generate, "report")
        report_filter_hint = QLabel(
            "Abre cada filtro para incluir o excluir varias opciones. Los cambios se confirmarán "
            "solo al pulsar Aplicar filtro."
        )
        report_filter_hint.setWordWrap(True)
        report_filter_hint.setStyleSheet(
            "background: #F7FAFE; color: #52657A; border-radius: 7px; padding: 8px 10px;"
        )

        cfg.addWidget(self.report_period_selector, 0, 0, 1, 4)
        cfg.addWidget(QLabel("ARS:"), 1, 0)
        cfg.addWidget(self.ars_filter, 1, 1)
        cfg.addWidget(QLabel("Facturador:"), 1, 2)
        cfg.addWidget(self.user_filter, 1, 3)
        cfg.addWidget(report_filter_hint, 2, 0, 1, 4)
        cfg.addWidget(self.btn_generate, 3, 0, 1, 4)
        layout.addWidget(config_box)

        preview = QGroupBox("Resumen antes de generar")
        preview_layout = QVBoxLayout(preview)
        self.report_preview = QLabel()
        self.report_preview.setWordWrap(True)
        self.report_preview.setStyleSheet("font-size: 11pt; padding: 12px; color: #31475F;")
        preview_layout.addWidget(self.report_preview)
        layout.addWidget(preview)
        layout.addStretch(1)
        self.tabs.addTab(tab, "Generar reporte PDF")

        self.report_period_selector.periodChanged.connect(self.update_report_preview)
        self.ars_filter.selectionChanged.connect(self.update_report_preview)
        self.user_filter.selectionChanged.connect(self.update_report_preview)
        self.btn_generate.clicked.connect(self.generate_selected_report)
        self.update_report_preview()

    def _build_history_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setContentsMargins(10, 14, 10, 10)
        tools = QHBoxLayout()
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("Buscar por tipo, fecha, usuario o archivo...")
        self.history_count = QLabel("0 reportes")
        self.history_count.setStyleSheet("font-weight: 800; color: #123F83;")
        self.btn_history_refresh = QPushButton("Actualizar")
        set_button_role(self.btn_history_refresh, "neutral")
        tools.addWidget(self.search_edit, 1)
        tools.addWidget(self.history_count)
        tools.addWidget(self.btn_history_refresh)
        layout.addLayout(tools)

        self.table = QTableWidget(0, 11)
        self.table.setHorizontalHeaderLabels(
            ["Tipo", "Desde", "Hasta", "Generado el", "Por", "Archivo", "JSON", "ID", "Tabla",
             "Período actual", "Comparado con"]
        )
        for column in (5, 6, 7, 8):
            self.table.setColumnHidden(column, True)
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Interactive)
        header.moveSection(header.visualIndex(9), 1)
        header.moveSection(header.visualIndex(10), 2)
        header.setSectionResizeMode(9, QHeaderView.Stretch)
        header.setSectionResizeMode(10, QHeaderView.Stretch)
        self.table.setColumnWidth(0, 125)
        self.table.setColumnWidth(1, 100)
        self.table.setColumnWidth(2, 100)
        self.table.setColumnWidth(3, 135)
        self.table.setColumnWidth(4, 130)
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setAlternatingRowColors(True)
        layout.addWidget(self.table, 1)

        buttons = QHBoxLayout()
        self.btn_open = QPushButton("Abrir PDF")
        self.btn_history_save = QPushButton("Guardar copia")
        self.btn_history_print = QPushButton("Vista previa e imprimir")
        self.btn_delete = QPushButton("Eliminar reporte")
        set_button_role(self.btn_open, "report")
        set_button_role(self.btn_history_save, "success")
        set_button_role(self.btn_history_print, "report")
        set_button_role(self.btn_delete, "danger")
        buttons.addWidget(self.btn_open)
        buttons.addWidget(self.btn_history_save)
        buttons.addWidget(self.btn_history_print)
        buttons.addWidget(self.btn_delete)
        buttons.addStretch(1)
        layout.addLayout(buttons)
        self.tabs.addTab(tab, "Historial de reportes")

        self.search_edit.textChanged.connect(self.filter_history)
        self.btn_history_refresh.clicked.connect(self.load_rows)
        self.btn_open.clicked.connect(self.open_selected)
        self.btn_history_save.clicked.connect(self.save_selected_report_copy)
        self.btn_history_print.clicked.connect(self.print_selected_report)
        self.btn_delete.clicked.connect(self.delete_selected_report)
        self.table.itemDoubleClicked.connect(self.open_selected)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_context_menu)

    def update_dashboard_period(self, _definition=None, refresh=True):
        self._update_dashboard_compare_summary()
        self._mark_dashboard_stale()
        if refresh:
            self.dashboard_status.setText("Período ajustado. Pulsa Actualizar panel para aplicarlo.")

    def _dashboard_compare_toggled(self, checked):
        self.btn_dashboard_compare.setText(
            "✓ Comparación activada" if checked else "Comparar período actual"
        )
        set_button_role(self.btn_dashboard_compare, "report" if checked else "neutral")
        self._update_dashboard_compare_summary()
        self._mark_dashboard_stale()

    def _update_dashboard_compare_summary(self):
        if not hasattr(self, "dashboard_compare_summary"):
            return
        checked = self.btn_dashboard_compare.isChecked()
        self.dashboard_compare_summary.setVisible(checked)
        if checked:
            definition = self.dashboard_period_selector.definition()
            self.dashboard_compare_summary.setText(
                f"Comparando: {definition['period_label']} "
                f"({self._display_iso_date(definition['start_date'])}–"
                f"{self._display_iso_date(definition['end_date'])}) contra "
                f"{definition['comparison_label']} "
                f"({self._display_iso_date(definition['comparison_start'])}–"
                f"{self._display_iso_date(definition['comparison_end'])})"
            )

    def _toggle_dashboard_advanced(self, checked):
        self.dashboard_advanced_box.setVisible(checked)
        self._update_dashboard_filter_preview()
        self.btn_dashboard_advanced.setText(
            ("▾ Más filtros" if checked else "▸ Más filtros") + self._advanced_filter_count_text()
        )

    def _advanced_filter_count(self):
        return sum((
            self.dashboard_category.currentIndex() > 0,
            self.dashboard_granularity.currentIndex() > 0,
        ))

    def _advanced_filter_count_text(self):
        count = self._advanced_filter_count()
        return f" · {count} activo{'s' if count != 1 else ''}" if count else ""

    @staticmethod
    def _join_filter_values(values):
        values = list(values)
        if len(values) <= 1:
            return "".join(values)
        return ", ".join(values[:-1]) + " y " + values[-1]

    def _selection_preview(self, selector, all_text, subject):
        values = selector.selected_values()
        if not values:
            return all_text
        joined = self._join_filter_values(values)
        if selector.mode() == "exclude":
            return f"{all_text} excepto {joined}"
        return f"{subject}: {joined}"

    def _period_preview(self):
        return self.dashboard_period_selector.definition()["period_label"]

    @staticmethod
    def _display_iso_date(value):
        parsed = QDate.fromString(str(value or ""), "yyyy-MM-dd")
        return parsed.toString("dd/MM/yyyy") if parsed.isValid() else str(value or "")

    def _update_dashboard_filter_preview(self):
        if not hasattr(self, "dashboard_filter_summary"):
            return
        ars_text = self._selection_preview(self.dashboard_ars, "Todas las ARS", "ARS")
        users_text = self._selection_preview(
            self.dashboard_user, "Todos los facturadores", "Facturadores"
        )
        coverage_text = {
            "Todas": "Todas las coberturas",
            "Asegurados": "Solo asegurados",
            "No asegurados": "Solo no asegurados",
        }.get(self.dashboard_coverage.currentText(), self.dashboard_coverage.currentText())
        advanced_count = self._advanced_filter_count()
        advanced_text = (
            f"{advanced_count} filtro{'s' if advanced_count != 1 else ''} adicional"
            f"{'es' if advanced_count != 1 else ''}"
            if advanced_count else "Sin filtros adicionales"
        )
        self.dashboard_filter_summary.setText(
            f"Mostrando: {self._period_preview()} · {ars_text} · {users_text} · "
            f"{coverage_text} · {advanced_text}"
        )
        self.btn_dashboard_advanced.setText(
            ("▾ Más filtros" if self.btn_dashboard_advanced.isChecked() else "▸ Más filtros")
            + self._advanced_filter_count_text()
        )

    def _dashboard_previous_period(self):
        definition = self.dashboard_period_selector.definition()
        return (
            QDate.fromString(definition["comparison_start"], "yyyy-MM-dd"),
            QDate.fromString(definition["comparison_end"], "yyyy-MM-dd"),
        )

    def _dashboard_trend_granularity(self):
        selected = self.dashboard_granularity.currentText()
        explicit = {"Diario": "day", "Semanal": "week", "Mensual": "month"}
        if selected in explicit:
            return explicit[selected]
        definition = self.dashboard_period_selector.definition()
        start = QDate.fromString(definition["start_date"], "yyyy-MM-dd")
        end = QDate.fromString(definition["end_date"], "yyyy-MM-dd")
        days = start.daysTo(end) + 1
        return "day" if days <= 45 else "week" if days <= 180 else "month"

    def refresh_dashboard(self, *_args):
        if not self._ensure_panel_access():
            return
        if self._dashboard_worker and self._dashboard_worker.isRunning():
            return
        definition = self.dashboard_period_selector.definition()
        start_date = definition["start_date"]
        end_date = definition["end_date"]
        if end_date < start_date:
            FloatingToast("La fecha final no puede ser anterior a la inicial", self, is_error=True).show()
            return
        previous_start, previous_end = self._dashboard_previous_period()
        parameters = (
            start_date,
            end_date,
            self.dashboard_ars.filter_data(),
            self.dashboard_user.filter_data(),
            "Todos los medicamentos",
            self.dashboard_category.currentText(),
            self._dashboard_trend_granularity(),
            self.dashboard_coverage.currentText(),
            self.btn_dashboard_compare.isChecked(),
            previous_start.toString("yyyy-MM-dd"),
            previous_end.toString("yyyy-MM-dd"),
        )
        self._set_dashboard_busy(True, "Consultando y agrupando los datos...")
        worker = DashboardLoadWorker(parameters, self)
        self._dashboard_worker = worker
        worker.loaded.connect(self._dashboard_loaded)
        worker.failed.connect(self._dashboard_failed)
        worker.finished.connect(lambda: self._set_dashboard_busy(False))
        worker.start()

    def _dashboard_loaded(self, data):
        data["period"] = copy.deepcopy(self.dashboard_period_selector.definition())
        self.dashboard_data = data
        self.render_dashboard()
        has_data = bool(data.get("summary", {}).get("receipts"))
        self.dashboard_results_tabs.setVisible(has_data)
        self.dashboard_results_state.setVisible(not has_data)
        if not has_data:
            self.dashboard_results_state.setText(
                "No se encontraron recibos para el período y los filtros seleccionados. "
                "Ajusta la consulta y pulsa Actualizar panel."
            )
            self.dashboard_results_state.setStyleSheet(
                "background: #FFF8E8; color: #7A4B00; border: 1px solid #E7C775; "
                "border-radius: 8px; padding: 11px 13px; font-weight: 700;"
            )
        self.dashboard_status.setText(
            f"Datos actualizados: {datetime.now():%d/%m/%Y %H:%M:%S}"
        )

    def _dashboard_failed(self, error):
        self.dashboard_data = None
        self.dashboard_results_tabs.setVisible(False)
        self.dashboard_results_state.setVisible(True)
        self.dashboard_results_state.setText(
            "No fue posible cargar las estadísticas. Verifica la conexión e inténtalo nuevamente. "
            "Código: PANEL-DATA-001"
        )
        self.dashboard_results_state.setStyleSheet(
            "background: #FDECEC; color: #A32121; border: 1px solid #E7B0B0; "
            "border-radius: 8px; padding: 11px 13px; font-weight: 700;"
        )
        self.dashboard_status.setText("No fue posible actualizar el panel.")
        self.dashboard_guide.setText(
            "No se pudieron cargar las estadísticas. Verifica la conexión y vuelve a intentarlo. "
            "Código: PANEL-DATA-001"
        )
        write_runtime_log(f"Panel de estadísticas [PANEL-DATA-001]: {error}")

    def _set_dashboard_busy(self, busy, message=None):
        self.btn_dashboard_refresh.setEnabled(not busy)
        self.btn_dashboard_compare.setEnabled(not busy)
        self.btn_dashboard_advanced.setEnabled(not busy)
        self.btn_dashboard_clear.setEnabled(not busy)
        can_export = bool(not busy and self.dashboard_data)
        self.btn_export_excel.setEnabled(can_export)
        self.btn_export_panel_pdf.setEnabled(can_export)
        self.btn_dashboard_export.setEnabled(can_export)
        self.btn_print_comparison.setEnabled(
            bool(can_export and self.btn_dashboard_compare.isChecked()
                 and self.dashboard_data.get("previous", {}).get("receipts"))
            if self.dashboard_data else False
        )
        self.dashboard_period_selector.setEnabled(not busy)
        for _label, field in self.dashboard_filter_fields:
            field.setEnabled(not busy)
        for field in (
            self.dashboard_comparison_metric,
            self.dashboard_evolution_metric,
        ):
            field.setEnabled(not busy)
        if message:
            self.dashboard_status.setText(message)
        if busy:
            self.dashboard_results_state.setVisible(True)
            self.dashboard_results_state.setText(
                message or "Cargando y organizando los resultados del período seleccionado..."
            )
            self.dashboard_results_state.setStyleSheet(
                "background: #EDF4FD; color: #24415F; border: 1px solid #C4D7EF; "
                "border-radius: 8px; padding: 11px 13px; font-weight: 700;"
            )
            if not self.dashboard_data:
                self.dashboard_results_tabs.setVisible(False)
        elif self.dashboard_data and self.dashboard_data.get("summary", {}).get("receipts"):
            self.dashboard_results_state.setVisible(False)
            self.dashboard_results_tabs.setVisible(True)

    def clear_dashboard_filters(self):
        today = QDate.currentDate()
        self.dashboard_period_selector.period_type.setCurrentText("Mensual")
        self.dashboard_period_selector.month.setCurrentIndex(today.month() - 1)
        self.dashboard_period_selector.year.setValue(today.year())
        self.dashboard_ars.clear_selection()
        self.dashboard_user.clear_selection()
        self.dashboard_ars.set_mode("include")
        self.dashboard_user.set_mode("include")
        self.dashboard_coverage.setCurrentIndex(0)
        self.dashboard_category.setCurrentIndex(0)
        self.dashboard_granularity.setCurrentIndex(0)
        self.btn_dashboard_advanced.setChecked(False)
        self.btn_dashboard_compare.setChecked(False)
        self.dashboard_comparison_metric.setCurrentIndex(0)
        self.dashboard_evolution_metric.setCurrentIndex(0)
        self.update_dashboard_period(refresh=False)
        self.refresh_dashboard()

    def _mark_dashboard_stale(self, *_args):
        self._update_dashboard_filter_preview()
        if self._dashboard_worker and self._dashboard_worker.isRunning():
            return
        self.btn_export_excel.setEnabled(False)
        self.btn_export_panel_pdf.setEnabled(False)
        self.btn_dashboard_export.setEnabled(False)
        self.btn_print_comparison.setEnabled(False)
        self.dashboard_status.setText("Filtros modificados. Pulsa Actualizar panel para aplicarlos.")
        self.dashboard_results_state.setVisible(True)
        self.dashboard_results_state.setText(
            "Los filtros cambiaron. Los resultados visibles corresponden a la consulta anterior; "
            "pulsa Actualizar panel para aplicar los cambios."
        )
        self.dashboard_results_state.setStyleSheet(
            "background: #FFF8E8; color: #7A4B00; border: 1px solid #E7C775; "
            "border-radius: 8px; padding: 11px 13px; font-weight: 700;"
        )
        if self.dashboard_data:
            self._arrange_dashboard_charts()

    def _comparison_metric_key(self):
        return "receipts" if self.dashboard_comparison_metric.currentText() == "Cantidad de recibos" else "total"

    def _evolution_metric_key(self):
        text = self.dashboard_evolution_metric.currentText()
        if text == "Total de recibos":
            return "receipts"
        if text == "Promedio por recibo":
            return "average"
        return "total"

    def render_dashboard(self, *_args):
        data = self.dashboard_data
        if not data:
            return
        summary = data["summary"]
        self.kpi_total.setText(f"RD$ {summary['total']:,.2f}")
        self.kpi_receipts.setText(f"{summary['receipts']:,}")
        self.kpi_average.setText(f"RD$ {summary['average']:,.2f}")
        self.kpi_room.setText(f"RD$ {summary.get('room', 0):,.2f}")

        previous = data.get("previous") or {}
        comparison_enabled = bool(data.get("filters", {}).get("compare_previous"))
        self.dashboard_comparison_table.setRowCount(0)
        self.dashboard_categories_comparison.setRowCount(0)

        def display_date(value):
            try:
                return datetime.strptime(value, "%Y-%m-%d").strftime("%d/%m/%Y")
            except (TypeError, ValueError):
                return str(value or "")

        def money_text(value, signed=False):
            value = float(value or 0)
            if signed:
                sign = "+" if value > 0 else "-" if value < 0 else ""
                return f"{sign}RD$ {abs(value):,.2f}"
            return f"RD$ {value:,.2f}"

        def variation_style(value):
            if value is None or abs(value) < 0.0001:
                return "•", QColor("#66788A"), QColor("#EEF2F6")
            if value > 0:
                return "▲", QColor("#0B7A5A"), QColor("#E6F5EF")
            return "▼", QColor("#C62828"), QColor("#FDECEC")

        if comparison_enabled and previous.get("receipts", 0):
            self.dashboard_comparison_table.setVisible(True)
            self.dashboard_categories_comparison_box.setVisible(True)
            metric_specs = [
                ("Recibos", "receipts", False),
                ("Total emitido", "total", True),
                ("Promedio por recibo", "average", True),
                ("Sala de emergencia", "room", True),
            ]
            period = data.get("period", {})
            self.dashboard_comparison_periods.setText(
                f"Actual: {period.get('period_label', 'Período actual')} · "
                f"{display_date(data['start_date'])} al {display_date(data['end_date'])}\n"
                f"Anterior: {period.get('comparison_label', 'Período anterior')} · "
                f"{display_date(data['previous_start'])} al {display_date(data['previous_end'])}"
            )
            self.dashboard_comparison_table.setRowCount(len(metric_specs))
            for row_index, (label, key, currency) in enumerate(metric_specs):
                current_value = float(summary.get(key, 0))
                previous_value = float(previous.get(key, 0))
                difference = current_value - previous_value
                variation = difference / previous_value * 100 if previous_value else None
                icon, color, background = variation_style(variation)
                values = [
                    label,
                    money_text(current_value) if currency else f"{int(current_value):,}",
                    money_text(previous_value) if currency else f"{int(previous_value):,}",
                    money_text(difference, signed=True) if currency else f"{difference:+,.0f}",
                    f"{icon} {variation:+.1f}%" if variation is not None else "• Sin base",
                ]
                for column, value in enumerate(values):
                    item = QTableWidgetItem(value)
                    item.setTextAlignment(
                        Qt.AlignLeft | Qt.AlignVCenter if column == 0
                        else Qt.AlignRight | Qt.AlignVCenter
                    )
                    if column in (3, 4):
                        item.setForeground(color)
                        item.setBackground(background)
                        font = item.font(); font.setBold(True); item.setFont(font)
                    self.dashboard_comparison_table.setItem(row_index, column, item)
                self.dashboard_comparison_table.setRowHeight(row_index, 36)

            category_rows = data.get("category_comparison", [])
            self.dashboard_categories_comparison.setRowCount(len(category_rows))
            self.dashboard_categories_comparison.setMinimumHeight(
                max(130, 42 + len(category_rows) * 32)
            )
            for row_index, row in enumerate(category_rows):
                current_value = float(row.get("current", 0))
                previous_value = float(row.get("previous", 0))
                variation = (
                    (current_value - previous_value) / previous_value * 100
                    if previous_value else None
                )
                icon, color, background = variation_style(variation)
                label_item = QTableWidgetItem(str(row.get("label", "")))
                variation_item = QTableWidgetItem(
                    f"{icon} {variation:+.1f}%" if variation is not None else "• Sin base comparable"
                )
                variation_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                variation_item.setForeground(color)
                variation_item.setBackground(background)
                font = variation_item.font(); font.setBold(True); variation_item.setFont(font)
                self.dashboard_categories_comparison.setItem(row_index, 0, label_item)
                self.dashboard_categories_comparison.setItem(row_index, 1, variation_item)
                self.dashboard_categories_comparison.setRowHeight(row_index, 32)
            self.dashboard_comparison_panel.setVisible(True)
        elif comparison_enabled:
            self.dashboard_comparison_periods.setText(
                "No existen datos en el período anterior para realizar una comparación válida."
            )
            self.dashboard_comparison_table.setVisible(False)
            self.dashboard_categories_comparison_box.setVisible(False)
            self.dashboard_comparison_panel.setVisible(True)
        else:
            self.dashboard_comparison_panel.setVisible(False)

        comparison_metric = self._comparison_metric_key()
        evolution_metric = self._evolution_metric_key()
        show_ars_comparison = bool(data.get("show_ars_comparison"))
        comparison_rows = list(data.get("summary_table", [])) if show_ars_comparison else []
        percentage_key = "receipt_percentage" if comparison_metric == "receipts" else "money_percentage"
        comparison_rows.sort(key=lambda row: row.get(percentage_key, 0), reverse=True)
        comparison_entries = [(row["label"], row.get(percentage_key, 0) * 100) for row in comparison_rows]
        trend_entries = [(row["label"], row.get(evolution_metric, 0)) for row in data["trend"]]
        category_entries = [(row["label"], row["total"]) for row in data["categories"]]
        distribution_entries = [
            (row["label"], row["total"]) for row in data.get("category_distribution", data["categories"])
        ]
        coverage_rows = list(data.get("coverage", []))
        coverage_total = sum(row.get("receipts", 0) for row in coverage_rows)
        coverage_entries = [
            (row["label"], row.get("receipts", 0) / coverage_total * 100 if coverage_total else 0)
            for row in coverage_rows
        ]
        evolution_currency = evolution_metric in ("total", "average")
        self.comparison_chart.set_entries(comparison_entries, percent=True)
        self.category_bar_chart.set_entries(category_entries, currency=True)
        self.line_chart.set_entries(trend_entries, currency=evolution_currency)
        self.donut_chart.set_entries(distribution_entries)
        self.coverage_chart.set_entries(coverage_entries, percent=True)
        self.comparison_box.setTitle("Distribución porcentual por ARS")
        granularity_title = {"day": "diaria", "week": "semanal", "month": "mensual"}.get(
            data.get("filters", {}).get("trend_granularity"), "temporal"
        )
        self.line_box.setTitle(f"Evolución {granularity_title}")
        self.coverage_box.setVisible(bool(coverage_rows))
        self.comparison_box.setVisible(show_ars_comparison)
        self.dashboard_table_box.setVisible(show_ars_comparison)
        self.dashboard_results_tabs.setTabVisible(
            self.dashboard_comparison_tab_index, comparison_enabled
        )
        self.dashboard_results_tabs.setTabVisible(
            self.dashboard_detail_tab_index, show_ars_comparison
        )

        data["comparison"] = comparison_rows
        data["view"] = {
            "show_ars_comparison": show_ars_comparison,
            "ars_metric": comparison_metric,
            "ars_metric_label": self.dashboard_comparison_metric.currentText(),
            "evolution_metric": evolution_metric,
            "evolution_label": self.dashboard_evolution_metric.currentText(),
            "coverage_visible": bool(coverage_rows),
        }
        self.btn_print_comparison.setEnabled(
            bool(comparison_enabled and previous.get("receipts", 0))
        )
        self._arrange_dashboard_charts()

        self.dashboard_table.setRowCount(0)
        for row in data.get("summary_table", []):
            table_row = self.dashboard_table.rowCount()
            self.dashboard_table.insertRow(table_row)
            values = [row["label"], f"RD$ {row['total']:,.2f}", f"{row['receipts']:,}",
                      f"{row['money_percentage'] * 100:.1f}%",
                      f"{row['receipt_percentage'] * 100:.1f}%",
                      f"RD$ {row['average']:,.2f}"]
            for column, value in enumerate(values):
                item = QTableWidgetItem(str(value))
                if column >= 1:
                    item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
                self.dashboard_table.setItem(table_row, column, item)

        best_period = max(data["trend"], key=lambda row: row.get(evolution_metric, 0), default=None)
        best_category = max(data["categories"], key=lambda row: row["total"], default=None)
        best_group = max(comparison_rows, key=lambda row: row.get(percentage_key, 0), default=None)
        sentences = [
            f"Entre {data['start_date']} y {data['end_date']} se emitieron "
            f"{summary['receipts']:,} recibos por RD$ {summary['total']:,.2f}."
        ]
        if best_period:
            sentences.append(
                f"El día con mayor {self.dashboard_evolution_metric.currentText().lower()} "
                f"fue {best_period['label']}."
            )
        if best_category:
            sentences.append(
                f"La categoría con mayor peso económico fue {best_category['label']} "
                f"con RD$ {best_category['total']:,.2f} "
                f"({best_category.get('percentage', 0) * 100:.1f}% del total por categorías)."
            )
        if best_group:
            sentences.append(
                f"La ARS con mayor participación según {self.dashboard_comparison_metric.currentText().lower()} "
                f"fue {best_group['label']} con {best_group[percentage_key] * 100:.1f}%."
            )
        uninsured = next((row for row in coverage_rows if row["label"] == "No asegurados"), None)
        if uninsured and coverage_total:
            sentences.append(
                f"Las atenciones no aseguradas fueron {uninsured['receipts']:,}, "
                f"equivalentes al {uninsured['receipts'] / coverage_total * 100:.1f}% del período."
            )
        self.dashboard_guide.setText(" ".join(sentences))
        self._update_dashboard_filter_preview()

    def _metric_text(self, value, currency):
        return f"RD$ {float(value):,.2f}" if currency else f"{int(value):,} recibos"

    def _arrange_dashboard_charts(self):
        compact = getattr(self, "_dashboard_compact", False)
        show_ars = bool(
            self.dashboard_data and self.dashboard_data.get("view", {}).get("show_ars_comparison")
        )
        show_coverage = bool(
            self.dashboard_data and self.dashboard_data.get("view", {}).get("coverage_visible")
        )
        for widget in self._dashboard_chart_widgets:
            self.dashboard_charts_layout.removeWidget(widget)
        if compact:
            row = 0
            for widget in (self.donut_box, self.category_box):
                self.dashboard_charts_layout.addWidget(widget, row, 0)
                row += 1
            if show_ars:
                self.dashboard_charts_layout.addWidget(self.comparison_box, row, 0)
                row += 1
            if show_coverage:
                self.dashboard_charts_layout.addWidget(self.coverage_box, row, 0)
        else:
            self.dashboard_charts_layout.addWidget(self.donut_box, 0, 0)
            self.dashboard_charts_layout.addWidget(self.category_box, 0, 1)
            if show_ars:
                self.dashboard_charts_layout.addWidget(self.comparison_box, 1, 0, 1, 2)
                next_row = 2
            else:
                next_row = 1
            if show_coverage:
                self.dashboard_charts_layout.addWidget(self.coverage_box, next_row, 0, 1, 2)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if not hasattr(self, "dashboard_charts_layout"):
            return
        compact = self.width() < 1080
        if compact == getattr(self, "_dashboard_compact", None):
            return
        self._dashboard_compact = compact
        self._arrange_dashboard_filters(compact)
        for card in self.dashboard_kpi_cards:
            self.dashboard_kpi_layout.removeWidget(card)
        if compact:
            for row, card in enumerate(self.dashboard_kpi_cards):
                self.dashboard_kpi_layout.addWidget(card, row, 0)
        else:
            for column, card in enumerate(self.dashboard_kpi_cards):
                self.dashboard_kpi_layout.addWidget(card, 0, column)
        self._arrange_dashboard_charts()

    def closeEvent(self, event):
        active = any(
            worker and worker.isRunning()
            for worker in (self._dashboard_worker, self._dashboard_export_worker)
        )
        if active:
            QMessageBox.information(
                self,
                "Operación en curso",
                "Espera a que termine la consulta o exportación antes de cerrar esta ventana.",
            )
            event.ignore()
            return
        super().closeEvent(event)

    def update_report_preview(self, *_args):
        if not hasattr(self, "report_preview"):
            return
        definition = self.report_period_selector.definition()
        ars_text = self._selection_preview(self.ars_filter, "Todas las ARS", "ARS")
        user_text = self._selection_preview(
            self.user_filter, "Todos los facturadores", "Facturadores"
        )
        self.report_preview.setText(
            f"Período: {definition['period_label']}\n"
            f"Fechas exactas: {self._display_iso_date(definition['start_date'])} al "
            f"{self._display_iso_date(definition['end_date'])}\n"
            f"ARS: {ars_text}\nFacturadores: {user_text}\n"
            "Criterio: fecha real de generación del recibo."
        )

    def generate_selected_report(self):
        definition = self.report_period_selector.definition()
        start_date = definition["start_date"]
        end_date = definition["end_date"]
        ars_filter = self.ars_filter.filter_data()
        user_filter = self.user_filter.filter_data()
        if end_date < start_date:
            FloatingToast("Fechas inválidas", self, is_error=True).show()
            return
        try:
            if definition["period_type"] == "Diario":
                path = generate_daily_report_pdf(
                    start_date, self.current_user["username"], None, ars_filter, user_filter
                )
            else:
                path = generate_period_report_pdf(
                    definition["period_label"], start_date, end_date,
                    self.current_user["username"], None, ars_filter, user_filter,
                    period_metadata=definition,
                )
            if not path:
                QMessageBox.information(self, "Reportes", "No hay datos para ese período.")
                return
            self.load_rows()
            FloatingToast("✅ Reporte generado", self).show()
            if not open_file_path(path):
                QMessageBox.warning(
                    self, "Reportes",
                    f"El reporte se generó, pero no se pudo abrir automáticamente:\n{path}",
                )
        except Exception as exc:
            QMessageBox.critical(self, "Reportes", f"No se pudo generar el reporte:\n{exc}")

    def load_rows(self):
        super().load_rows()
        for row in range(self.table.rowCount()):
            raw_type = self.table.item(row, 0).text() if self.table.item(row, 0) else "Reporte PDF"
            raw_json = self.table.item(row, 6).text() if self.table.item(row, 6) else "{}"
            try:
                metadata = json.loads(raw_json or "{}")
            except (TypeError, ValueError, json.JSONDecodeError):
                metadata = {}
            period = metadata.get("period") or metadata.get("_period") or {}
            is_comparison = raw_type.startswith("Reporte comparativo")
            display_type = "Comparativo" if is_comparison else (
                "Diario" if raw_type.startswith("Diario") else "Reporte PDF"
            )
            current_period = period.get("period_label") or raw_type.split(" (")[0]
            compared_period = period.get("comparison_label", "") if is_comparison else ""
            type_item = QTableWidgetItem(display_type)
            type_item.setToolTip(raw_type)
            current_item = QTableWidgetItem(current_period)
            current_item.setToolTip(current_period)
            compared_item = QTableWidgetItem(compared_period or "No aplica")
            compared_item.setToolTip(compared_period or "Este reporte no es comparativo")
            self.table.setItem(row, 0, type_item)
            self.table.setItem(row, 9, current_item)
            self.table.setItem(row, 10, compared_item)
            self.table.setRowHeight(row, 34)
        if hasattr(self, "history_count"):
            self.history_count.setText(f"{self.table.rowCount():,} reportes")

    def _selected_report_path(self):
        row = self.table.currentRow()
        if row < 0:
            FloatingToast("Selecciona un reporte", self, is_error=True).show()
            return ""
        stored_path = self.table.item(row, 5).text().strip() if self.table.item(row, 5) else ""
        filename = os.path.basename(stored_path)
        if not filename:
            QMessageBox.warning(self, "Reportes", "Este reporte no tiene un archivo PDF asociado.")
            return ""
        path = stored_path if os.path.exists(stored_path) else stable_storage_path(REPORTS_DIR, filename)
        if not os.path.exists(path):
            with db_connect() as con:
                result = con.execute(
                    "SELECT file_data FROM pdf_storage WHERE filename=%s", (filename,)
                ).fetchone()
            if result:
                os.makedirs(os.path.dirname(path), exist_ok=True)
                with open(path, "wb") as pdf_file:
                    pdf_file.write(bytes(result["file_data"]))
        if not os.path.exists(path):
            QMessageBox.warning(self, "Reportes", "No se encontró el PDF de este reporte.")
            return ""
        return path

    def save_selected_report_copy(self):
        path = self._selected_report_path()
        if not path:
            return
        destination, _ = QFileDialog.getSaveFileName(
            self, "Guardar una copia del reporte", os.path.basename(path), "PDF (*.pdf)"
        )
        if not destination:
            return
        if not destination.lower().endswith(".pdf"):
            destination += ".pdf"
        try:
            shutil.copy2(path, destination)
            FloatingToast("Copia del reporte guardada", self).show()
        except Exception as exc:
            QMessageBox.critical(self, "Guardar copia", f"No se pudo guardar la copia:\n{exc}")

    def print_selected_report(self):
        path = self._selected_report_path()
        if not path:
            return
        ComparisonPdfDialog(
            path,
            self,
            dialog_title="Reporte listo para revisar",
            detail_text=(
                "Verifica el documento antes de imprimir. También puedes guardar una copia "
                "en otra ubicación."
            ),
        ).exec()

    def show_context_menu(self, pos):
        row = self.table.rowAt(pos.y())
        if row < 0:
            return
        self.table.selectRow(row)
        menu = QMenu(self)
        actions = [
            ("Abrir PDF", self.open_selected),
            ("Guardar una copia", self.save_selected_report_copy),
            ("Vista previa e imprimir", self.print_selected_report),
        ]
        for label, callback in actions:
            action = menu.addAction(label)
            action.triggered.connect(callback)
        menu.addSeparator()
        delete_action = menu.addAction("Eliminar reporte")
        delete_action.triggered.connect(self.delete_selected_report)
        menu.exec(self.table.viewport().mapToGlobal(pos))

    def filter_history(self, text):
        filter_table_widget(self.table, text)
        visible = sum(not self.table.isRowHidden(row) for row in range(self.table.rowCount()))
        self.history_count.setText(f"{visible:,} de {self.table.rowCount():,} reportes")

    def export_dashboard_excel(self):
        if not self._ensure_panel_access() or not self.dashboard_data:
            if self.panel_access and not self.dashboard_data:
                FloatingToast("Actualiza el panel antes de exportar", self, is_error=True).show()
            return
        default_name = f"panel_estadistico_{self.dashboard_data['start_date']}_{self.dashboard_data['end_date']}.xlsx"
        path, _ = QFileDialog.getSaveFileName(self, "Exportar panel a Excel", default_name, "Excel (*.xlsx)")
        if not path:
            return
        if not path.lower().endswith(".xlsx"):
            path += ".xlsx"
        self._start_dashboard_export("xlsx", path)

    def print_dashboard_comparison(self):
        if not self._ensure_panel_access() or not self.dashboard_data:
            FloatingToast("Actualiza el panel antes de imprimir", self, is_error=True).show()
            return
        if not self.btn_dashboard_compare.isChecked():
            FloatingToast("Activa la comparación y actualiza el panel", self, is_error=True).show()
            return
        if not self.dashboard_data.get("previous", {}).get("receipts"):
            FloatingToast(
                "No existen datos en el período anterior para imprimir una comparación válida",
                self,
                is_error=True,
            ).show()
            return
        try:
            snapshot = copy.deepcopy(self.dashboard_data)
            path = generate_comparison_report_pdf(
                snapshot, self.current_user.get("username", "Usuario")
            )
            self.load_rows()
            ComparisonPdfDialog(path, self).exec()
        except Exception as exc:
            write_runtime_log(f"Reporte comparativo: {exc}")
            QMessageBox.critical(
                self, "Imprimir comparación", f"No se pudo generar el reporte comparativo:\n{exc}"
            )

    def export_dashboard_pdf(self):
        if not self._ensure_panel_access() or not self.dashboard_data:
            if self.panel_access and not self.dashboard_data:
                FloatingToast("Actualiza el panel antes de exportar", self, is_error=True).show()
            return
        default_name = f"panel_estadistico_{self.dashboard_data['start_date']}_{self.dashboard_data['end_date']}.pdf"
        path, _ = QFileDialog.getSaveFileName(self, "Exportar panel a PDF", default_name, "PDF (*.pdf)")
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"
        self._start_dashboard_export("pdf", path)

    def _start_dashboard_export(self, mode, path):
        if self._dashboard_export_worker and self._dashboard_export_worker.isRunning():
            FloatingToast("Ya existe una exportación en proceso", self, is_error=True).show()
            return
        self.render_dashboard()
        self._set_dashboard_busy(
            True,
            "Creando el archivo Excel..." if mode == "xlsx" else "Creando el PDF institucional...",
        )
        worker = DashboardExportWorker(
            mode,
            copy.deepcopy(self.dashboard_data),
            path,
            self.current_user.get("username", "Sistema"),
            LOGO_PATH or "",
            self,
        )
        self._dashboard_export_worker = worker
        worker.exported.connect(
            lambda result: self._dashboard_exported(result, mode)
        )
        worker.failed.connect(
            lambda error: self._dashboard_export_failed(error, mode)
        )
        worker.finished.connect(lambda: self._set_dashboard_busy(False))
        worker.start()

    def _dashboard_exported(self, result, mode):
        label = "Excel" if mode == "xlsx" else "PDF"
        self.dashboard_status.setText(f"Exportación {label} completada: {datetime.now():%H:%M:%S}")
        FloatingToast(f"✅ Panel exportado a {label}", self).show()
        open_file_path(result)

    def _dashboard_export_failed(self, error, mode):
        code = "PANEL-XLSX-001" if mode == "xlsx" else "PANEL-PDF-001"
        label = "Excel" if mode == "xlsx" else "PDF"
        write_runtime_log(f"Exportación {label} [{code}]: {error}")
        QMessageBox.critical(
            self,
            f"Exportar {label}",
            f"No se pudo crear el archivo {label}.\nEl archivo no fue guardado.\nCódigo: {code}",
        )


class HonorarioSelectorDialog(QDialog):
    def __init__(self, honorarios: dict, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Agregar honorario")
        self.setMinimumSize(560, 360)
        self.selected_item = None

        lay = QVBoxLayout(self)
        title = QLabel("Selecciona un honorario para agregar al recibo:")
        title.setWordWrap(True)
        title.setStyleSheet("font-weight: bold; font-size: 12pt;")
        lay.addWidget(title)

        def sort_key(item):
            name, _price = item
            n = remove_accents(name)
            if "interconsulta" in n or "inter consulta" in n:
                priority = 2
            elif "noche" in n or "nocturn" in n or "6pm" in n or "9pm" in n or " pm" in f" {n}":
                priority = 1
            elif "dia" in n or "diurno" in n or "emergencia" in n or "consulta" in n:
                priority = 0
            else:
                priority = 3
            return (priority, n)

        self.ordered_items = sorted(honorarios.items(), key=sort_key)
        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QAbstractItemView.SingleSelection)
        self.list_widget.setAlternatingRowColors(True)
        self.list_widget.setStyleSheet("QListWidget::item { padding: 10px; font-size: 11pt; }")

        for name, price in self.ordered_items:
            item = QListWidgetItem(f"{name}   —   ${float(price):,.2f}")
            item.setData(Qt.UserRole, (name, float(price)))
            font = item.font()
            font.setPointSize(11)
            font.setBold(True)
            item.setFont(font)
            self.list_widget.addItem(item)

        if self.list_widget.count() > 0:
            self.list_widget.setCurrentRow(0)

        self.list_widget.itemDoubleClicked.connect(lambda _item: self._accept())
        lay.addWidget(self.list_widget, 1)

        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.button(QDialogButtonBox.Ok).setText("Agregar honorario")
        btns.button(QDialogButtonBox.Cancel).setText("Omitir")
        btns.accepted.connect(self._accept)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)

        for btn in btns.buttons():
            if btn is btns.button(QDialogButtonBox.Ok):
                set_button_role(btn, "success")
            else:
                set_button_role(btn, "neutral")

    def _accept(self):
        item = self.list_widget.currentItem()
        if item is None:
            FloatingToast("Selecciona un honorario o pulsa Omitir", self, is_error=True).show()
            return
        self.selected_item = item.data(Qt.UserRole)
        self.accept()

    def values(self):
        return [self.selected_item] if self.selected_item else []

class MascarillaSelectorDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Seleccionar mascarilla para nebulizar")
        self.setFixedSize(420, 300)
        self.selected_name = None
        root = QVBoxLayout(self)
        root.setContentsMargins(18, 18, 18, 18)
        root.setSpacing(10)
        lbl = QLabel("Se detectó \"nebulizar con\". Selecciona la mascarilla a agregar:")
        lbl.setWordWrap(True)
        lbl.setStyleSheet("font-weight: bold; font-size: 11pt;")
        root.addWidget(lbl)
        self.list_widget = QListWidget()
        self.list_widget.setAlternatingRowColors(True)
        self.list_widget.setStyleSheet("QListWidget::item { padding: 10px; } QListWidget::item:selected { background-color: #1565c0; color: #ffffff; }")
        root.addWidget(self.list_widget, 1)
        btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        btns.button(QDialogButtonBox.Ok).setText("Agregar")
        btns.button(QDialogButtonBox.Cancel).setText("Omitir")
        btns.accepted.connect(self._accept)
        btns.rejected.connect(self.reject)
        root.addWidget(btns)
        for btn in btns.buttons():
            set_button_role(btn, "success" if btn is btns.button(QDialogButtonBox.Ok) else "neutral")
        self.list_widget.itemDoubleClicked.connect(self._accept)

    def load_items(self, items: dict):
        self.list_widget.clear()
        for n, p in sorted(items.items(), key=lambda x: x[0]):
            it = QListWidgetItem(f"{n}\n${p:,.2f}")
            it.setData(Qt.UserRole, n)
            font = it.font(); font.setPointSize(10); it.setFont(font)
            self.list_widget.addItem(it)
        if self.list_widget.count() > 0:
            self.list_widget.setCurrentRow(0)

    def _accept(self):
        item = self.list_widget.currentItem()
        if item:
            self.selected_name = item.data(Qt.UserRole)
        self.accept()

class IdleEventFilter(QObject):
    activity = Signal()
    def eventFilter(self, obj, event):
        if event.type() in (QEvent.MouseMove, QEvent.MouseButtonPress, QEvent.MouseButtonRelease, QEvent.KeyPress, QEvent.Wheel, QEvent.FocusIn, QEvent.TouchBegin):
            self.activity.emit()
        return super().eventFilter(obj, event)

class MainWindow(QMainWindow):
    logout_requested = Signal()
    theme_toggled = Signal(bool)

    def __init__(self, current_user: dict):
        super().__init__()
        self.current_user = current_user
        self.setWindowTitle(APP_TITLE)
        self.setMinimumSize(800, 600) 
        self.current_ars = ""
        self.locked_ars = None 
        self.universal = {cat: get_universal(cat) for cat in UNIVERSAL_CATEGORIES}
        self.ars_cache = {cat: {} for cat in ARS_CATEGORIES}
        self.last_activity = datetime.now()
        self.warned_idle = False
        self.editing_recibo_id = None
        self.editing_recibo_numero = None
        self.is_dark_mode = False
        self.bajante_added_for_solucion = False
        self._reverting_ars_change = False
        self.honorario_prompt_done = False
        self.mascarilla_prompt_done = False
        self.session_started_at = now_str()
        self.session_id = make_session_id()
        self._last_generate_time = 0.0
        try:
            register_active_session(self.current_user.get("username", ""), self.session_id)
        except Exception as e:
            print(f"[SESIONES] No se pudo registrar la sesión activa: {e}")
        screen = QApplication.primaryScreen()
        screen_w = screen.availableGeometry().width() if screen else 1366
        self.catalog_font_size = 12 if screen_w >= 1800 else 11 if screen_w >= 1500 else 10
        self.catalog_left_width = 560 if screen_w >= 1800 else 500 if screen_w >= 1500 else 420
        self.catalog_tab_font_size = 10 if screen_w >= 1800 else 9 if screen_w >= 1500 else 8
        self._responsive_mode = None

        username = self.current_user.get("username", "") if self.current_user else ""
        self.preferences = get_user_preferences(username)
        if self.preferences.get("theme") == "oscuro":
            self.is_dark_mode = True

        self._start_pdf_services()
        self._build_ui()
        self._setup_hotkeys()
        self._build_timers()
        self.apply_button_colors()
        self.on_ars_changed(self.ars_combo.currentText())
        self.refresh_picker()
        self.safe_startup_load()

    def _start_pdf_services(self):
        """Precarga Chromium y mantiene una cola separada para la nube."""
        self.pdf_sync_worker = PDFStorageSyncWorker(self)
        self.pdf_sync_worker.start()

        self.pdf_worker = PDFDatabaseWorker()
        self.pdf_worker.signals.finished_signal.connect(self.on_pdf_generated)
        self.pdf_worker.signals.sync_requested.connect(self.pdf_sync_worker.enqueue)
        self.pdf_worker.start()

    def _build_ui(self):
        central = QWidget(); self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        
        header_widget = QWidget()
        header_widget.setObjectName("HeaderWidget")
        header = QHBoxLayout(header_widget)
        header.setContentsMargins(20, 15, 20, 15)
        
        if LOGO_PATH and os.path.exists(LOGO_PATH):
            logo = QLabel(); pix = QPixmap(LOGO_PATH)
            if not pix.isNull(): logo.setPixmap(pix.scaled(120, 70, Qt.KeepAspectRatio, Qt.SmoothTransformation))
            header.addWidget(logo)
        else:
            header.addWidget(QLabel("🏥"))
            
        title_wrap = QVBoxLayout()
        title = QLabel(APP_TITLE)
        title.setStyleSheet("font-size: 24pt; font-weight: 800;")
        subtitle = QLabel("Hospital Provincial Dr. Ángel Contreras Mejía · Facturación de medicamentos y servicios médicos")
        subtitle.setStyleSheet("font-size: 11pt; font-weight: normal;")
        self.header_title = title
        self.header_subtitle = subtitle
        title_wrap.addWidget(title); title_wrap.addWidget(subtitle); header.addLayout(title_wrap, 1)
        
        top_right = QVBoxLayout()
        top_right_h = QHBoxLayout()
        
        self.lbl_modo_visual = QLabel("Cambiar Tema:")
        self.lbl_modo_visual.setStyleSheet("font-size: 11pt; font-weight: bold;")
        
        self.btn_theme_toggle = QPushButton("🌙")
        self.btn_theme_toggle.setFixedSize(36, 36)
        self.btn_theme_toggle.setCursor(Qt.PointingHandCursor)
        self.btn_theme_toggle.setStyleSheet("background-color: transparent; font-size: 20pt; border: none;")
        self.btn_theme_toggle.clicked.connect(self.toggle_theme)
        
        self.lbl_user_top = QLabel(f"👤 {self.current_user.get('full_name')} | 🏷️ {self.current_user.get('role').capitalize()}")
        self.lbl_user_top.setStyleSheet("font-size: 11pt; font-weight: bold;")
        
        top_right_h.addWidget(self.lbl_modo_visual)
        top_right_h.addWidget(self.btn_theme_toggle)
        top_right_h.addSpacing(15)
        self.btn_preferences = QPushButton("⚙️")
        self.btn_preferences.setFixedSize(36, 36)
        self.btn_preferences.setCursor(Qt.PointingHandCursor)
        self.btn_preferences.setStyleSheet("background-color: transparent; font-size: 18pt; border: none;")
        self.btn_preferences.clicked.connect(self.open_preferences_dialog)
        top_right_h.addWidget(self.btn_preferences)
        top_right_h.addSpacing(15)
        top_right_h.addWidget(self.lbl_user_top)
        
        top_right.addLayout(top_right_h)
        top_right.addStretch(1)
        header.addLayout(top_right)
        root.addWidget(header_widget)

        nav_widget = QWidget()
        nav_widget.setObjectName("NavWidget")
        nav_lay = QHBoxLayout(nav_widget)
        nav_lay.setContentsMargins(20, 5, 20, 5)
        
        self.btn_view_history = QPushButton("Historial de Acciones") if user_is_admin(self.current_user) else None
        self.btn_receipts_history = QPushButton("Recibos Guardados")
        self.btn_view_reports = QPushButton("Reportes")
        self.btn_admin_users = QPushButton("Gestión de Usuarios") if user_can_manage_sessions(self.current_user) else None
        self.btn_logout = QPushButton("🚪 Cerrar Sesión")

        for b in [self.btn_view_history, self.btn_receipts_history, self.btn_view_reports, self.btn_admin_users]:
            if b: 
                b.setStyleSheet("""
                    QPushButton {
                        background-color: #F8F9FA; 
                        color: #212529; 
                        font-weight: bold; 
                        font-size: 10pt; 
                        padding: 8px 15px;
                        border-radius: 4px;
                    }
                    QPushButton:hover {
                        background-color: #E2E6EA;
                    }
                """)
                nav_lay.addWidget(b)
        
        nav_lay.addStretch(1)
        self.btn_logout.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                font-weight: bold;
                padding: 8px 15px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        self.btn_logout.clicked.connect(self.force_logout)
        nav_lay.addWidget(self.btn_logout)
        root.addWidget(nav_widget)

        content_widget = QWidget()
        content_lay = QVBoxLayout(content_widget)
        content_lay.setContentsMargins(20, 15, 20, 10)

        billing_group = QGroupBox("Datos de Facturación")
        billing_lay = QGridLayout(billing_group)
        billing_lay.setHorizontalSpacing(8)
        billing_lay.setVerticalSpacing(8)
        # Las etiquetas conservan solo su ancho natural; los campos reciben
        # el espacio sobrante. Esto evita grandes vacíos en monitores anchos.
        billing_lay.setColumnStretch(0, 0)
        billing_lay.setColumnStretch(1, 5)
        billing_lay.setColumnStretch(2, 0)
        billing_lay.setColumnStretch(3, 0)
        billing_lay.setColumnStretch(4, 0)
        billing_lay.setColumnStretch(5, 1)

        self.name_edit = QLineEdit(); self.name_edit.setPlaceholderText("Nombre del paciente")
        self.date_edit = QDateEdit(); self.date_edit.setCalendarPopup(True); self.date_edit.setDisplayFormat("yyyy-MM-dd"); self.date_edit.setDate(QDate.currentDate()); self.date_edit.setMinimumWidth(150)
        self.dx_edit = QLineEdit(); self.dx_edit.setPlaceholderText("Diagnóstico (DX)")
        billing_lay.addWidget(QLabel("Paciente:"), 0, 0); billing_lay.addWidget(self.name_edit, 0, 1, 1, 3)
        billing_lay.addWidget(QLabel("Fecha:"), 0, 4); billing_lay.addWidget(self.date_edit, 0, 5)
        billing_lay.addWidget(QLabel("Diagnóstico:"), 1, 0); billing_lay.addWidget(self.dx_edit, 1, 1, 1, 3)

        self.ars_combo = QComboBox(); self.ars_combo.addItems(ars_list()); self.ars_combo.setMaximumWidth(240)
        self.coverage_combo = QComboBox()
        self.coverage_combo.addItems(["Asegurado", "No asegurado"])
        self.sala_spin = QDoubleSpinBox(); self.sala_spin.setRange(0, 1_000_000); self.sala_spin.setDecimals(2); self.sala_spin.setSingleStep(SALA_STEP)
        self.sala_spin.valueChanged.connect(lambda v: self.update_totals())

        self.btn_add_catalog_item = QPushButton("Nuevo Ítem")
        self.btn_manage_catalog = QPushButton("Catálogo")
        self.btn_ars_mgmt = QPushButton("Gestión ARS")
        self.btn_import_meds = QPushButton("Importar Word")
        
        billing_lay.addWidget(QLabel("ARS:"), 1, 4); billing_lay.addWidget(self.ars_combo, 1, 5)
        billing_lay.addWidget(QLabel("Sala de emergencia:"), 2, 0); billing_lay.addWidget(self.sala_spin, 2, 1)
        billing_lay.addWidget(QLabel("Cobertura:"), 2, 4); billing_lay.addWidget(self.coverage_combo, 2, 5)
        actions = QHBoxLayout()
        actions.addStretch(1)
        actions.addWidget(self.btn_add_catalog_item); actions.addWidget(self.btn_manage_catalog)
        actions.addWidget(self.btn_ars_mgmt); actions.addWidget(self.btn_import_meds)
        billing_lay.addLayout(actions, 3, 0, 1, 6)
        
        content_lay.addWidget(billing_group)

        main_split = QSplitter(Qt.Horizontal)
        left = QWidget(); left_v = QVBoxLayout(left); left_v.setContentsMargins(0, 10, 10, 0)
        middle = QWidget(); middle_v = QVBoxLayout(middle); middle_v.setContentsMargins(10, 10, 0, 0)
        left.setMinimumWidth(300)
        middle.setMinimumWidth(390)
        self.main_split = main_split
        main_split.addWidget(left); main_split.addWidget(middle)
        main_split.setChildrenCollapsible(False)
        main_split.setStretchFactor(0, 1) 
        main_split.setStretchFactor(1, 3)
        main_split.setSizes([self.catalog_left_width, max(760, self.width() - self.catalog_left_width)]) 
        main_split.setStyleSheet("QSplitter::handle { width: 1px; background: #d1d9e6; }")
        if self.is_dark_mode:
            main_split.setStyleSheet("QSplitter::handle { width: 1px; background: #333333; }")
        content_lay.addWidget(main_split, 1)
        root.addWidget(content_widget, 1)

        self.tabs = QTabWidget()
        self.tabs.setUsesScrollButtons(False)
        self.tabs.setElideMode(Qt.ElideNone)
        self.tabs.tabBar().setExpanding(False)
        self.tabs.setStyleSheet(f"QTabBar::tab {{ padding: 6px 8px; font-size: {self.catalog_tab_font_size}pt; }}")
        self.compact_tab_labels = {
            "Medicamentos": "💊 Med.",
            "Materiales": "📦 Mat.",
            "Laboratorios": "🩸 Lab.",
            "Imágenes": "📷 Img.",
            "Procedimientos": "🩺 Proc.",
            "Honorarios": "👨‍⚕️ Hon.",
        }
        self.source_lists = {cat: QListWidget() for cat in ALL_CATEGORIES}

        def crear_callback_menu(cat_name):
            return lambda pos: self.show_catalog_context_menu(pos, cat_name)

        for cat in ALL_CATEGORIES:
            tab_idx = self.tabs.addTab(self.source_lists[cat], self.compact_tab_labels.get(cat, f"{CAT_EMOJIS.get(cat, '')} {cat}"))
            self.tabs.setTabToolTip(tab_idx, cat)
            self.source_lists[cat].setContextMenuPolicy(Qt.CustomContextMenu)
            self.source_lists[cat].customContextMenuRequested.connect(crear_callback_menu(cat))

        left_v.addWidget(self.tabs)
        search_row = QHBoxLayout(); self.search = QLineEdit(); self.search.setPlaceholderText("🔍 Buscar en catálogo (F3)...")
        search_row.addWidget(self.search, 1); left_v.addLayout(search_row)
        
        qty_row = QHBoxLayout()
        
        self.qty = QSpinBox()
        self.qty.setRange(1, 300)
        self.qty.setValue(1)
        self.qty.setButtonSymbols(QAbstractSpinBox.NoButtons)
        self.qty.setMinimumWidth(60)

        btn_qty_up = QToolButton()
        btn_qty_up.setText("▲") 
        btn_qty_up.setObjectName("SpinArrowBtn") 
        btn_qty_up.setFixedSize(30, 30) 
        btn_qty_up.clicked.connect(lambda: self.qty.setValue(self.qty.value() + 1))

        btn_qty_down = QToolButton()
        btn_qty_down.setText("▼") 
        btn_qty_down.setObjectName("SpinArrowBtn") 
        btn_qty_down.setFixedSize(30, 30)
        btn_qty_down.clicked.connect(lambda: self.qty.setValue(max(1, self.qty.value() - 1)))

        qty_container = QWidget()
        qty_lay = QHBoxLayout(qty_container)
        qty_lay.setContentsMargins(0, 0, 0, 0)
        qty_lay.setSpacing(2)
        qty_lay.addWidget(self.qty)
        qty_lay.addWidget(btn_qty_up)
        qty_lay.addWidget(btn_qty_down)

        self.btn_add = QPushButton("➡️ Añadir al Recibo (Enter)")
        qty_row.addWidget(QLabel("Cant:"))
        qty_row.addWidget(qty_container)
        qty_row.addStretch(1)
        qty_row.addWidget(self.btn_add)
        left_v.addLayout(qty_row)

        self.shortcut_return = QShortcut(QKeySequence("Return"), self)
        self.shortcut_return.activated.connect(self.btn_add.click)
        
        self.shortcut_enter = QShortcut(QKeySequence("Enter"), self)
        self.shortcut_enter.activated.connect(self.btn_add.click)

        cart_group = QGroupBox("🛒 Recibo de Facturación")
        cart_lay = QVBoxLayout(cart_group)
        self.cart_table = QTableWidget(0, 5)
        self.cart_table.setHorizontalHeaderLabels(["Categoría", "Ítem", "Cant.", "Precio", "Subtotal"])
        
        header_v = self.cart_table.horizontalHeader()
        header_v.setSectionResizeMode(0, QHeaderView.Interactive)
        header_v.setSectionResizeMode(1, QHeaderView.Stretch)
        header_v.setSectionResizeMode(2, QHeaderView.Interactive)
        header_v.setSectionResizeMode(3, QHeaderView.Interactive)
        header_v.setSectionResizeMode(4, QHeaderView.Interactive)
        self.cart_table.setColumnWidth(0, 130)
        self.cart_table.setColumnWidth(2, 60)
        self.cart_table.setColumnWidth(3, 80)
        self.cart_table.setColumnWidth(4, 90)
        
        self.cart_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.cart_table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.cart_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.cart_table.setAlternatingRowColors(True)
        self.cart_table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.cart_table.customContextMenuRequested.connect(self.show_cart_context_menu)
        cart_lay.addWidget(self.cart_table)
        
        self.btn_remove_from_cart = QPushButton("❌ Quitar Ítem (Supr)")
        
        cart_actions = QHBoxLayout()
        cart_actions.setContentsMargins(0, 5, 10, 5) 
        self.lbl_sub_medicamentos = QLabel("Medicamentos: RD$ 0.00")
        self.lbl_sub_medicamentos.setStyleSheet("font-size: 10pt; color: #555;")
        self.lbl_sub_materiales = QLabel("Materiales: RD$ 0.00")
        self.lbl_sub_materiales.setStyleSheet("font-size: 10pt; color: #555;")
        cart_actions.addWidget(self.lbl_sub_medicamentos)
        cart_actions.addSpacing(20)
        cart_actions.addWidget(self.lbl_sub_materiales)
        cart_actions.addStretch(1); cart_actions.addWidget(self.btn_remove_from_cart)
        cart_lay.addLayout(cart_actions)
        middle_v.addWidget(cart_group)

        bottom_widget = QWidget()
        bottom_widget.setObjectName("BottomBar")
        bottom_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        bottom_widget.setMinimumHeight(80)
        
        bottom = QHBoxLayout(bottom_widget)
        bottom.setContentsMargins(20, 10, 20, 10)
        
        self.lbl_edit_mode = QLabel("")
        self.lbl_edit_mode.setStyleSheet("color: #c62828; font-weight: bold; font-size: 12pt; border: none;")
        
        self.lbl_total = QLabel("Total: RD$ 0.00")
        self.lbl_total.setObjectName("TotalLabel")
        
        self.btn_generate = QPushButton("🖨️ GENERAR RECIBO PDF (F5)")
        self.btn_generate.setMinimumHeight(55)
        font = self.btn_generate.font(); font.setPointSize(12); font.setBold(True); self.btn_generate.setFont(font)
        
        self.btn_reset = QPushButton("Limpiar Todo")
        self.btn_reset.setMinimumHeight(55)
        self.btn_cancel_edit = QPushButton("↩️ Volver al inicio")
        self.btn_cancel_edit.setMinimumHeight(55)
        self.btn_cancel_edit.hide()
        self.btn_cancel_edit.setToolTip("Cancelar edición y volver al inicio")
        
        bottom.addWidget(self.lbl_edit_mode); bottom.addStretch(1)
        bottom.addWidget(self.lbl_total); bottom.addSpacing(20)
        bottom.addWidget(self.btn_cancel_edit); bottom.addWidget(self.btn_reset); bottom.addWidget(self.btn_generate)
        root.addWidget(bottom_widget)
        
        self.ars_combo.currentTextChanged.connect(self.on_ars_changed)
        self.coverage_combo.currentTextChanged.connect(self.on_coverage_changed)
        self.btn_add_catalog_item.clicked.connect(self.add_catalog_item_inline)
        self.btn_manage_catalog.clicked.connect(self.manage_current_catalog)
        self.btn_ars_mgmt.clicked.connect(self.open_ars_manager)
        self.search.textChanged.connect(self.search_and_maybe_switch_tab)
        self.btn_add.clicked.connect(self.add_selected_item)
        self.btn_remove_from_cart.clicked.connect(self.remove_current_cart_selection)
        self.btn_generate.clicked.connect(self.generate_pdf)
        self.btn_reset.clicked.connect(self.reset_all)
        self.btn_cancel_edit.clicked.connect(self.cancel_edit)
        self.btn_import_meds.clicked.connect(self.import_meds_mats_from_word)
        self.btn_receipts_history.clicked.connect(self.open_receipts_history_dialog)
        if self.btn_view_history: self.btn_view_history.clicked.connect(self.open_history_dialog)
        if self.btn_view_reports: self.btn_view_reports.clicked.connect(self.open_reports_dialog)
        if self.btn_admin_users: self.btn_admin_users.clicked.connect(self.open_users_admin)
        
        for cat, widget in self.source_lists.items():
            widget.itemDoubleClicked.connect(lambda _item, c=cat: self.add_selected_item(category_override=c))
            
        self.cart_table.itemDoubleClicked.connect(self.remove_current_cart_selection)
        QTimer.singleShot(0, self._update_responsive_ui)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        if hasattr(self, "tabs"):
            self._update_responsive_ui()

    def _update_responsive_ui(self):
        width = self.width()
        mode = "wide" if width >= 1500 else "medium" if width >= 1100 else "compact"
        if mode == self._responsive_mode:
            return
        self._responsive_mode = mode

        for index, category in enumerate(ALL_CATEGORIES):
            icon = CAT_EMOJIS.get(category, "")
            if mode == "wide":
                label = f"{icon} {category}"
            elif mode == "medium":
                label = self.compact_tab_labels.get(category, category)
            else:
                label = icon or category[:3]
            self.tabs.setTabText(index, label)
            self.tabs.setTabToolTip(index, category)

        compact = mode == "compact"
        self.lbl_modo_visual.setVisible(not compact)
        self.header_subtitle.setVisible(not compact)
        self.cart_table.setColumnHidden(3, compact)
        self.btn_import_meds.setVisible(not compact)
        self.btn_ars_mgmt.setText("ARS" if compact else "Gestión ARS")
        self.btn_add_catalog_item.setText("+ Ítem" if compact else "Nuevo ítem")
        self.btn_generate.setText("Generar PDF (F5)" if compact else "GENERAR RECIBO PDF (F5)")
        self.btn_reset.setText("Limpiar" if compact else "Limpiar todo")
        full_name = self.current_user.get("full_name", "Usuario")
        role = self.current_user.get("role", "").capitalize()
        self.lbl_user_top.setText(full_name if compact else f"{full_name} · {role}")

    def _setup_hotkeys(self):
        QShortcut(QKeySequence("F3"), self).activated.connect(self.focus_search)
        QShortcut(QKeySequence("Ctrl+F"), self).activated.connect(self.focus_search)
        self.search.returnPressed.connect(self.add_first_search_result)
        QShortcut(QKeySequence("F5"), self).activated.connect(self.generate_pdf)
        QShortcut(QKeySequence("Ctrl+P"), self).activated.connect(self.generate_pdf)
        QShortcut(QKeySequence("Del"), self.cart_table).activated.connect(self.remove_current_cart_selection)
        QShortcut(QKeySequence("Backspace"), self.cart_table).activated.connect(self.remove_current_cart_selection)

    def toggle_theme(self):
        self.is_dark_mode = not self.is_dark_mode
        self.btn_theme_toggle.setText("☀️" if self.is_dark_mode else "🌙")
        QApplication.instance().setStyleSheet(get_stylesheet(self.is_dark_mode))
        self.theme_toggled.emit(self.is_dark_mode)

    def open_preferences_dialog(self):
        username = self.current_user.get("username", "") if self.current_user else ""
        prefs = dict(self.preferences)
        dlg = PreferencesDialog(prefs, self)
        if dlg.exec() == QDialog.Accepted:
            new_prefs = dlg.values()
            self.preferences.update(new_prefs)
            upsert_user_preferences(
                username,
                auto_add_guantes=new_prefs.get("auto_add_guantes"),
                auto_print=new_prefs.get("auto_print"),
                auto_add_bajante_cateter=new_prefs.get("auto_add_bajante_cateter"),
            )
            # Handle theme change from preferences
            if new_prefs.get("theme") == "oscuro" and not self.is_dark_mode:
                self.toggle_theme()
            elif new_prefs.get("theme") == "claro" and self.is_dark_mode:
                self.toggle_theme()

    def focus_search(self):
        self.search.setFocus()
        self.search.selectAll()

    def add_first_search_result(self):
        cat = self.get_current_category()
        active_list = self.source_lists[cat]
        for i in range(active_list.count()):
            item = active_list.item(i)
            if not item.isHidden():
                active_list.setCurrentItem(item)
                self.add_selected_item()
                break

    def _build_timers(self):
        self.idle_timer = QTimer(self); self.idle_timer.timeout.connect(self.check_idle_timeout); self.idle_timer.start(60_000)
        self.remote_logout_timer = QTimer(self); self.remote_logout_timer.timeout.connect(self.check_remote_logout); self.remote_logout_timer.start(15_000)
        self.session_heartbeat_timer = QTimer(self); self.session_heartbeat_timer.timeout.connect(self.update_session_heartbeat); self.session_heartbeat_timer.start(15_000)
        self.update_session_heartbeat()

    def apply_button_colors(self):
        set_button_role(self.btn_add_catalog_item, 'success')
        set_button_role(self.btn_manage_catalog, 'info')
        set_button_role(self.btn_ars_mgmt, 'info')
        set_button_role(self.btn_add, 'success')
        set_button_role(self.btn_remove_from_cart, 'danger')
        set_button_role(self.btn_import_meds, 'warning')
        set_button_role(self.btn_generate, 'report')
        set_button_role(self.btn_reset, 'neutral')
        set_button_role(self.btn_cancel_edit, 'info')

    def mark_activity(self):
        self.last_activity = datetime.now(); self.warned_idle = False

    def get_current_category(self):
        return ALL_CATEGORIES[self.tabs.currentIndex()]

    def cart_has_ars_items(self) -> bool:
        for r in range(self.cart_table.rowCount()):
            item = self.cart_table.item(r, 0)
            cat_text = item.text() if item else ""
            cat = cat_text.split(" ", 1)[-1].strip() if " " in cat_text else cat_text.strip()
            if cat in ARS_CATEGORIES:
                return True
        return False


    def cart_has_category(self, category: str) -> bool:
        for r in range(self.cart_table.rowCount()):
            item = self.cart_table.item(r, 0)
            cat_text = item.text() if item else ""
            cat = cat_text.split(" ", 1)[-1].strip() if " " in cat_text else cat_text.strip()
            if cat == category:
                return True
        return False

    def insert_or_increment_cart_item(self, category: str, name: str, price: float, qty: int = 1) -> str:
        qty = int(qty or 1)
        price = float(price or 0.0)

        # Buscar si el item ya existe en el carrito
        for r in range(self.cart_table.rowCount()):
            if self.cart_table.item(r, 1).text() == name:
                # Sumar cantidad al item existente
                existing_qty = int(self.cart_table.item(r, 2).text())
                new_qty = existing_qty + qty
                self.cart_table.item(r, 2).setText(str(new_qty))
                new_sub = price * new_qty
                self.cart_table.item(r, 4).setText(f"${new_sub:,.2f}")
                return "incremented"
        
        # Agregar como nuevo item
        r = self.cart_table.rowCount()
        self.cart_table.insertRow(r)

        cat_item = QTableWidgetItem(f"{CAT_EMOJIS.get(category, '')} {category}")
        cat_item.setForeground(QColor(CAT_COLORS.get(category, "#333")))

        qty_item = QTableWidgetItem(str(qty))
        qty_item.setTextAlignment(Qt.AlignCenter)

        price_item = QTableWidgetItem(f"${price:,.2f}")
        price_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

        sub_item = QTableWidgetItem(f"${price * qty:,.2f}")
        sub_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)

        self.cart_table.setItem(r, 0, cat_item)
        self.cart_table.setItem(r, 1, QTableWidgetItem(name))
        self.cart_table.setItem(r, 2, qty_item)
        self.cart_table.setItem(r, 3, price_item)
        self.cart_table.setItem(r, 4, sub_item)
        return "added"

    def maybe_prompt_honorario(self, trigger_category: str):
        if trigger_category == "Honorarios":
            return
        if trigger_category not in ["Medicamentos", "Materiales", "Laboratorios", "Imágenes", "Procedimientos"]:
            return
        if not self.current_ars or self.honorario_prompt_done or self.cart_has_category("Honorarios"):
            return

        honorarios = get_ars_items("Honorarios", self.current_ars)
        if not honorarios:
            return

        self.honorario_prompt_done = True

        dlg = HonorarioSelectorDialog(honorarios, self)
        if dlg.exec() != QDialog.Accepted:
            return

        if self.locked_ars is None:
            self.locked_ars = self.current_ars

        selected = dlg.values()
        if not selected:
            return

        name, price = selected[0]
        self.insert_or_increment_cart_item("Honorarios", name, price, 1)
        self.update_totals()
        FloatingToast("✅ Honorario agregado", self).show()

    def maybe_prompt_mascarilla(self, trigger_category: str, trigger_name: str):
        if trigger_category not in ["Medicamentos", "Materiales", "Laboratorios", "Imágenes", "Procedimientos", "Honorarios"]:
            return
        needle = "nebulizar con"
        if needle not in (trigger_name or "").casefold():
            return
        if getattr(self, "mascarilla_prompt_done", False):
            return
        self.mascarilla_prompt_done = True
        items = self.universal.get("Materiales", {})
        candidates = {n: p for n, p in items.items() if ("mascarilla" in n.casefold() or "nebuliz" in n.casefold()) and "cpap" not in n.casefold()}
        if not candidates:
            return
        dlg = MascarillaSelectorDialog(self)
        dlg.load_items(candidates)
        if dlg.exec() != QDialog.Accepted or not dlg.selected_name:
            self.mascarilla_prompt_done = False
            return
        price = candidates[dlg.selected_name]
        self.insert_or_increment_cart_item("Materiales", dlg.selected_name, price, 1)
        self.update_totals()
        FloatingToast("😷 Mascarilla agregada", self).show()

    def maybe_auto_add_guantes(self, trigger_category: str):
        if trigger_category != "Medicamentos":
            return
        if not self.preferences.get("auto_add_guantes", True):
            return
        if getattr(self, "guantes_auto_added", False):
            return

        guantes_name = "Guantes"
        materiales = self.universal.get("Materiales", {})
        if guantes_name not in materiales:
            for candidate in materiales:
                if candidate.strip().casefold() == "guantes":
                    guantes_name = candidate
                    break
            else:
                return

        guantes_price = materiales[guantes_name]
        self.insert_or_increment_cart_item("Materiales", guantes_name, guantes_price, 1)
        self.guantes_auto_added = True
        self.update_totals()
        FloatingToast("🧤 Guantes agregados automáticamente", self).show()

    def _is_in_cart(self, name: str) -> bool:
        for r in range(self.cart_table.rowCount()):
            if self.cart_table.item(r, 1).text() == name:
                return True
        return False

    def maybe_auto_add_bajante_cateter(self, trigger_name: str):
        if not self.preferences.get("auto_add_bajante_cateter", True):
            return
        needle = trigger_name or ""
        needle_clean = needle.replace(" ", "").casefold()
        is_solucion = "sol." in needle.casefold() or "solucion" in needle.casefold() or "solución" in needle.casefold()
        is_paracetamol_1gr = "paracetamol" in needle.casefold() and "1gr" in needle_clean
        
        if not is_solucion and not is_paracetamol_1gr:
            return

        materiales = self.universal.get("Materiales", {})
        bajante_name = next((n for n in materiales if "bajante" in n.casefold() and ("suero" in n.casefold() or "sol." in n.casefold() or "solucion" in n.casefold() or "solución" in n.casefold())), None)
        cateter_name = next((n for n in materiales if "cateter" in n.casefold() or "catéter" in n.casefold()), None)

        added = 0
        if is_paracetamol_1gr:
            # Paracetamol 1gr: SIEMPRE agregar bajante, y catéter (si no está)
            if bajante_name:
                self.insert_or_increment_cart_item("Materiales", bajante_name, materiales[bajante_name], 1)
                added += 1
            if cateter_name and not self._is_in_cart(cateter_name):
                self.insert_or_increment_cart_item("Materiales", cateter_name, materiales[cateter_name], 1)
                added += 1
        elif is_solucion and not self.bajante_added_for_solucion:
            # Solución: Agregar bajante Y catéter la primera vez (si no están)
            if bajante_name and not self._is_in_cart(bajante_name):
                self.insert_or_increment_cart_item("Materiales", bajante_name, materiales[bajante_name], 1)
                added += 1
            if cateter_name and not self._is_in_cart(cateter_name):
                self.insert_or_increment_cart_item("Materiales", cateter_name, materiales[cateter_name], 1)
                added += 1
            if added > 0:
                self.bajante_added_for_solucion = True

        if added:
            self.bajante_cateter_auto_added = True
            self.update_totals()
            FloatingToast("🩺 Bajante de suero y catéter agregados automáticamente", self).show()

    def _fill_list(self, widget: QListWidget, items: dict, term: str) -> int:
        widget.clear()
        count = 0
        category = next((cat for cat, lst in self.source_lists.items() if lst is widget), None)
        color = CAT_COLORS.get(category, "#333333")
        term_norm = remove_accents(term)

        for name, price in sorted(items.items(), key=lambda x: x[0]):
            if term_norm and term_norm not in remove_accents(name): continue
            effective_price = get_effective_price(category, price) if category else float(price)
            
            display_text = f"{name}    ·    RD$ {effective_price:,.2f}"
            it = QListWidgetItem(display_text)
            it.setData(Qt.UserRole, (name, effective_price))
            
            font = it.font()
            font.setPointSize(self.catalog_font_size)
            font.setBold(True)
            it.setFont(font)
            it.setForeground(QColor(color))
            
            widget.addItem(it)
            count += 1
        return count

    def show_catalog_context_menu(self, pos, category):
        if not user_can_manage_catalog(self.current_user): return
        widget = self.source_lists[category]
        item = widget.itemAt(pos)
        if not item: return
        name, effective_price = item.data(Qt.UserRole)
        menu = QMenu(self)
        
        action_edit = QAction("✏️  Editar nombre o precio", self)
        action_edit.triggered.connect(lambda: self.edit_catalog_item_inline(category, name, effective_price))
        menu.addAction(action_edit)

        if user_is_admin(self.current_user):
            action_move = QMenu("➡️  Mover a otra categoría...", self)
            target_categories = ARS_CATEGORIES if category in ARS_CATEGORIES else UNIVERSAL_CATEGORIES
            for t_cat in target_categories:
                if t_cat != category:
                    sub_action = QAction(f"{CAT_EMOJIS.get(t_cat, '')}  {t_cat}", self)
                    sub_action.triggered.connect(lambda checked, target=t_cat: self.move_catalog_item_inline(category, name, target))
                    action_move.addAction(sub_action)

            action_delete = QAction("❌  Eliminar del catálogo", self)
            action_delete.triggered.connect(lambda: self.delete_catalog_item_inline(category, name))

            menu.addMenu(action_move)
            menu.addSeparator()
            menu.addAction(action_delete)
        
        menu.exec(widget.mapToGlobal(pos))

    def edit_catalog_item_inline(self, category, old_name, effective_price):
        self.mark_activity()
        dlg = AddCatalogItemDialog(
            category,
            self.current_ars if category in ARS_CATEGORIES else "",
            self,
            allow_category_navigation=False,
        )
        dlg.setWindowTitle(f"Editar {category}")
        dlg.name_edit.setText(old_name)
        dlg.price_spin.setValue(effective_price)
        
        if dlg.exec() == QDialog.Accepted:
            new_name, new_effective_price = dlg.values()
            if not new_name:
                FloatingToast("El nombre no puede estar vacío", self, is_error=True).show()
                return
            price_to_db = new_effective_price
            if category == 'Medicamentos': price_to_db = round(float(new_effective_price) / 1.20, 2)
            try:
                if category in UNIVERSAL_CATEGORIES:
                    edit_universal_item(old_name, new_name, price_to_db, category)
                    self.universal[category] = get_universal(category)
                else:
                    edit_ars_item(self.current_ars, old_name, new_name, price_to_db, category)
                    self.ars_cache[category] = get_ars_items(category, self.current_ars)
                log_action(self.current_user["username"], f"Editar ítem {category}", f"{old_name} -> {new_name}")
                self.search_and_maybe_switch_tab()
                FloatingToast("✅ Ítem actualizado", self).show()
            except psycopg2.IntegrityError:
                FloatingToast("Ya existe un ítem con ese nombre", self, is_error=True).show()
                
    def delete_catalog_item_inline(self, category, name):
        self.mark_activity()
        res = QMessageBox.question(self, "Eliminar", f"¿Seguro que deseas eliminar '{name}' de {category}?", QMessageBox.Yes | QMessageBox.No)
        if res != QMessageBox.Yes: return
        if category in UNIVERSAL_CATEGORIES:
            delete_universal(category, name)
            self.universal[category] = get_universal(category)
        else:
            delete_ars_item(category, self.current_ars, name)
            self.ars_cache[category] = get_ars_items(category, self.current_ars)
        log_action(self.current_user["username"], f"Eliminar ítem {category}", name)
        self.search_and_maybe_switch_tab()
        FloatingToast("✅ Ítem eliminado", self).show()

    def move_catalog_item_inline(self, old_category, name, new_category):
        self.mark_activity()
        try:
            if old_category in UNIVERSAL_CATEGORIES:
                move_universal_item(name, old_category, new_category)
                self.universal[old_category] = get_universal(old_category)
                self.universal[new_category] = get_universal(new_category)
            else:
                move_ars_item(self.current_ars, name, old_category, new_category)
                self.ars_cache[old_category] = get_ars_items(old_category, self.current_ars)
                self.ars_cache[new_category] = get_ars_items(new_category, self.current_ars)
            log_action(self.current_user["username"], "Mover ítem", f"{name} de {old_category} a {new_category}")
            self.search_and_maybe_switch_tab()
            FloatingToast("✅ Ítem movido", self).show()
        except psycopg2.IntegrityError:
            FloatingToast(f"Ya existe '{name}' en {new_category}", self, is_error=True).show()

    def refresh_picker(self):
        term = (self.search.text() or "").strip()
        self._fill_list(self.source_lists["Medicamentos"], self.universal["Medicamentos"], term)
        self._fill_list(self.source_lists["Materiales"], self.universal["Materiales"], term)
        for cat in ARS_CATEGORIES:
            self._fill_list(self.source_lists[cat], self.ars_cache.get(cat, {}), term)

    def search_and_maybe_switch_tab(self):
        term = (self.search.text() or "").strip()
        matches = []
        for cat in ALL_CATEGORIES:
            source_data = self.universal[cat] if cat in UNIVERSAL_CATEGORIES else self.ars_cache.get(cat, {})
            matches.append(self._fill_list(self.source_lists[cat], source_data, term))
        cats = [i for i, m in enumerate(matches) if m > 0]
        if term and len(cats) == 1 and self.tabs.currentIndex() != cats[0]:
            self.tabs.setCurrentIndex(cats[0])

    def bump_sala(self, delta: float):
        self.mark_activity(); self.sala_spin.setValue(max(0.0, self.sala_spin.value() + delta)); self.update_totals()

    def on_ars_changed(self, name):
        self.mark_activity()
        proposed_ars = name or ""
        previous_ars = self.current_ars

        # Si ya se agregó algún Laboratorio, Imagen, Procedimiento u Honorario,
        # se bloquea el cambio de ARS para evitar mezclar tarifas por error.
        if (
            self.locked_ars
            and proposed_ars != self.locked_ars
            and self.cart_has_ars_items()
            and not getattr(self, "_reverting_ars_change", False)
        ):
            FloatingToast(
                f"No puedes cambiar de ARS mientras el recibo tiene ítems de {self.locked_ars}. "
                "Quita esos ítems o limpia el recibo.",
                self,
                is_error=True
            ).show()
            self._reverting_ars_change = True
            try:
                self.ars_combo.blockSignals(True)
                self.ars_combo.setCurrentText(self.locked_ars)
                self.ars_combo.blockSignals(False)
            finally:
                self._reverting_ars_change = False
            return

        if proposed_ars != previous_ars:
            self.honorario_prompt_done = False
            self.mascarilla_prompt_done = False
            self.bajante_cateter_auto_added = False
            self.bajante_added_for_solucion = False

        self.current_ars = proposed_ars
        self.sala_spin.setValue(get_emergency_price(self.current_ars) if self.current_ars else 0.0)
        for cat in ARS_CATEGORIES:
            self.ars_cache[cat] = get_ars_items(cat, self.current_ars) if self.current_ars else {}
        self.refresh_picker()
        self.update_totals()


    def open_ars_manager(self):
        self.mark_activity(); ARSManagerDialog(self.current_user, self).exec()
        names = ars_list(); self.ars_combo.clear(); self.ars_combo.addItems(names)
        if self.current_ars not in names and names: self.ars_combo.setCurrentIndex(0)
        self.on_ars_changed(self.ars_combo.currentText())

    def save_sala(self):
        self.mark_activity()
        if not self.current_ars:
            FloatingToast("Seleccione una ARS primero", self, is_error=True).show()
            return
        set_emergency_price(self.current_ars, self.sala_spin.value())
        log_action(self.current_user["username"], "Guardar sala ARS", f"{self.current_ars} -> {self.sala_spin.value():.2f}")
        FloatingToast("✅ Precio de sala guardado", self).show()
        self.update_totals()

    def add_catalog_item_inline(self):
        self.mark_activity()
        if not user_can_manage_catalog(self.current_user):
            FloatingToast("No tienes permisos para esto", self, is_error=True).show()
            return
        category = self.get_current_category()
        if category in ARS_CATEGORIES and not self.current_ars:
            category = "Medicamentos"
        dlg = AddCatalogItemDialog(
            category,
            self.current_ars,
            self,
            allow_category_navigation=True,
        )
        if dlg.exec() == QDialog.Accepted:
            category = dlg.category
            name, price = dlg.values()
            if not name: return
            if category in UNIVERSAL_CATEGORIES:
                if category == 'Medicamentos': price = round(float(price) / 1.20, 2)
                upsert_universal(category, name, price)
            else:
                upsert_ars_item(category, self.current_ars, name, price)
            log_action(self.current_user["username"], f"Añadido/actualizado {category}", f"{name}")
            if category in UNIVERSAL_CATEGORIES: self.universal[category] = get_universal(category)
            else: self.ars_cache[category] = get_ars_items(category, self.current_ars)
            self.search_and_maybe_switch_tab()
            FloatingToast("✅ Ítem agregado", self).show()

    def manage_current_catalog(self):
        self.mark_activity()
        if not user_can_manage_catalog(self.current_user):
            FloatingToast("No tienes permisos para esto", self, is_error=True).show()
            return

        category = self.get_current_category()
        ars_name = self.current_ars or ""

        if category in ARS_CATEGORIES and not ars_name:
            FloatingToast("Seleccione una ARS primero", self, is_error=True).show()
            return

        dlg = CatalogEditorDialog(category, ars_name, self)
        dlg.exec()

        # Refresca todos los catálogos porque ahora se puede navegar y editar
        # más de una categoría desde la misma ventana.
        for cat in UNIVERSAL_CATEGORIES:
            self.universal[cat] = get_universal(cat)
        if self.current_ars:
            for cat in ARS_CATEGORIES:
                self.ars_cache[cat] = get_ars_items(cat, self.current_ars)

        if getattr(dlg, "has_changes", False):
            log_action(self.current_user["username"], "Editar catálogo", f"Ventana de catálogo navegable | ARS: {ars_name or 'N/A'}")

        self.search_and_maybe_switch_tab()

    def on_coverage_changed(self, coverage):
        """Separa la condición de cobertura de las aseguradoras reales."""
        uninsured = coverage == "No asegurado"
        self.ars_combo.setEnabled(not uninsured)
        self.btn_ars_mgmt.setEnabled(not uninsured)
        if uninsured:
            self.ars_combo.blockSignals(True)
            self.ars_combo.setCurrentIndex(-1)
            self.ars_combo.blockSignals(False)
            self.current_ars = ""
            self.locked_ars = None
            for category in ARS_CATEGORIES:
                self.ars_cache[category] = {}
            self.refresh_picker()
            self.sala_spin.setValue(0.0)
        elif self.ars_combo.count() > 0:
            self.ars_combo.setCurrentIndex(0)
            self.on_ars_changed(self.ars_combo.currentText())
        self.update_totals()

    def add_selected_item(self, category_override=None):
        self.mark_activity()
        category = category_override or self.get_current_category()
        src_list = self.source_lists[category]

        if category in ARS_CATEGORIES and not self.current_ars:
            FloatingToast("Seleccione una ARS primero", self, is_error=True).show()
            return

        # ---> LÓGICA DE BLOQUEO DE ARS <---
        if category in ARS_CATEGORIES:
            if self.locked_ars is None:
                self.locked_ars = self.current_ars
            elif self.locked_ars != self.current_ars:
                FloatingToast(f"Ya tienes ítems de {self.locked_ars}. No puedes mezclar ARS.", self, is_error=True).show()
                return

        item = src_list.currentItem()
        if not item:
            FloatingToast("Seleccione un ítem", self, is_error=True).show()
            return

        name, price = item.data(Qt.UserRole)
        qty = int(self.qty.value())

        self.insert_or_increment_cart_item(category, name, price, qty)

        self.update_totals()
        self.search.clear()
        self.search.setFocus()
        self.qty.setValue(1)

        FloatingToast("✅ Añadido al recibo", self).show()

        self.maybe_auto_add_guantes(category)
        self.maybe_auto_add_bajante_cateter(name)
        self.maybe_prompt_honorario(category)
        self.maybe_prompt_mascarilla(category, name)

    def show_cart_context_menu(self, pos):
        row = self.cart_table.rowAt(pos.y())
        if row < 0:
            return
        self.cart_table.selectRow(row)
        menu = QMenu(self)
        a_edit_qty = QAction("✏️ Modificar cantidad", self)
        a_del = QAction("❌ Quitar ítem", self)
        a_edit_qty.triggered.connect(lambda: self.edit_cart_item_qty(row))
        a_del.triggered.connect(lambda: self.remove_current_cart_selection(row))
        menu.addAction(a_edit_qty)
        menu.addAction(a_del)
        menu.exec(self.cart_table.viewport().mapToGlobal(pos))

    def edit_cart_item_qty(self, row):
        self.mark_activity()
        if row < 0 or row >= self.cart_table.rowCount():
            return
        current_qty = int(self.cart_table.item(row, 2).text())
        dlg = QtyDialog(current_qty, self)
        if dlg.exec() != QDialog.Accepted:
            return
        new_qty = dlg.spin.value()
        price = float(self.cart_table.item(row, 3).text().replace('$', '').replace(',', ''))
        
        # Actualizar directamente la cantidad sin eliminar la fila
        qty_item = QTableWidgetItem(str(new_qty))
        qty_item.setTextAlignment(Qt.AlignCenter)
        self.cart_table.setItem(row, 2, qty_item)
        
        # Actualizar subtotal
        sub_item = QTableWidgetItem(f"${price * new_qty:,.2f}")
        sub_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
        self.cart_table.setItem(row, 4, sub_item)
        
        self.update_totals()
        FloatingToast(f"✅ Cantidad actualizada a {new_qty}", self).show()

    def remove_current_cart_selection(self, _item=None):
        self.mark_activity()
        row = self.cart_table.currentRow()
        if row >= 0:
            self.cart_table.removeRow(row)
            self.update_totals()
            
            # ---> LIBERAR ARS SI YA NO QUEDAN ÍTEMS DE ARS <---
            if not self.cart_has_ars_items():
                self.locked_ars = None
            FloatingToast("🗑️ Ítem removido", self).show()

    def update_totals(self):
        total = 0.0
        medicamentos_sub = 0.0
        materiales_sub = 0.0
        
        for r in range(self.cart_table.rowCount()):
            cat = self.cart_table.item(r, 0).text()
            sub = float(self.cart_table.item(r, 4).text().replace('$', '').replace(',', ''))
            total += sub
            if "Medicamentos" in cat:
                medicamentos_sub += sub
            elif "Materiales" in cat:
                materiales_sub += sub
        
        total += self.sala_spin.value()
        
        txt = f"Total: RD$ {total:,.2f}"
        self.lbl_total.setText(txt)
        self.lbl_sub_medicamentos.setText(f"Medicamentos: RD$ {medicamentos_sub:,.2f}")
        self.lbl_sub_materiales.setText(f"Materiales: RD$ {materiales_sub:,.2f}")

    def reset_all(self):
        self.mark_activity()
        self.search.clear()
        self.qty.setValue(1)
        self.locked_ars = None # LIBERAR ARS AL LIMPIAR
        self.honorario_prompt_done = False
        self.guantes_auto_added = False
        self.bajante_cateter_auto_added = False
        self.bajante_added_for_solucion = False
        self.mascarilla_prompt_done = False
        self.editing_recibo_id = None
        self.editing_recibo_numero = None
        self.lbl_edit_mode.setText("")
        self.btn_generate.setText("GENERAR RECIBO PDF (F5)")
        self.btn_cancel_edit.hide()
        self.cart_table.setRowCount(0)
        self.lbl_total.setText("Total: RD$ 0.00")
        self.lbl_sub_medicamentos.setText("Medicamentos: RD$ 0.00")
        self.lbl_sub_materiales.setText("Materiales: RD$ 0.00")
        self.name_edit.clear()
        self.dx_edit.clear()
        self.date_edit.setDate(QDate.currentDate())
        self.coverage_combo.setCurrentText("Asegurado")
        self._responsive_mode = None
        self._update_responsive_ui()

    def cancel_edit(self):
        self.reset_all()
        FloatingToast("ℹ️ Edición cancelada", self).show()

    def load_recibo_for_editing(self, recibo_id: int):
        self.reset_all()
        self.btn_cancel_edit.show()
        data = get_recibo_data(recibo_id)
        self.editing_recibo_id = recibo_id
        self.editing_recibo_numero = data["numero"]

        self.lbl_edit_mode.setText(f"✏️ EDICIÓN: RECIBO N° {data['numero']}")
        self.btn_generate.setText("💾 GUARDAR CAMBIOS PDF (F5)")

        self.name_edit.setText(data["nombre"])
        self.dx_edit.setText(data["dx"])
        try:
            self.date_edit.setDate(QDate.fromString(data["fecha"], "yyyy-MM-dd"))
        except Exception:
            pass

        coverage = data.get("tipo_cobertura") or ("NO_ASEGURADO" if not data.get("ars") else "ASEGURADO")
        self.coverage_combo.setCurrentText("No asegurado" if coverage == "NO_ASEGURADO" else "Asegurado")
        ars_name = data["ars"]
        if ars_name in [self.ars_combo.itemText(i) for i in range(self.ars_combo.count())]:
            self.ars_combo.setCurrentText(ars_name)
        else:
            self.ars_combo.setCurrentIndex(0)

        self.sala_spin.setValue(float(data["sala"]))

        for item in data["items"]:
            r = self.cart_table.rowCount()
            self.cart_table.insertRow(r)
            cat = item["categoria"]
            
            cat_item = QTableWidgetItem(f"{CAT_EMOJIS.get(cat, '')} {cat}")
            cat_item.setForeground(QColor(CAT_COLORS.get(cat, "#333")))
            
            qty_item = QTableWidgetItem(str(item["cantidad"]))
            qty_item.setTextAlignment(Qt.AlignCenter)
            
            price_item = QTableWidgetItem(f"${float(item['precio_unit']):,.2f}")
            price_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            
            sub_item = QTableWidgetItem(f"${float(item['total']):,.2f}")
            sub_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            
            self.cart_table.setItem(r, 0, cat_item)
            self.cart_table.setItem(r, 1, QTableWidgetItem(item["nombre"]))
            self.cart_table.setItem(r, 2, qty_item)
            self.cart_table.setItem(r, 3, price_item)
            self.cart_table.setItem(r, 4, sub_item)

        self.locked_ars = ars_name if self.cart_has_ars_items() else None
        self.update_totals()

    def generate_pdf(self):
        self.mark_activity()
        patient = self.name_edit.text().strip()
        if not patient:
            FloatingToast("❌ El nombre del paciente es obligatorio", self, is_error=True).show()
            self.name_edit.setFocus()
            return
        now = datetime.now().timestamp()
        if now - self._last_generate_time < 2.0:
            FloatingToast("⏳ Espera un momento para generar otro recibo", self, is_error=True).show()
            return
        self._last_generate_time = now
        date_str = self.date_edit.date().toString("yyyy-MM-dd")
        dx_raw = self.dx_edit.text().strip()

        today = QDate.currentDate()
        selected_date = self.date_edit.date()
        days_past = selected_date.daysTo(today)

        if days_past < 0:
            FloatingToast("❌ No puedes crear recibos con fechas futuras", self, is_error=True).show()
            return

        is_backdated = 0
        if days_past > 2:
            is_backdated = 1
            if self.current_user.get('role') not in (ROLE_ADMIN, ROLE_AUDIT):
                FloatingToast("❌ Fechas mayores a 48h requieren rol Administrador o Facturador de Auditoría", self, is_error=True).show()
                return

        grouped_dict = {cat: [] for cat in ALL_CATEGORIES}
        for r in range(self.cart_table.rowCount()):
            cat_raw = self.cart_table.item(r, 0).text()
            cat = cat_raw.split(" ")[-1] if " " in cat_raw else cat_raw 
            if cat not in ALL_CATEGORIES: cat = "Procedimientos"
            name = self.cart_table.item(r, 1).text()
            qty = int(self.cart_table.item(r, 2).text())
            pu = float(self.cart_table.item(r, 3).text().replace('$', '').replace(',', ''))
            sub = float(self.cart_table.item(r, 4).text().replace('$', '').replace(',', ''))
            grouped_dict[cat].append((name, pu, qty, sub, cat))
            
        grouped = [(c, grouped_dict[c]) for c in ALL_CATEGORIES if grouped_dict[c]]

        sala = self.sala_spin.value()
        subtotales = {label: sum(sub for _, _, _, sub, _ in lst) for label, lst in grouped}
        total_general = sum(subtotales.values()) + sala

        if total_general <= 0:
            FloatingToast("No hay ítems para facturar", self, is_error=True).show()
            return

        self.btn_generate.setEnabled(False)
        self.btn_generate.setText("⏳ Generando y guardando...")

        self.pdf_worker.submit({
            "patient": patient,
            "date_str": date_str,
            "dx_raw": dx_raw,
            "ars_name": self.current_ars,
            "coverage": "NO_ASEGURADO" if self.coverage_combo.currentText() == "No asegurado" else "ASEGURADO",
            "sala": sala,
            "grouped": grouped,
            "total_general": total_general,
            "editing_id": self.editing_recibo_id,
            "editing_num": self.editing_recibo_numero,
            "current_user": dict(self.current_user),
            "is_backdated": is_backdated,
        })

    def _send_ctrl_p(self):
        try:
            import time
            time.sleep(0.6)
            import ctypes
            ctypes.windll.user32.keybd_event(0x11, 0, 0, 0)
            ctypes.windll.user32.keybd_event(0x50, 0, 0, 0)
            ctypes.windll.user32.keybd_event(0x50, 0, 2, 0)
            ctypes.windll.user32.keybd_event(0x11, 0, 2, 0)
        except Exception as e:
            print(f"[PDF] Error al enviar Ctrl+P: {e}")

    def _open_pdf_after_generation(self, pdf_path: str, auto_print: bool = False):
        try:
            if not os.path.exists(pdf_path):
                print(f"[PDF] Archivo no encontrado: {pdf_path}")
                return
            if sys.platform == "win32":
                try:
                    os.startfile(pdf_path)
                except Exception as e:
                    print(f"[PDF] os.startfile falló: {e}")
                    try:
                        startupinfo = subprocess.STARTUPINFO()
                        startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                        subprocess.Popen(["cmd", "/c", "start", "", pdf_path], startupinfo=startupinfo)
                    except Exception as e2:
                        print(f"[PDF] Error al abrir PDF (fallback): {e2}")
                        return
                if auto_print:
                    QTimer.singleShot(800, self._send_ctrl_p)
                    FloatingToast("🖨️ Abriendo panel de impresión...", self).show()
            else:
                open_file_path(pdf_path)
        except Exception as e:
            print(f"[PDF] Error al abrir/Imprimir PDF: {e}")
            try:
                import webbrowser
                webbrowser.open(f"file:///{pdf_path}")
            except Exception as e2:
                print(f"[PDF] Error al abrir PDF (fallback): {e2}")

    def on_pdf_generated(self, success, message, pdf_path, recibo_number):
        selected_date_before_reset = self.date_edit.date()

        self.btn_generate.setEnabled(True)
        self.btn_generate.setText("🖨️ GENERAR RECIBO PDF (F5)")
        
        if success:
            self.reset_all()
            self.date_edit.setDate(selected_date_before_reset)
            FloatingToast(f"✅ Recibo N° {recibo_number} guardado con éxito", self).show()
            auto_print = bool(self.preferences.get("auto_print"))
            self._open_pdf_after_generation(pdf_path, auto_print=auto_print)
        else:
            QMessageBox.critical(self, "Error", f"Fallo al generar el recibo:\n{message}")

    def import_meds_mats_from_word(self):
        self.mark_activity()
        if Document is None:
            QMessageBox.warning(self, "Importar Word", "Falta dependencia: python-docx.\nInstala con:\n\npip install python-docx")
            return
        current_cat = self.get_current_category()
        path, _ = QFileDialog.getOpenFileName(self, "Seleccionar archivo Word", "", "Word (*.docx)")
        if not path: return
        try:
            if current_cat in UNIVERSAL_CATEGORIES:
                summary = import_word_to_universal_catalog(path)
                self.universal["Medicamentos"] = get_universal("Medicamentos")
                self.universal["Materiales"] = get_universal("Materiales")
                log_action(self.current_user['username'], "Importar Word Universal", f"Archivo: {os.path.basename(path)}")
                QMessageBox.information(
                    self, "Importación completada",
                    (f"Medicamentos — Actualizados: {summary['Medicamentos']['updated']} | Nuevos: {summary['Medicamentos']['inserted']}\n"
                     f"Materiales   — Actualizados: {summary['Materiales']['updated']} | Nuevos: {summary['Materiales']['inserted']}\n\n"
                     f"Regla aplicada: Medicamentos guardados base en BD y mostrados/calculados con +20% en GUI.")
                )
            else:
                if not self.current_ars:
                    FloatingToast("Selecciona una ARS primero", self, is_error=True).show()
                    return
                summary = import_word_to_ars_catalog(path, self.current_ars)
                for cat in ARS_CATEGORIES:
                    self.ars_cache[cat] = get_ars_items(cat, self.current_ars)
                log_action(self.current_user['username'], "Importar Word ARS", f"ARS: {self.current_ars} | Archivo: {os.path.basename(path)}")
                QMessageBox.information(
                    self, "Importación completada",
                    (f"ARS: {self.current_ars}\n\n"
                     f"Laboratorios — Actualizados: {summary['Laboratorios']['updated']} | Nuevos: {summary['Laboratorios']['inserted']}\n"
                     f"Imágenes     — Actualizados: {summary['Imágenes']['updated']} | Nuevos: {summary['Imágenes']['inserted']}\n"
                     f"Procedimientos — Actualizados: {summary['Procedimientos']['updated']} | Nuevos: {summary['Procedimientos']['inserted']}\n"
                     f"Honorarios     — Actualizados: {summary['Honorarios']['updated']} | Nuevos: {summary['Honorarios']['inserted']}")
                )
            self.refresh_picker()
            self.search_and_maybe_switch_tab()
        except Exception as e:
            QMessageBox.critical(self, "Importar Word", f"Error al procesar el archivo:\n{e}")

    def open_history_dialog(self):
        self.mark_activity()
        HistoryDialog(self).exec()

    def open_reports_dialog(self):
        self.mark_activity()
        if not self.current_user:
            FloatingToast("Debes iniciar sesión para acceder a reportes", self, is_error=True).show()
            return
        ReportsDialog(self.current_user, self).exec()

    def open_receipts_history_dialog(self):
        self.mark_activity()
        ReceiptHistoryDialog(main_window=self, parent=self).exec()

    def open_trash_dialog(self):
        self.mark_activity()
        ReceiptTrashDialog(main_window=self, parent=self).exec()
        self._purge_old_deleted_receipts()

    def _purge_old_deleted_receipts(self):
        try:
            purged = purge_old_deleted_receipts(30)
            if purged:
                FloatingToast(f"🧹 {len(purged)} recibo(s) antiguos purgados (30 días)", self).show()
        except Exception as e:
            print(f"[PAPELERA] Error en purga automática: {e}")

    def open_users_admin(self):
        self.mark_activity()
        if not user_can_manage_sessions(self.current_user):
            FloatingToast("No tienes permisos para esta sección", self, is_error=True).show()
            return
        UsersAdminDialog(self.current_user, self).exec()

    def check_idle_timeout(self):
        elapsed = datetime.now() - self.last_activity
        minutes = elapsed.total_seconds() / 60.0
        if minutes >= (IDLE_TIMEOUT_MINUTES - WARNING_BEFORE_TIMEOUT_MINUTES) and not self.warned_idle:
            self.warned_idle = True
            FloatingToast("La sesión se cerrará por inactividad pronto.", self, is_error=True).show()
        if minutes >= IDLE_TIMEOUT_MINUTES:
            self.force_logout(reason="Sesión cerrada por inactividad")

    def update_session_heartbeat(self):
        try:
            username = self.current_user.get('username') if isinstance(self.current_user, dict) else ""
            heartbeat_active_session(username, getattr(self, "session_id", ""))
        except Exception as e:
            print(f"[SESIONES] Error actualizando sesión activa: {e}")

    def check_remote_logout(self):
        try:
            username = self.current_user.get('username') if isinstance(self.current_user, dict) else ""
            signal = get_remote_logout_signal(username)
            force_at = (signal or {}).get("force_logout_at") if signal else ""
            if force_at and str(force_at) > str(self.session_started_at):
                self.force_logout(reason="Sesión cerrada")
        except Exception as e:
            print(f"[SESIONES] Error verificando cierre remoto: {e}")

    def force_logout(self, reason="Sesión cerrada"):
        if hasattr(self, 'idle_timer'):
            self.idle_timer.stop()
        if hasattr(self, 'remote_logout_timer'):
            self.remote_logout_timer.stop()
        if hasattr(self, 'session_heartbeat_timer'):
            self.session_heartbeat_timer.stop()

        try:
            end_active_session(
                self.current_user.get('username') if isinstance(self.current_user, dict) else "",
                getattr(self, "session_id", "")
            )
        except Exception as e:
            print(f"[SESIONES] Error cerrando sesión activa: {e}")

        if isinstance(reason, bool) or reason is None:
            safe_reason = "Sesión cerrada"
        else:
            safe_reason = str(reason).strip() or "Sesión cerrada"
            if safe_reason in ("False", "True", "Cierre de sesión manual", "Cierre de sesión"):
                safe_reason = "Sesión cerrada"

        log_action(self.current_user.get('username') if isinstance(self.current_user, dict) else "Sistema", "Cerrar sesión", safe_reason)
        QMessageBox.information(self, "Sesión", f"{safe_reason}.")
        self.logout_requested.emit()
        self.close()
        
    def closeEvent(self, event):
        if hasattr(self, 'idle_timer'):
            self.idle_timer.stop()
        if hasattr(self, 'remote_logout_timer'):
            self.remote_logout_timer.stop()
        if hasattr(self, 'session_heartbeat_timer'):
            self.session_heartbeat_timer.stop()
        if hasattr(self, 'pdf_sync_worker'):
            self.pdf_sync_worker.stop()
            self.pdf_sync_worker.wait(3000)
        if hasattr(self, 'pdf_worker') and self.pdf_worker.is_alive():
            self.pdf_worker.stop()
            self.pdf_worker.join(timeout=5.0)
        try:
            end_active_session(
                self.current_user.get('username') if isinstance(self.current_user, dict) else "",
                getattr(self, "session_id", "")
            )
        except Exception:
            pass
        super().closeEvent(event)

    def safe_startup_load(self):
        try:
            generated, errors = safe_generate_pending_reports(self.current_user['username'])
            if errors: FloatingToast("Hubo errores leves al sincronizar", self, is_error=True).show()
        except Exception as e:
            print(f"[INICIO] Error controlado al cargar reportes pendientes: {e}")
            try: log_action(self.current_user['username'], "Error de arranque", f"Carga inicial de reportes: {e}")
            except Exception: pass

class AppController(QObject):
    def __init__(self, app: QApplication):
        super().__init__()
        self.app = app
        self.main_window = None
        self.idle_filter = IdleEventFilter()
        self.idle_filter.activity.connect(self._on_activity)
        self.app.installEventFilter(self.idle_filter)
        
        self.app.setStyleSheet(get_stylesheet(False))

    def run(self):
        while True:
            dlg = LoginDialog()
            if dlg.exec() != QDialog.Accepted: return 0
            user = dlg.user
            self.main_window = MainWindow(user)
            self.main_window.theme_toggled.connect(self._on_theme_toggled)
            self.main_window.logout_requested.connect(self._on_logout)
            self.main_window.showMaximized() 
            result = self.app.exec()
            if self.main_window is None: continue
            return result


    def _on_theme_toggled(self, is_dark):
        self.app.setStyleSheet(get_stylesheet(is_dark))

    def _on_activity(self):
        if self.main_window is not None: self.main_window.mark_activity()

    def _on_logout(self):
        self.main_window = None
        self.app.quit()

# =========================================================
# LIMPIEZA AUTOMÁTICA DE VERSIONES VIEJAS
# =========================================================
def cleanup_old_updates():
    """Busca y elimina versiones antiguas del ejecutable trz una actualización."""
    if getattr(sys, 'frozen', False) and sys.executable.endswith('.exe'):
        import os
        current_exe = sys.executable
        exe_dir = os.path.dirname(current_exe)
        current_name = os.path.basename(current_exe).lower()
        
        # Escanea la carpeta buscando basura del pasado
        for file in os.listdir(exe_dir):
            file_lower = file.lower()
            
            # Si el archivo termina en .old, O es un .exe descargado temporal que se atascó
            if file_lower.endswith(".old") or (file_lower.endswith(".exe") and file_lower != current_name and "actualizado" in file_lower):
                try:
                    # Lo borra del disco duro permanentemente
                    os.remove(os.path.join(exe_dir, file))
                except Exception:
                    pass # Si está bloqueado, lo ignora y lo borrará la próxima vez

class PreferencesDialog(QDialog):
    def __init__(self, current_prefs: dict, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Preferencias del Usuario")
        self.setFixedSize(460, 260)
        self.parent_ref = parent

        root = QVBoxLayout(self)

        self.chk_guantes = QCheckBox("Autoagregar guantes al agregar medicamentos")
        self.chk_guantes.setChecked(bool(current_prefs.get("auto_add_guantes", True)))
        note_guantes = QLabel("Nota: se agrega automáticamente el item 'Guantes' de Materiales, solo una vez por recibo.")
        note_guantes.setWordWrap(True)
        note_guantes.setStyleSheet("color: #666666; font-size: 9pt;")

        self.chk_auto_print = QCheckBox("Intentar imprimir automáticamente al generar el PDF")
        self.chk_auto_print.setChecked(bool(current_prefs.get("auto_print", False)))
        note_print = QLabel("Nota: abre el diálogo de impresión, solo presiona Enter para confirmar.")
        note_print.setWordWrap(True)
        note_print.setStyleSheet("color: #666666; font-size: 9pt;")

        self.chk_bajante_cateter = QCheckBox("Autoagregar bajante y catéter al agregar soluciones")
        self.chk_bajante_cateter.setChecked(bool(current_prefs.get("auto_add_bajante_cateter", True)))
        note_bc = QLabel("Nota: al agregar ítems con 'sol.' o 'solución', se agregan bajante y catéter de Materiales.")
        note_bc.setWordWrap(True)
        note_bc.setStyleSheet("color: #666666; font-size: 9pt;")

        self.chk_theme_dark = QCheckBox("Usar tema oscuro")
        self.chk_theme_dark.setChecked(bool(current_prefs.get("theme") == "oscuro"))

        root.addWidget(self.chk_guantes)
        root.addWidget(note_guantes)
        root.addSpacing(12)
        root.addWidget(self.chk_auto_print)
        root.addWidget(note_print)
        root.addSpacing(12)
        root.addWidget(self.chk_bajante_cateter)
        root.addWidget(note_bc)
        root.addSpacing(12)
        root.addWidget(self.chk_theme_dark)
        root.addStretch(1)

        btns = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        root.addWidget(btns)
        for btn in btns.buttons():
            set_button_role(btn, "success" if btn is btns.button(QDialogButtonBox.Save) else "neutral")

    def values(self) -> dict:
        return {
            "auto_add_guantes": self.chk_guantes.isChecked(),
            "auto_print": self.chk_auto_print.isChecked(),
            "auto_add_bajante_cateter": self.chk_bajante_cateter.isChecked(),
            "theme": "oscuro" if self.chk_theme_dark.isChecked() else "claro",
        }


def run_pdf_self_test(output_path: str) -> int:
    """Prueba el motor y los recursos embebidos sin conectarse a la BD."""
    try:
        from pdf_engine import ReceiptPDFRenderer

        test_data = {
            "numero": "SELF-TEST",
            "fecha": datetime.now().strftime("%Y-%m-%d"),
            "paciente": "PRUEBA DE EMPAQUETADO",
            "diagnostico": "Verificación del motor PDF",
            "ars": "N/A",
            "sala": 100.0,
            "categorias": [
                {
                    "nombre": "Medicamentos",
                    "items": [
                        {
                            "descripcion": "Elemento de prueba",
                            "cantidad": 1,
                            "precio": 100.0,
                            "total": 100.0,
                        }
                    ],
                }
            ],
            "total_general": 200.0,
            "total_letras": "DOSCIENTOS PESOS",
            "usuario": "Autodiagnóstico",
            "logo_path": LOGO_PATH,
        }
        renderer = ReceiptPDFRenderer(persistent=True)
        try:
            renderer.render_pdf(
                test_data,
                os.path.abspath(output_path),
                save_html_preview=False,
            )
        finally:
            renderer.close()
        return 0
    except Exception as exc:
        write_runtime_log(f"Autodiagnóstico PDF falló: {exc}")
        return 1


def run_report_exports_self_test(output_dir: str) -> int:
    """Valida los recursos de reportes, Excel y Playwright del paquete onedir."""
    try:
        output_dir = os.path.abspath(output_dir)
        os.makedirs(output_dir, exist_ok=True)
        categories = [
            {"label": "Medicamentos", "receipts": 18, "quantity": 45, "total": 18000.0, "percentage": 0.60},
            {"label": "Materiales", "receipts": 12, "quantity": 30, "total": 8000.0, "percentage": 0.2667},
            {"label": "Honorarios", "receipts": 4, "quantity": 4, "total": 4000.0, "percentage": 0.1333},
        ]
        trend = [
            {"label": "2026-07-09", "receipts": 7, "total": 7000.0, "average": 1000.0},
            {"label": "2026-07-10", "receipts": 10, "total": 10000.0, "average": 1000.0},
            {"label": "2026-07-11", "receipts": 13, "total": 13000.0, "average": 1000.0},
        ]
        for row in categories:
            row["average"] = row["total"] / row["receipts"]
        comparison = [
            {"label": "SENASA", "receipts": 18, "total": 18000.0, "average": 1000.0, "money_percentage": .60, "receipt_percentage": .60},
            {"label": "HUMANO", "receipts": 8, "total": 8000.0, "average": 1000.0, "money_percentage": .2667, "receipt_percentage": .2667},
            {"label": "PRIMERA ARS", "receipts": 4, "total": 4000.0, "average": 1000.0, "money_percentage": .1333, "receipt_percentage": .1333},
        ]
        data = {
            "start_date": "2026-05-01",
            "end_date": "2026-07-31",
            "filters": {
                "ars": "Todas las ARS", "user": "Todos los Usuarios",
                "medication": "Todos los medicamentos", "category": "Todas las categorías",
            },
            "summary": {
                "receipts": 30,
                "total": 30000.0,
                "average": 1000.0,
                "room": 2500.0,
                "medications_registered": 187,
                "top_medication": "Elemento de prueba",
                "top_medication_quantity": 22,
            },
            "previous": {"receipts": 24, "total": 24000.0},
            "trend": trend,
            "monthly": trend,
            "categories": categories,
            "category_distribution": categories,
            "users": [{"label": "Autodiagnóstico", "receipts": 30, "quantity": 79, "total": 30000.0}],
            "ars": comparison,
            "ars_breakdown": comparison,
            "show_ars_comparison": True,
            "bar": comparison,
            "comparison": comparison,
            "breakdown": comparison,
            "breakdown_type": "ars",
            "summary_table": [{"type": "ars", **row} for row in comparison],
            "view": {
                "show_ars_comparison": True, "ars_metric": "total",
                "ars_metric_label": "Total emitido", "evolution_metric": "total",
                "evolution_label": "Total emitido",
            },
            "details": [{
                "receipt": 1,
                "created_at": "2026-07-11 12:00:00",
                "username": "Autodiagnóstico",
                "ars": "N/A",
                "category": "Medicamentos",
                "item": "Elemento de prueba",
                "quantity": 1,
                "unit_price": 1000.0,
                "total": 1000.0,
            }],
        }
        renderer = ReportHTMLRenderer()
        renderer.render_pdf({
            "mode": "panel",
            "title": "PANEL DE REPORTES Y GRÁFICOS",
            "subtitle": f"Período analizado: {data['start_date']} al {data['end_date']}",
            "generated_by": "Autodiagnóstico",
            "logo_path": LOGO_PATH,
            "data": data,
            "ars_metric": "total",
            "evolution_metric": "total",
        },
            os.path.join(output_dir, "panel_autodiagnostico.pdf"),
            landscape=True,
        )
        renderer.render_pdf({
            "mode": "standard",
            "title": "REPORTE DE AUTODIAGNÓSTICO",
            "subtitle": f"Período por fecha de generación: {data['start_date']} al {data['end_date']}",
            "generated_by": "Autodiagnóstico",
            "logo_path": LOGO_PATH,
            "totals": {"_total_recibos": 30, "Total General": 30000.0, "Sala Emergencia": 2500.0},
            "category_rows": [(row["label"], row["total"]) for row in categories],
            "ars_rows": [("N/A", 30)],
            "user_rows": [("Autodiagnóstico", 30)],
        },
            os.path.join(output_dir, "reporte_autodiagnostico.pdf"),
            landscape=False,
        )
        export_panel_xlsx(
            data,
            os.path.join(output_dir, "panel_autodiagnostico.xlsx"),
            "Autodiagnóstico",
            LOGO_PATH,
        )
        specific_data = copy.deepcopy(data)
        specific_data["filters"]["ars"] = "SENASA"
        specific_data["show_ars_comparison"] = False
        specific_data["ars"] = []
        specific_data["ars_breakdown"] = []
        specific_data["comparison"] = []
        specific_data["summary_table"] = []
        specific_data["view"]["show_ars_comparison"] = False
        renderer.render_pdf({
            "mode": "panel",
            "title": "PANEL DE REPORTES Y GRÁFICOS",
            "subtitle": f"Período analizado: {data['start_date']} al {data['end_date']}",
            "generated_by": "Autodiagnóstico",
            "logo_path": LOGO_PATH,
            "data": specific_data,
            "ars_metric": "total",
            "evolution_metric": "total",
        }, os.path.join(output_dir, "panel_ars_especifica.pdf"), landscape=True)
        export_panel_xlsx(
            specific_data,
            os.path.join(output_dir, "panel_ars_especifica.xlsx"),
            "Autodiagnóstico",
            LOGO_PATH,
        )
        return 0
    except Exception as exc:
        write_runtime_log(f"Autodiagnóstico de reportes falló: {exc}")
        return 1


def main():
    cleanup_old_updates()

    app = QApplication.instance() or QApplication(sys.argv)
    if ICON_PATH:
        app.setWindowIcon(QIcon(ICON_PATH))

    if not psycopg2 or DB_URL == "AQUI_TU_URL_DE_POSTGRESQL":
        QMessageBox.critical(None, "Error de Configuración",
            "Para conectar este sistema a la nube necesitas:\n\n"
            "1. Instalar las librerías: pip install psycopg2-binary python-dotenv\n"
            "2. Colocar tu enlace de base de datos en un archivo .env\n\n"
        )
        return 1

    try:
        db_init()
    except Exception as exc:
        QMessageBox.critical(
            None,
            "Error de conexión",
            "No se pudo inicializar la base de datos.\n\n"
            f"Detalle: {exc}",
        )
        return 1

    controller = AppController(app)
    return int(controller.run())

if __name__ == "__main__":
    if len(sys.argv) == 3 and sys.argv[1] == "--self-test-pdf":
        raise SystemExit(run_pdf_self_test(sys.argv[2]))
    if len(sys.argv) == 3 and sys.argv[1] == "--self-test-reports":
        raise SystemExit(run_report_exports_self_test(sys.argv[2]))
    raise SystemExit(main())
